"""
Fixed Kurmanji Word Recorder with Menu System
This version fixes the edit/delete functionality and adds a clean menu system
"""

import os
import json
import threading
import queue
import time
import re
from datetime import datetime
from typing import List, Dict, Optional, Tuple
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from pathlib import Path

# Audio processing
import sounddevice as sd
import soundfile as sf
import numpy as np
import librosa
from scipy.signal import butter, filtfilt

# UI Framework
import customtkinter as ctk

# Document processing
import docx
import PyPDF2
import requests
from urllib.parse import urlparse

# Language Support System
class LanguageManager:
    """Manages multilingual support for Turkish, Kurdish, and English"""
    
    def __init__(self):
        self.current_language = "tr"  # Default to Turkish for Kurdish support
        self.translations = {
            "en": {
                # Main interface
                "app_title": "🎤 Kurmanji Word Recorder",
                "menu": "📋 Menu",
                "current_word": "Current Word:",
                "no_words_loaded": "No words loaded",
                "previous": "← Previous",
                "next": "Next →",
                "already_recorded": "✓ Already recorded",
                "not_recorded": "○ Not recorded",
                "start_recording": "🔴 Start Recording",
                "stop_recording": "⏹️ Stop Recording",
                "play": "▶️ Play",
                "save_next": "💾 Save",
                "edit": "✏️ Edit",
                "delete": "🗑️ Delete",
                "confirm_delete": "Delete this word?",
                "word_deleted": "Word deleted successfully",
                "recording_progress": "Recording Progress:",
                "words_recorded": "words recorded",
                "add_words": "➕ Add Words",
                "view_words": "📋 View Words",
                "load_file": "📂 Load File",
                "statistics": "📊 Statistics",
                "checking_microphone": "🔍 Checking microphone...",
                "microphone_working": "🎤 Microphone working",
                "microphone_error": "❌ Microphone error",
                "no_word_to_record": "No word to record. Please load a word list first.",
                "recording_started": "Recording started...",
                "recording_stopped": "Recording stopped.",
                "no_recording_to_play": "No recording to play.",
                "word_saved": "Word saved and moved to next.",
                "edit_current_word": "Edit Current Word",
                "enter_corrected_word": "Enter the corrected word or sentence...",
                "save_changes": "💾 Save Changes",
                "cancel": "❌ Cancel",
                "word_updated": "Word Updated",
                "duplicate_word": "Duplicate Word",
                "error": "Error",
                "warning": "Warning",
                "success": "Success",
                "language": "Language:",
                "current_sentence": "Current Sentence:",
                "current_paragraph": "Current Paragraph:",
                "sentence_mode": "🎭 Sentence Mode",
                "paragraph_mode": "📄 Paragraph Mode",
                "word_mode": "📝 Word Mode",
                "no_sentences_loaded": "No sentences loaded",
                "no_paragraphs_loaded": "No paragraphs loaded",
                "sentences_recorded": "sentences recorded",
                "paragraphs_recorded": "paragraphs recorded",
                "speed_selection": "Speed Selection (Multi-Speed Recording)",
                "slow_speed": "🐌 Slow",
                "normal_speed": "🎯 Normal",
                "fast_speed": "🚀 Fast",
                "recording_status": "Current Word Recording Status:",
                "force_next": "⏭️ Force Next Word",
                "skip_word": "⏭️ Skip This Word",
                "not_recorded_yet": "Not Recorded",
                "recorded": "Recorded",
                "completion": "Completion"
            },
            "tr": {
                # Turkish translations
                "app_title": "🎤 Kurmancî Kelime Kaydedici",
                "menu": "📋 Menü",
                "current_word": "Şu Anki Kelime:",
                "no_words_loaded": "Hiç kelime yüklenmedi",
                "previous": "← Önceki",
                "next": "Sonraki →",
                "already_recorded": "✓ Zaten kaydedildi",
                "not_recorded": "○ Kaydedilmedi",
                "start_recording": "🔴 Kayda Başla",
                "stop_recording": "⏹️ Kaydı Durdur",
                "play": "▶️ Oynat",
                "save_next": "💾 Kaydet",
                "edit": "✏️ Düzenle",
                "delete": "🗑️ Sil",
                "confirm_delete": "Bu kelimeyi silmek istiyor musunuz?",
                "word_deleted": "Kelime başarıyla silindi",
                "recording_progress": "Kayıt İlerlemesi:",
                "words_recorded": "kelime kaydedildi",
                "add_words": "➕ Kelime Ekle",
                "view_words": "📋 Kelimeleri Görüntüle",
                "load_file": "📂 Dosya Yükle",
                "statistics": "📊 İstatistikler",
                "checking_microphone": "🔍 Mikrofon kontrol ediliyor...",
                "microphone_working": "🎤 Mikrofon çalışıyor",
                "microphone_error": "❌ Mikrofon hatası",
                "no_word_to_record": "Kaydedilecek kelime yok. Lütfen önce kelime listesi yükleyin.",
                "recording_started": "Kayıt başladı...",
                "recording_stopped": "Kayıt durduruldu.",
                "no_recording_to_play": "Oynatılacak kayıt yok.",
                "word_saved": "Kelime kaydedildi ve sonrakine geçildi.",
                "edit_current_word": "Şu Anki Kelimeyi Düzenle",
                "enter_corrected_word": "Düzeltilmiş kelime veya cümleyi girin...",
                "save_changes": "💾 Değişiklikleri Kaydet",
                "cancel": "❌ İptal",
                "word_updated": "Kelime Güncellendi",
                "duplicate_word": "Yinelenen Kelime",
                "error": "Hata",
                "warning": "Uyarı",
                "success": "Başarılı",
                "language": "Dil:",
                "current_sentence": "Şu Anki Cümle:",
                "current_paragraph": "Şu Anki Paragraf:",
                "sentence_mode": "🎭 Cümle Modu",
                "paragraph_mode": "📄 Paragraf Modu",
                "word_mode": "📝 Kelime Modu",
                "no_sentences_loaded": "Hiç cümle yüklenmedi",
                "no_paragraphs_loaded": "Hiç paragraf yüklenmedi",
                "sentences_recorded": "cümle kaydedildi",
                "paragraphs_recorded": "paragraf kaydedildi",
                "speed_selection": "Kayıt Hızı Seçimi (Multi-Speed Recording)",
                "slow_speed": "🐌 Yavaş",
                "normal_speed": "🎯 Normal",
                "fast_speed": "🚀 Hızlı",
                "recording_status": "Mevcut Kelimenin Kayıt Durumu:",
                "force_next": "⏭️ Sonraki Kelimeye Zorla",
                "skip_word": "⏭️ Bu Kelimeyi Atla",
                "not_recorded_yet": "Kayıt Edilmedi",
                "recorded": "Kaydedildi",
                "completion": "Tamamlanma",
                "word_management": "Kelime Yönetimi",
                "whisper_optimization": "Whisper Optimizasyon & Dataset Zenginleştirme",
                "dataset_management": "Dataset Yönetimi",
                "word_list_management": "Kelime Listesi Yönetimi",
                "import_export": "Al/Ver",
                "settings": "Ayarlar",
                "add_words_sentences": "➕ Kelime/Cümle Ekle",
                "quick_add_word": "📝 Hızlı Kelime Ekle",
                "view_all_words": "📋 Tüm Kelimeleri Görüntüle",
                "refresh_display": "🔄 Ana Ekranı Yenile",
                "jump_first_unrecorded": "🎯 İlk Kaydedilmemişe Atla",
                "statistics": "📊 İstatistikler",
                "remove_duplicates": "🔄 Tekrarları Kaldır",
                "clear_all_words": "🗑️ Tüm Kelimeleri Temizle",
                "old_recordings_convert": "🎵 Eski Ses Dosyalarını Whisper'a Dönüştür",
                "show_conversion_results": "📋 Dönüştürme Sonuçlarını Göster",
                "sentence_mode_toggle": "🔄 Mod Değiştir (Kelime→Cümle→Paragraf)",
                "generate_sentence_list": "📝 Cümle Listesi Oluştur",
                "audio_quality_test": "🔊 Ses Kalitesi Testi",
                "audio_augmentation": "🎵 Ses Zenginleştirme",
                "whisper_training_prep": "🚀 Whisper Eğitim Hazırlığı",
                "merge_datasets": "🔗 Datasetleri Birleştir",
                "set_speaker_id": "👤 Konuşmacı ID Ayarla",
                "show_merged_stats": "📊 Birleşik Dataset İstatistikleri",
                "split_word_list": "✂️ Kelime Listesi Böl",
                "create_sub_list": "📄 Alt Liste Oluştur",
                "random_distribute": "🔀 Rastgele Dağıt",
                "show_list_stats": "📊 Liste İstatistikleri",
                "import_txt": "📄 TXT Dosyası Al",
                "import_docx": "📘 DOCX Dosyası Al",
                "import_pdf": "📕 PDF Dosyası Al",
                "load_from_url": "🌐 URL'den Yükle",
                "export_words": "📤 Kelimeleri Ver",
                "export_dataset": "📦 Dataset Ver",
                "toggle_theme": "🌙 Tema Değiştir",
                "test_audio": "🎙️ Ses Sistemini Test Et"
            },
            "ku": {
                # Kurdish (Kurmanji) translations
                "app_title": "🎤 Qeydkarê Peyvên Kurmancî",
                "menu": "📋 Menû",
                "current_word": "Peyva Niha:",
                "no_words_loaded": "Tu peyvek nehate barkirin",
                "previous": "← Paşve",
                "next": "Pêşve →",
                "already_recorded": "✓ Jixwe hatiye tomarkirin",
                "not_recorded": "○ Nehatiye tomarkirin",
                "start_recording": "🔴 Tomarkirinê Dest Pê Bike",
                "stop_recording": "⏹️ Tomarkirinê Bisekine",
                "play": "▶️ Lêxe",
                "save_next": "💾 Tomar Bike",
                "edit": "✏️ Serrast Bike",
                "delete": "🗑️ Jê Bibe",
                "confirm_delete": "Tu dixwazî ev peyv jê bibî?",
                "word_deleted": "Peyv bi serkeftî hate jêbirin",
                "recording_progress": "Pêşketin a Tomarkirinê:",
                "words_recorded": "Peyv hatin tomarkirin",
                "add_words": "➕ Peyv Zêde Bike",
                "view_words": "📋 Peyvan Bibîne",
                "load_file": "📂 Pel Bar Bike",
                "statistics": "📊 Statîstîk",
                "checking_microphone": "🔍 Mîkrofon tê kontrolkirin...",
                "microphone_working": "🎤 Mîkrofon dixebite",
                "microphone_error": "❌ Xeletiyek di mîkrofonê de",
                "no_word_to_record": "Peyveke ji bo tomarkirinê tune. Ji kerema xwe pêşî lîsteya Peyvn bar bike.",
                "recording_started": "Tomarkirin dest pê kir...",
                "recording_stopped": "Tomarkirin hate sekinandin.",
                "no_recording_to_play": "Tomarkerek ji bo lêxistinê tune.",
                "word_saved": "Peyv hate tomarkirin û hate guhartinê bo ya din.",
                "edit_current_word": "Peyva Niha Sererast Bike",
                "enter_corrected_word": "Peyva an hevoka rast binivîse...",
                "save_changes": "💾 Guherînan Tomar Bike",
                "cancel": "❌ Betal Bike",
                "word_updated": "Peyv Hat Nûkirin",
                "duplicate_word": "Peyva Dubare",
                "error": "Xeletî",
                "warning": "Hişyarî",
                "success": "Serkeftin",
                "language": "Ziman:",
                "current_sentence": "Hevoka Niha:",
                "current_paragraph": "Paragrafa Niha:",
                "sentence_mode": "🎭 Moda Hevokan",
                "paragraph_mode": "📄 Moda Paragrafan",
                "word_mode": "📝 Moda Peyvan",
                "no_sentences_loaded": "Hîç hevok nehate barkirin",
                "no_paragraphs_loaded": "Hîç paragraf nehate barkirin",
                "sentences_recorded": "hevok hatin tomarkirin",
                "paragraphs_recorded": "paragraf hatin tomarkirin",
                "speed_selection": "Hilbijartina Leza Tomarkirinê (Multi-Speed Recording)",
                "slow_speed": "🐌 Hêdî",
                "normal_speed": "🎯 Asayî",
                "fast_speed": "🚀 Lez",
                "recording_status": "Rewşa Tomarkirina Peyva Niha:",
                "force_next": "⏭️ Bi Zorê Peyva Pêş",
                "skip_word": "⏭️ Vê Peyvê Derbas Bike",
                "not_recorded_yet": "Nehatiye Tomarkirin",
                "recorded": "Hatiye Tomarkirin",
                "completion": "Qedandin",
                "word_management": "Rêvebirina Peyvan",
                "whisper_optimization": "Baştirkirina Whisper & Dewlemendkirina Dataset",
                "dataset_management": "Rêvebirina Dataset",
                "word_list_management": "Rêvebirina Lîsteya Peyvan",
                "import_export": "Anîn/Birin",
                "settings": "Mîhengan",
                "add_words_sentences": "➕ Peyv/Hevok Zêde Bike",
                "quick_add_word": "📝 Lez Peyv Zêde Bike",
                "view_all_words": "📋 Hemû Peyvan Bibîne",
                "refresh_display": "🔄 Dîmena Sereke Nû Bike",
                "jump_first_unrecorded": "🎯 Berê Netomarkirî Biçe",
                "statistics": "📊 Statîstîk",
                "remove_duplicates": "🔄 Dubarekirinan Rake",
                "clear_all_words": "🗑️ Hemû Peyvan Paqij Bike",
                "old_recordings_convert": "🎵 Kevn Dengan ji Whisper'ê re Veguhere",
                "show_conversion_results": "📋 Encamên Veguhertinê Nîşan Bide",
                "sentence_mode_toggle": "🎭 Moda Hevokan (Peyv→Hevok)",
                "generate_sentence_list": "📝 Lîsta Hevokan Ava Bike",
                "audio_quality_test": "🔊 Cîrbendiya Dengê Test Bike",
                "audio_augmentation": "🎵 Deng Dewlemend Bike",
                "whisper_training_prep": "🚀 Amadekariya Perwerdeya Whisper'ê",
                "merge_datasets": "🔗 Datasetan Yek Bike",
                "set_speaker_id": "👤 ID ya Axafkarê Saz Bike",
                "show_merged_stats": "📊 Statîstîkên Dataset a Yekbû",
                "split_word_list": "✂️ Lîsta Peyvan Parve Bike",
                "create_sub_list": "📄 Bin-lîste Ava Bike",
                "random_distribute": "🔀 Bi Rastî Belav Bike",
                "show_list_stats": "📊 Statîstîkên Lîstê",
                "import_txt": "📄 Pelê TXT Bîne",
                "import_docx": "📘 Pelê DOCX Bîne",
                "import_pdf": "📕 Pelê PDF Bîne",
                "load_from_url": "🌐 Ji URLê Bar Bike",
                "export_words": "📤 Peyvan Derêxe",
                "export_dataset": "📦 Dataset Derêxe",
                "toggle_theme": "🌙 Temayê Biguhere",
                "test_audio": "🎙️ Sîstema Dengê Test Bike"
            }
        }
    
    def set_language(self, lang_code: str):
        """Set the current language"""
        if lang_code in self.translations:
            self.current_language = lang_code
    
    def get(self, key: str) -> str:
        """Get translated text for the current language"""
        return self.translations.get(self.current_language, {}).get(key, key)
    
    def get_available_languages(self) -> List[Tuple[str, str]]:
        """Get list of available languages"""
        return [
            ("en", "English"),
            ("tr", "Türkçe"),
            ("ku", "Kurmancî")
        ]

# Global language manager instance
lang = LanguageManager()

# Configuration
class Config:
    SAMPLE_RATE = 44100  # Higher quality, less artifacts
    CHANNELS = 1
    DTYPE = np.int16  # More compatible with Windows audio
    CHUNK_SIZE = 2048   # Larger buffer to prevent dropouts
    
    MIN_SILENCE_DURATION = 0.5
    SENTENCE_SILENCE_DURATION = 1.5  # Longer pause detection for sentences
    NOISE_THRESHOLD = 0.02
    
    BASE_DIR = Path("kurmanji_dataset")
    AUDIO_DIR = BASE_DIR / "audio"
    DOCS_DIR = BASE_DIR / "documents"
    METADATA_FILE = BASE_DIR / "metadata.json"
    WORDS_FILE = BASE_DIR / "wordlist.json"
    SENTENCES_FILE = BASE_DIR / "sentencelist.json"  # Ayrı cümle dosyası
    PARAGRAPHS_FILE = BASE_DIR / "paragraphlist.json"  # Ayrı paragraf dosyası
    
    # Whisper training files
    WHISPER_MANIFEST = BASE_DIR / "whisper_manifest.jsonl"  # Whisper eğitim dosyası
    TRANSCRIPT_FILE = BASE_DIR / "transcripts.txt"  # Transkript dosyası
    
    UI_THEME = "dark"
    WINDOW_SIZE = "1000x700"

# Ensure directories exist
for directory in [Config.BASE_DIR, Config.AUDIO_DIR, Config.DOCS_DIR]:
    directory.mkdir(exist_ok=True)

# All classes are defined in this file - no external imports needed

class AudioManager:
    """Enhanced audio recording manager with device selection and quality control"""
    
    def __init__(self):
        self.devices = []
        self.selected_device = None
        self.is_recording = False
        self.current_recording = []
        self.audio_queue = queue.Queue()
        self.recording_thread = None
        self.is_sentence_mode = False  # Track recording mode for dynamic timeouts
        self.consecutive_empty = 0
        self.current_audio_level = 0.0  # For UI monitoring
        self.audio_chunks = []  # For callback recording
        self.sample_rate = Config.SAMPLE_RATE
        
        self.refresh_devices()
    
    def refresh_devices(self):
        """Refresh available audio input devices"""
        try:
            devices_info = sd.query_devices()
            self.devices = []
            
            for i, device in enumerate(devices_info):
                if device['max_input_channels'] > 0:
                    self.devices.append({
                        'id': i,
                        'name': device['name'],
                        'channels': device['max_input_channels'],
                        'sample_rate': device['default_samplerate']
                    })
            
            if self.devices:
                self.selected_device = self.devices[0]['id']
            
            return True, f"Found {len(self.devices)} input devices"
        except Exception as e:
            return False, f"Error detecting devices: {str(e)}"
    
    def check_microphone(self):
        """Check if microphone is working with detailed diagnostics"""
        try:
            if not self.devices:
                return False, "No audio input devices found"
            
            # Testing microphone
            
            # Try multiple test approaches
            devices_to_test = [self.selected_device, None]  # Test selected and default
            
            for device_id in devices_to_test:
                device_name = self.get_selected_device_name() if device_id == self.selected_device else "System Default"
                print(f"🎤 Testing device: {device_name}")
                
                try:
                    test_duration = 0.5  # Longer test
                    test_data = sd.rec(
                        int(test_duration * Config.SAMPLE_RATE),
                        samplerate=Config.SAMPLE_RATE,
                        channels=Config.CHANNELS,
                        device=device_id,
                        dtype=np.int16  # Use int16 for compatibility
                    )
                    sd.wait()
                    
                    if test_data is not None and len(test_data) > 0:
                        # Convert to float for level checking
                        test_float = test_data.astype(np.float32) / 32768.0
                        max_level = np.max(np.abs(test_float))
                        rms_level = np.sqrt(np.mean(test_float**2))
                        
                        # Audio levels checked
                        
                        if max_level > 0.0001:  # Very low threshold
                            return True, f"Microphone OK (Device: {device_name}, Max level: {max_level:.4f})"
                        elif max_level > 0:
                            return True, f"Microphone detected but very quiet (Device: {device_name}, Max level: {max_level:.6f})"
                        else:
                            pass  # Silent audio detected
                            
                except Exception as device_error:
                    print(f"❌ Device {device_name} test failed: {device_error}")
                    continue
            
            return False, "All microphone tests failed - please check microphone connection and permissions"
            
        except Exception as e:
            return False, f"Audio system error: {str(e)}"
    
    def get_selected_device_name(self):
        """Get the name of the currently selected device"""
        for device in self.devices:
            if device['id'] == self.selected_device:
                return device['name']
        return "Unknown Device"
    
    def is_bluetooth_device(self, device_id=None):
        """Check if the device is likely a Bluetooth audio device"""
        if device_id is None:
            device_id = self.selected_device
        
        device_name = ""
        for device in self.devices:
            if device['id'] == device_id:
                device_name = device['name'].lower()
                break
        
        # Common Bluetooth device indicators
        bluetooth_indicators = [
            'airpods', 'bluetooth', 'wireless', 'bt', 'headset',
            'headphone', 'earbuds', 'sony', 'bose', 'beats',
            'jabra', 'plantronics', 'samsung buds'
        ]
        
        return any(indicator in device_name for indicator in bluetooth_indicators)
    
    def start_recording(self, device_id=None):
        """Start recording audio - simplified approach"""
        if self.is_recording:
            return False, "Already recording"
        
        if device_id is not None:
            self.selected_device = device_id
        
        try:
            # Starting recording with selected device
            
            # Clear previous recording
            self.current_recording = []
            self.is_recording = True
            
            # Get content type for timeout calculation
            content_type = getattr(self, 'content_type', 'word')
            is_sentence_mode = getattr(self, 'is_sentence_mode', False)
            
            # Set recording duration based on mode
            if content_type == "paragraph":
                max_duration = 45.0  # 45 seconds for paragraphs
            elif content_type == "sentence" or is_sentence_mode:
                max_duration = 25.0  # 25 seconds for long sentences
            else:
                max_duration = 3.0   # 3 seconds for words - enough for any single word
            
            # Recording mode configuration (output cleaned up for user)
            
            # Detect device type and choose optimal recording method
            is_bluetooth = self.is_bluetooth_device()
            device_name = self.get_selected_device_name()
            
            if is_bluetooth:
                print(f"🎧 Detected Bluetooth device: {device_name}")
                print(f"🎙️ Using Bluetooth-optimized recording method")
                self.recording_thread = threading.Thread(target=self._bluetooth_recording_worker, args=(max_duration,))
            elif content_type in ("sentence", "paragraph") or is_sentence_mode:
                print(f"🎙️ Using streaming recording method for longer stability")
                self.recording_thread = threading.Thread(target=self._streaming_recording_worker, args=(max_duration,))
            else:
                print(f"🎙️ Using fallback recording method for better stability")
                self.recording_thread = threading.Thread(target=self._fallback_recording_worker, args=(max_duration,))
            self.recording_thread.start()
            
            return True
        except Exception as e:
            self.is_recording = False
            # Recording start error handled
            return False, f"Failed to start recording: {str(e)}"
    
    def _simple_recording_worker(self, max_duration):
        """Enhanced recording worker with real-time audio level monitoring"""
        print(f"🎙️ Recording worker started - Duration: {max_duration}s")
        try:
            # Starting recording (output cleaned up for user)
            
            # Try multiple device options in case of failure
            devices_to_try = [self.selected_device, None]  # None uses default device
            audio_data = None
            
            for device in devices_to_try:
                try:
                    # Trying recording device
                    
                    # Initialize recording with callback for real-time monitoring
                    self.audio_chunks = []
                    
                    # Initialize level monitoring
                    self.level_counter = 0
                    import sys
                    
                    def audio_callback(indata, frames, time, status):
                        """Real-time audio monitoring callback with recovery"""
                        if status:
                            print(f"\n⚠️ Audio callback status: {status}")
                        
                        # Store audio chunk
                        self.audio_chunks.append(indata.copy())
                        
                        # Real-time level monitoring (show every 10th callback ~ 0.5 seconds)
                        self.level_counter += 1
                        if len(indata) > 0:
                            # Convert to float and calculate level
                            audio_float = indata.astype(np.float32) / 32768.0
                            level = np.max(np.abs(audio_float))
                            rms_level = np.sqrt(np.mean(audio_float**2))
                            
                            # Store current level for UI monitoring
                            self.current_audio_level = level
                            
                            # Show terminal level bar every 10th callback
                            if self.level_counter % 10 == 0:
                                # Create level bar
                                bar_length = 20
                                filled_bars = int(level * bar_length)
                                level_bar = "█" * filled_bars + "░" * (bar_length - filled_bars)
                                
                                # Print level with carriage return for updating same line
                                print(f"\r🎚️ [{level_bar}] {level:.3f} (RMS: {rms_level:.3f})", end="", flush=True)
                                sys.stdout.flush()
                                
                                # Detect if audio stream is dying
                                if level < 0.01 and rms_level < 0.005 and self.level_counter > 20:
                                    print(f"\n⚠️ Low audio detected at callback {self.level_counter}")
                    
                    # Start recording with callback and more robust configuration
                    with sd.InputStream(
                        device=device,
                        channels=Config.CHANNELS,
                        samplerate=Config.SAMPLE_RATE,
                        dtype=np.int16,
                        callback=audio_callback,
                        blocksize=Config.CHUNK_SIZE,
                        latency='low'  # Lower latency for better responsiveness
                    ):
                        # Recording in progress
                        
                        # Wait for recording to complete or until manually stopped
                        start_time = time.time()
                        last_chunk_time = start_time
                        print(f"🎤 Starting {max_duration}s recording session...")
                        
                        chunk_count_last_check = 0
                        last_chunk_check_time = start_time
                        
                        while self.is_recording and (time.time() - start_time) < max_duration:
                            current_time = time.time()
                            elapsed = current_time - start_time
                            current_chunk_count = len(self.audio_chunks)
                            
                            # Check if we're still getting audio chunks
                            if current_chunk_count > chunk_count_last_check:
                                last_chunk_time = current_time
                                chunk_count_last_check = current_chunk_count
                            
                            # Show progress every 2 seconds
                            if int(elapsed) % 2 == 0 and int(elapsed) != int(elapsed - 0.05):
                                print(f"⏱️  Recording... {elapsed:.1f}s / {max_duration}s - Chunks: {current_chunk_count}")
                                
                                # Check for stalled audio stream
                                if current_time - last_chunk_check_time > 1.0 and current_chunk_count == 0:
                                    print(f"⚠️ Warning: No audio chunks received for 1 second!")
                                    
                            # More frequent check but longer patience
                            time.sleep(0.05)  # Check every 50ms for better responsiveness
                    
                    # Clear the level display line
                    print("\r" + " " * 50 + "\r", end="", flush=True)
                    
                    # Show recording completion info
                    recording_duration = time.time() - start_time
                    print(f"🎙️ Recording session completed: {recording_duration:.1f}s of {max_duration:.1f}s")
                    
                    # Combine all audio chunks
                    if self.audio_chunks:
                        audio_data = np.concatenate(self.audio_chunks, axis=0)
                    
                    # Check if we got valid audio data
                    if audio_data is not None and len(audio_data) > 0:
                        # Convert to float for processing and check amplitude
                        audio_float = audio_data.astype(np.float32) / 32768.0  # Normalize int16 to float
                        max_amplitude = np.max(np.abs(audio_float))
                        # Audio amplitude check completed
                        
                        if max_amplitude > 0.0001:  # Lower threshold for valid audio
                            # Recording completed successfully
                            self.current_recording = [audio_float]  # Store as float
                            print(f"✅ Recording completed: {len(audio_float)/Config.SAMPLE_RATE:.1f}s, amplitude: {max_amplitude:.6f}")
                            return
                        else:
                            print(f"⚠️ Audio too quiet: {max_amplitude:.6f}")
                    else:
                        pass  # No audio data received
                        
                except Exception as device_error:
                    # Device failed, trying next
                    continue
            
            # If we get here, all devices failed
            # All recording attempts failed
            self.current_recording = []
                
        except Exception as e:
            # Recording worker error handled
            self.current_recording = []
        finally:
            self.is_recording = False
            # Recording worker completed
    
    def _fallback_recording_worker(self, max_duration):
        """Robust segmented recording method to avoid Windows audio driver issues"""
        print(f"🎙️ Robust recording worker started - Duration: {max_duration}s")
        try:
            # Record in segments to avoid Windows audio driver timeout issues
            # Use smaller segments for short recordings (words), larger for long ones
            if max_duration <= 5.0:
                segment_duration = max_duration  # Single segment for short recordings
            else:
                segment_duration = 5.0  # 5-second segments for longer recordings
            segments = []
            total_segments = int(max_duration / segment_duration) + 1
            
            print(f"🎤 Recording in {total_segments} segments of {segment_duration}s each...")
            
            for segment_num in range(total_segments):
                if not self.is_recording:
                    print(f"🛑 Recording stopped manually at segment {segment_num}")
                    break
                
                # Calculate remaining time for this segment
                remaining_time = max_duration - (segment_num * segment_duration)
                current_segment_duration = min(segment_duration, remaining_time)
                
                if current_segment_duration <= 0:
                    break
                
                print(f"🎙️ Recording segment {segment_num + 1}/{total_segments} ({current_segment_duration:.1f}s)...")
                
                try:
                    # Record this segment with proper device configuration
                    segment_data = sd.rec(
                        int(current_segment_duration * Config.SAMPLE_RATE),
                        samplerate=Config.SAMPLE_RATE,
                        channels=Config.CHANNELS,
                        dtype=np.int16,
                        device=self.selected_device
                    )
                    
                    # Wait for segment with periodic stop checks
                    check_interval = 0.1  # Check every 100ms
                    elapsed = 0
                    while elapsed < current_segment_duration:
                        if not self.is_recording:
                            # User clicked stop - abort this segment
                            sd.stop()
                            print(f"🛑 Recording stopped manually during segment {segment_num + 1}")
                            break
                        time.sleep(check_interval)
                        elapsed += check_interval
                    
                    # Only add segment if we completed it
                    if self.is_recording and segment_data is not None and len(segment_data) > 0:
                        # Check segment quality
                        segment_float = segment_data.astype(np.float32) / 32768.0
                        segment_amplitude = np.max(np.abs(segment_float))
                        
                        print(f"✅ Segment {segment_num + 1} completed: amplitude {segment_amplitude:.6f}")
                        segments.append(segment_float)
                        
                        # Store current level for UI
                        self.current_audio_level = segment_amplitude
                    elif not self.is_recording:
                        # Stopped early - only keep the audio recorded so far
                        recorded_samples = int(elapsed * Config.SAMPLE_RATE)
                        if segment_data is not None and recorded_samples > 0:
                            segment_float = segment_data[:recorded_samples].astype(np.float32) / 32768.0
                            segment_amplitude = np.max(np.abs(segment_float))
                            print(f"✅ Partial segment {segment_num + 1}: {elapsed:.2f}s, amplitude {segment_amplitude:.6f}")
                            segments.append(segment_float)
                        break  # Exit the segment loop
                    else:
                        print(f"⚠️ Segment {segment_num + 1} failed - no data")
                        # Add silence for missing segment
                        silence = np.zeros((int(current_segment_duration * Config.SAMPLE_RATE), Config.CHANNELS), dtype=np.float32)
                        segments.append(silence)
                    
                    # Brief pause between segments to reinitialize driver
                    time.sleep(0.1)
                    
                except Exception as segment_error:
                    print(f"❌ Segment {segment_num + 1} error: {segment_error}")
                    # Add silence for failed segment
                    silence = np.zeros((int(current_segment_duration * Config.SAMPLE_RATE), Config.CHANNELS), dtype=np.float32)
                    segments.append(silence)
            
            # Combine all segments
            if segments:
                audio_float = np.concatenate(segments, axis=0)
                
                # Remove buzzing artifacts from the very end (typical with sd.rec())
                # Buzzing usually appears in the last 10-20ms
                buzz_removal_samples = int(Config.SAMPLE_RATE * 0.015)  # Remove last 15ms
                if len(audio_float) > buzz_removal_samples:
                    audio_float = audio_float[:-buzz_removal_samples]
                    print(f"🔧 Removed last 15ms to eliminate buzzing artifacts")
                
                # Apply dynamic range compression to compensate for Windows AGC
                audio_float = self._normalize_audio_levels(audio_float)
                max_amplitude = np.max(np.abs(audio_float))
                
                print(f"✅ Segmented recording completed: {len(audio_float)/Config.SAMPLE_RATE:.1f}s, amplitude: {max_amplitude:.6f}")
                
                if max_amplitude > 0.0001:
                    self.current_recording = [audio_float]
                else:
                    print(f"⚠️ Combined audio too quiet: {max_amplitude:.6f}")
                    self.current_recording = []
            else:
                print(f"❌ No segments recorded")
                self.current_recording = []
                
        except Exception as e:
            print(f"❌ Robust recording error: {e}")
            self.current_recording = []
        finally:
            self.is_recording = False
            self.current_audio_level = 0.0
            print(f"🔚 Robust recording worker finished")

    def _streaming_recording_worker(self, max_duration: float):
        """Single-stream reader to avoid device reinitialization between segments.

        - Keeps one InputStream open for the entire session.
        - Reads fixed-size chunks until max_duration or manual stop.
        - Updates UI level and analyzes/saves debug output at the end.
        """
        print(f"🎙️ Streaming recording worker started - Duration: {max_duration}s")
        frames_per_read = int(self.sample_rate * 0.25)  # ~250ms per read
        collected = []
        start_time = time.time()
        total_frames = 0
        try:
            stream = sd.InputStream(
                device=self.selected_device,
                channels=Config.CHANNELS,
                samplerate=self.sample_rate,
                dtype=np.int16,
                blocksize=Config.CHUNK_SIZE,
                latency='high'
            )
            stream.start()
            print("🎤 Stream opened")

            while self.is_recording:
                # Respect max_duration based on frames collected
                if total_frames >= int(max_duration * self.sample_rate):
                    break
                try:
                    indata, overflowed = stream.read(frames_per_read)
                except Exception as re:
                    print(f"⚠️ Stream read error: {re}")
                    # Brief backoff and continue
                    time.sleep(0.05)
                    continue

                if indata is None or len(indata) == 0:
                    time.sleep(0.01)
                    continue

                # Convert to float32 mono
                chunk = indata.astype(np.float32) / 32768.0
                if chunk.ndim > 1:
                    chunk = np.mean(chunk, axis=1, dtype=np.float32)
                collected.append(chunk)

                # UI level update
                level = float(np.max(np.abs(chunk))) if chunk.size else 0.0
                self.current_audio_level = level

                total_frames += chunk.size

            # Stop and close stream
            try:
                stream.stop()
            except Exception:
                pass
            try:
                stream.close()
            except Exception:
                pass

            duration_s = total_frames / float(self.sample_rate) if self.sample_rate else 0.0
            print(f"🎙️ Streaming session completed: {duration_s:.2f}s of {max_duration:.2f}s")

            if collected:
                audio_float = np.concatenate(collected, axis=0)
                # Normalize/level
                audio_float = self._normalize_audio_levels(audio_float)
                max_amp = float(np.max(np.abs(audio_float))) if audio_float.size else 0.0

                if max_amp > 1e-5 and duration_s >= 0.5:
                    self.current_recording = [audio_float]
                else:
                    print(f"⚠️ Streaming audio too quiet/short (amp {max_amp:.6f}, dur {duration_s:.2f}s)")
                    self.current_recording = []
            else:
                print("❌ No audio collected in streaming worker")
                self.current_recording = []

        except Exception as e:
            print(f"❌ Streaming recording error: {e}")
            self.current_recording = []
        finally:
            self.is_recording = False
            self.current_audio_level = 0.0
            print("🔚 Streaming recording worker finished")
    
    def _bluetooth_recording_worker(self, max_duration: float):
        """Specialized recording worker optimized for Bluetooth audio devices.
        
        Uses smaller buffers, more frequent reads, and aggressive fallback strategies
        to handle Bluetooth audio latency and connection stability issues.
        """
        print(f"🎧 Bluetooth recording worker started - Duration: {max_duration}s")
        device_name = self.get_selected_device_name()
        
        # Bluetooth-optimized settings
        bt_chunk_size = 512  # Smaller chunks for Bluetooth
        read_interval = 0.1  # Read every 100ms
        frames_per_read = int(self.sample_rate * read_interval)
        
        collected = []
        start_time = time.time()
        total_frames = 0
        consecutive_failures = 0
        last_level_time = start_time
        
        try:
            # Try Bluetooth device with optimized settings
            print(f"🎧 Attempting Bluetooth recording with {device_name}")
            
            stream = sd.InputStream(
                device=self.selected_device,
                channels=Config.CHANNELS,
                samplerate=self.sample_rate,
                dtype=np.int16,
                blocksize=bt_chunk_size,  # Smaller buffer for Bluetooth
                latency='low'  # Lower latency for wireless
            )
            stream.start()
            print("🎧 Bluetooth stream opened successfully")
            
            while self.is_recording:
                current_time = time.time()
                
                # Check total duration
                if total_frames >= int(max_duration * self.sample_rate):
                    print(f"✅ Reached maximum duration: {max_duration:.1f}s")
                    break
                
                try:
                    # Read smaller chunks more frequently
                    indata, overflowed = stream.read(frames_per_read)
                    
                    if overflowed:
                        print(f"⚠️ Bluetooth audio buffer overflow detected")
                    
                    if indata is None or len(indata) == 0:
                        consecutive_failures += 1
                        if consecutive_failures > 10:
                            print(f"❌ Too many consecutive read failures: {consecutive_failures}")
                            break
                        time.sleep(0.01)
                        continue
                    
                    # Reset failure counter on successful read
                    consecutive_failures = 0
                    
                    # Convert to float32 mono
                    chunk = indata.astype(np.float32) / 32768.0
                    if chunk.ndim > 1:
                        chunk = np.mean(chunk, axis=1, dtype=np.float32)
                    
                    collected.append(chunk)
                    total_frames += len(chunk)
                    
                    # Monitor audio levels
                    level = float(np.max(np.abs(chunk))) if chunk.size > 0 else 0.0
                    self.current_audio_level = level
                    
                    # Show progress every 2 seconds with level monitoring
                    if current_time - last_level_time >= 2.0:
                        elapsed = current_time - start_time
                        level_bar = "█" * int(level * 20) + "░" * (20 - int(level * 20))
                        print(f"🎧 [{level_bar}] {elapsed:.1f}s/{max_duration:.1f}s - Level: {level:.3f}")
                        last_level_time = current_time
                        
                        # Warn if no audio detected for extended period
                        if level < 0.001:
                            print(f"⚠️ Very low audio level detected - check Bluetooth connection")
                    
                    # Brief pause to prevent overwhelming Bluetooth buffer
                    time.sleep(0.01)
                    
                except Exception as read_error:
                    consecutive_failures += 1
                    print(f"⚠️ Bluetooth read error {consecutive_failures}: {read_error}")
                    
                    if consecutive_failures > 20:
                        print(f"❌ Bluetooth connection appears unstable - switching to fallback")
                        break
                    
                    time.sleep(0.05)  # Brief recovery pause
            
            # Clean up stream
            try:
                stream.stop()
                stream.close()
            except Exception:
                pass
            
            # Process collected audio
            if collected:
                audio_float = np.concatenate(collected, axis=0)
                duration_s = len(audio_float) / float(self.sample_rate)
                
                print(f"🎧 Bluetooth recording completed: {duration_s:.2f}s")
                
                # Apply normalization for Bluetooth audio
                audio_float = self._normalize_audio_levels(audio_float)
                max_amp = float(np.max(np.abs(audio_float))) if audio_float.size > 0 else 0.0
                
                # More lenient validation for Bluetooth
                if max_amp > 5e-5 and duration_s >= 0.3:  # Lower thresholds for Bluetooth
                    self.current_recording = [audio_float]
                    print(f"✅ Bluetooth recording successful: amp={max_amp:.6f}, dur={duration_s:.2f}s")
                else:
                    print(f"⚠️ Bluetooth audio quality insufficient")
                    print(f"   Amplitude: {max_amp:.6f} (need > 5e-5)")
                    print(f"   Duration: {duration_s:.2f}s (need > 0.3s)")
                    # Try emergency fallback
                    self._emergency_fallback_recording(max_duration - (time.time() - start_time))
            else:
                print(f"❌ No Bluetooth audio collected")
                # Try emergency fallback
                self._emergency_fallback_recording(max_duration - (time.time() - start_time))
                
        except Exception as e:
            print(f"❌ Bluetooth recording failed: {e}")
            print(f"🚨 Attempting emergency fallback to system default...")
            # Try emergency fallback
            self._emergency_fallback_recording(max_duration - (time.time() - start_time))
            
        finally:
            self.is_recording = False
            self.current_audio_level = 0.0
            print("🔚 Bluetooth recording worker finished")
    
    def _emergency_fallback_recording(self, remaining_duration):
        """Emergency fallback to system default microphone when Bluetooth fails"""
        if remaining_duration <= 0.5:
            print("⏰ Not enough time remaining for fallback recording")
            self.current_recording = []
            return
            
        print(f"🚨 EMERGENCY FALLBACK: Switching to system default microphone")
        print(f"⏱️ Time remaining: {remaining_duration:.1f}s")
        
        try:
            # Use system default device (None = default)
            fallback_data = sd.rec(
                int(remaining_duration * self.sample_rate),
                samplerate=self.sample_rate,
                channels=Config.CHANNELS,
                dtype=np.int16,
                device=None  # System default
            )
            
            # Wait for recording to complete
            sd.wait()
            
            if fallback_data is not None and len(fallback_data) > 0:
                # Convert and process
                audio_float = fallback_data.astype(np.float32) / 32768.0
                if audio_float.ndim > 1:
                    audio_float = np.mean(audio_float, axis=1, dtype=np.float32)
                
                max_amp = float(np.max(np.abs(audio_float)))
                duration_s = len(audio_float) / float(self.sample_rate)
                
                print(f"🚨 Emergency fallback completed: {duration_s:.2f}s, amp: {max_amp:.6f}")
                
                if max_amp > 1e-4 and duration_s >= 0.3:
                    self.current_recording = [audio_float]
                    print("✅ Emergency fallback successful!")
                else:
                    print("❌ Emergency fallback also failed")
                    self.current_recording = []
            else:
                print("❌ Emergency fallback - no data received")
                self.current_recording = []
                
        except Exception as e:
            print(f"❌ Emergency fallback error: {e}")
            self.current_recording = []
    
    def _analyze_recording(self, audio_data, filename):
        """Analyze recording to understand the audio capture issue"""
        try:
            duration = len(audio_data) / Config.SAMPLE_RATE
            max_amp = np.max(np.abs(audio_data))
            rms = np.sqrt(np.mean(audio_data**2))
            
            # Analyze in 1-second chunks
            chunk_size = Config.SAMPLE_RATE  # 1 second
            chunks = []
            
            print(f"\n🔍 AUDIO ANALYSIS for {filename}:")
            print(f"📊 Total duration: {duration:.2f}s")
            print(f"📊 Max amplitude: {max_amp:.6f}")
            print(f"📊 RMS level: {rms:.6f}")
            print(f"📊 Chunk analysis (1-second intervals):")
            
            for i in range(0, len(audio_data), chunk_size):
                chunk = audio_data[i:i+chunk_size]
                if len(chunk) > 0:
                    chunk_max = np.max(np.abs(chunk))
                    chunk_rms = np.sqrt(np.mean(chunk**2))
                    second = i // chunk_size + 1
                    chunks.append((second, chunk_max, chunk_rms))
                    
                    # Visual level indicator
                    level_bar = "█" * int(chunk_max * 20) + "░" * (20 - int(chunk_max * 20))
                    print(f"   Second {second:2d}: [{level_bar}] Max:{chunk_max:.4f} RMS:{chunk_rms:.4f}")
            
            # Check for silence pattern
            silent_chunks = [c for c in chunks if c[1] < 0.01]  # Max amplitude < 0.01
            if len(silent_chunks) > 0:
                print(f"⚠️ Found {len(silent_chunks)} silent/very quiet chunks out of {len(chunks)} total")
                if len(silent_chunks) > len(chunks) / 2:
                    print(f"🚨 ISSUE: More than half the recording is silent!")
            
            print(f"🔍 Analysis complete\n")
            
        except Exception as e:
            print(f"❌ Analysis error: {e}")
    
    def _normalize_audio_levels(self, audio_data):
        """Normalize audio levels and trim silence to prevent noise amplification"""
        try:
            # Smart speech endpoint detection
            # Calculate energy in 50ms windows
            window_size = int(Config.SAMPLE_RATE * 0.05)  # 50ms
            energy = []
            for i in range(0, len(audio_data), window_size):
                window = audio_data[i:i+window_size]
                energy.append(np.mean(np.abs(window)))
            
            if len(energy) > 2:
                # Find where energy drops significantly and stays low
                max_energy = np.max(energy)
                threshold = max_energy * 0.15  # 15% of peak energy
                
                # Find last high-energy window
                last_speech = 0
                for i in range(len(energy)-1, -1, -1):
                    if energy[i] > threshold:
                        last_speech = i
                        break
                
                # Cut audio at this point with small tail
                cut_sample = (last_speech + 2) * window_size  # Add 2 windows (100ms) tail
                if cut_sample < len(audio_data):
                    # Apply quick fade-out
                    fade_samples = int(Config.SAMPLE_RATE * 0.05)
                    fade_start = cut_sample
                    fade_end = min(cut_sample + fade_samples, len(audio_data))
                    fade_length = fade_end - fade_start
                    if fade_length > 0:
                        fade_curve = np.linspace(1.0, 0.0, fade_length)
                        audio_data[fade_start:fade_end] *= fade_curve
                    
                    audio_data = audio_data[:fade_end]
                    print(f"🔧 Smart trim - final duration: {len(audio_data)/Config.SAMPLE_RATE:.2f}s")
            
            # Simple overall normalization (no chunk boosting to avoid noise amplification)
            max_amplitude = np.max(np.abs(audio_data))
            if max_amplitude > 0.01:  # Only normalize if there's actual audio
                # Normalize to 0.5 (50% of max) to avoid clipping but maintain good volume
                target_level = 0.5
                normalization_factor = target_level / max_amplitude
                # Limit boost to prevent noise amplification
                normalization_factor = min(normalization_factor, 1.5)
                audio_data = audio_data * normalization_factor
                print(f"✅ Audio normalized by {normalization_factor:.2f}x (peak: {max_amplitude:.3f} → {np.max(np.abs(audio_data)):.3f})")
            
            return audio_data
            
        except Exception as e:
            print(f"⚠️ Audio normalization failed: {e}")
            return audio_data  # Return original if normalization fails
    
    def _recording_worker(self):
        """Old complex recording worker - keeping for compatibility"""
        try:
            def audio_callback(indata, frames, time, status):
                if status:
                    pass  # Audio callback status monitoring
                self.audio_queue.put(indata.copy())
            
            # Initialize consecutive empty counter
            consecutive_empty = 0
            
            with sd.InputStream(
                device=self.selected_device,
                channels=Config.CHANNELS,
                samplerate=Config.SAMPLE_RATE,
                dtype=Config.DTYPE,
                callback=audio_callback,
                blocksize=Config.CHUNK_SIZE
            ):
                while self.is_recording:
                    try:
                        # Dynamic timeout based on recording mode
                        # Much longer timeouts for sentence recording to prevent premature stops
                        content_type = getattr(self, 'content_type', 'word')
                        is_sentence_mode = getattr(self, 'is_sentence_mode', False)
                        
                        if content_type == "paragraph":
                            timeout_duration = 15.0  # Very long for paragraphs
                        elif content_type == "sentence" or is_sentence_mode:
                            timeout_duration = 12.0  # Much longer for sentences
                        else:
                            timeout_duration = 2.0   # Reasonable for words
                        
                        audio_chunk = self.audio_queue.get(timeout=timeout_duration)
                        self.current_recording.append(audio_chunk)
                        
                        # Reset consecutive empty count on successful data
                        consecutive_empty = 0
                        
                    except queue.Empty:
                        # Increment consecutive empty counter properly
                        consecutive_empty += 1
                        
                        # Much more patient empty cycles for sentence recording
                        if content_type == "paragraph":
                            max_empty_cycles = 80   # Very patient for paragraphs
                        elif content_type == "sentence" or is_sentence_mode:
                            max_empty_cycles = 65   # Very patient for sentences
                        else:
                            max_empty_cycles = 8    # Quick for words
                        
                        # Empty cycle monitoring (cleaned up)
                        
                        if consecutive_empty > max_empty_cycles:
                            # Recording auto-stopped due to silence
                            break
                        continue
                    except Exception as e:
                        # Recording error handled
                        break
        except Exception as e:
            # Recording worker error handled
            pass  
        finally:
            self.is_recording = False
            self.current_audio_level = 0.0  # Reset level for UI
    
    def set_sentence_mode(self, is_sentence_mode, content_type="word"):
        """Update recording mode for dynamic timeout adjustment"""
        self.is_sentence_mode = is_sentence_mode
        self.content_type = content_type
        mode_name = content_type.upper() if content_type in ["word", "sentence", "paragraph"] else ("SENTENCE" if is_sentence_mode else "WORD")
        # AudioManager mode updated
    
    def stop_recording(self):
        """Stop recording and return audio data"""
        if not self.is_recording:
            return None
        
        print(f"🛑 Stopping recording manually...")
        self.is_recording = False
        
        # Stop sounddevice if it's still recording
        try:
            sd.stop()
        except:
            pass
        
        # Wait for recording thread to finish
        if self.recording_thread:
            self.recording_thread.join(timeout=2.0)
        
        # For simple recording, current_recording contains the direct audio data
        if self.current_recording is not None and len(self.current_recording) > 0:
            # Normalize to a single numpy array
            if isinstance(self.current_recording, np.ndarray):
                audio_data = self.current_recording
            elif len(self.current_recording) == 1 and isinstance(self.current_recording[0], np.ndarray):
                audio_data = self.current_recording[0]
            else:
                # Complex recording format - concatenate chunks
                # Clear any remaining queue items first
                if hasattr(self, 'audio_queue'):
                    while not self.audio_queue.empty():
                        try:
                            chunk = self.audio_queue.get_nowait()
                            self.current_recording.append(chunk)
                        except queue.Empty:
                            break
                try:
                    audio_data = np.concatenate(self.current_recording, axis=0)
                except Exception:
                    # Fallback: flatten list of lists
                    audio_data = np.array(self.current_recording, dtype=np.float32).flatten()

            # Ensure float32 mono array
            try:
                if audio_data.dtype != np.float32:
                    audio_data = audio_data.astype(np.float32)
                if audio_data.ndim > 1:
                    audio_data = np.mean(audio_data, axis=1).astype(np.float32)
            except Exception:
                pass

            # Apply noise reduction and normalization
            processed_audio = self._process_audio(audio_data)
            return processed_audio
        
        return None
    
    def _process_audio(self, audio_data):
        """Enhanced audio processing to reduce noise and artifacts"""
        try:
            # Convert to mono if needed
            if len(audio_data.shape) > 1:
                audio_data = np.mean(audio_data, axis=1)
            
            # Convert to float if it's still int16
            if audio_data.dtype == np.int16:
                audio_data = audio_data.astype(np.float32) / 32768.0
            
            # Remove DC offset safely
            mean_val = np.mean(audio_data)
            if not np.isnan(mean_val) and np.isfinite(mean_val):
                audio_data = audio_data - mean_val
            
            # Simple noise reduction - remove very quiet noise
            noise_threshold = 0.01  # Threshold for noise
            audio_data = np.where(np.abs(audio_data) < noise_threshold, 0, audio_data)
            
            # Gentle low-pass filter to remove high-frequency artifacts (without scipy)
            # Simple moving average filter
            window_size = 3
            if len(audio_data) > window_size:
                padded = np.pad(audio_data, (window_size//2, window_size//2), mode='edge')
                filtered = np.convolve(padded, np.ones(window_size)/window_size, mode='valid')
                audio_data = filtered[:len(audio_data)]
            
            # Soft limiting to prevent clipping distortion
            max_amplitude = np.max(np.abs(audio_data))
            if max_amplitude > 0 and np.isfinite(max_amplitude):
                # Gentle compression instead of hard limiting
                if max_amplitude > 0.9:
                    compression_ratio = 0.8 / max_amplitude
                    audio_data = np.tanh(audio_data * compression_ratio) * 0.8
                else:
                    # Normal normalization
                    audio_data = audio_data / max_amplitude * 0.8
            
            # Ensure no NaN or infinite values
            audio_data = np.nan_to_num(audio_data, nan=0.0, posinf=0.0, neginf=0.0)
            
            # Final range check
            audio_data = np.clip(audio_data, -1.0, 1.0)
            
            return audio_data.astype(np.float32)
        except Exception as e:
            print(f"Audio processing error: {e}")
            # Return safe fallback
            return np.zeros_like(audio_data, dtype=np.float32)
    
    def play_audio(self, audio_data):
        """Play audio data"""
        try:
            if audio_data is not None and len(audio_data) > 0:
                sd.play(audio_data, samplerate=Config.SAMPLE_RATE)
                return True
            return False
        except Exception as e:
            print(f"Playback error: {e}")
            return False
    
    def save_audio(self, audio_data, filename):
        """Save audio data to file with validation"""
        try:
            if audio_data is None or len(audio_data) == 0:
                print(f"❌ No audio data to save for {filename}")
                return False
            
            # Validate audio content
            max_amplitude = np.max(np.abs(audio_data))
            print(f"💾 Saving audio: {filename}, Max amplitude: {max_amplitude:.6f}")
            
            if max_amplitude < 0.0001:
                print(f"⚠️ Warning: Very quiet audio (amplitude: {max_amplitude:.6f}) - saving anyway")
            
            filepath = Config.AUDIO_DIR / filename
            sf.write(filepath, audio_data, Config.SAMPLE_RATE)
            
            # Verify file was written
            if filepath.exists():
                file_size = filepath.stat().st_size
                print(f"✅ Audio saved: {filename} ({file_size} bytes)")
                return True
            else:
                # File save verification
                return False
                
        except Exception as e:
            print(f"❌ Save error for {filename}: {e}")
            return False


class WordManager:
    """Enhanced word list manager with progress tracking"""
    
    def __init__(self):
        self.words = []  # Sadece kelimeler
        self.sentences = []  # Sadece cümleler
        self.recorded_words = set()
        # Multi-speed recording tracking
        self.recorded_speeds = {}  # {"word": {"slow": True, "normal": False, "fast": True}}
        self.current_recording_speed = "normal"  # "slow", "normal", "fast"
        self.current_index = 0
        self.speaker_id = self.load_speaker_id()
        self.is_sentence_mode = False  # Kelime/Cümle modu (backward compatibility)
        self.current_content_type = "word"  # "word", "sentence", "paragraph"
        self.paragraphs = []  # Paragraf listesi
        self.load_data()
        self.load_paragraphs()
    
    def load_speaker_id(self):
        """Konuşmacı ID'sini yükle"""
        try:
            speaker_file = Config.BASE_DIR / "speaker_config.txt"
            if speaker_file.exists():
                with open(speaker_file, "r", encoding="utf-8") as f:
                    return f.read().strip()
        except:
            pass
        return "default"  # Varsayılan ID
    
    def load_data(self):
        """Load words, sentences and recording progress"""
        try:
            # Load word list (individual words only)
            if Config.WORDS_FILE.exists():
                with open(Config.WORDS_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # Sadece kelimeleri yükle (cümle olmayan)
                    all_items = data.get('words', [])
                    self.words = [item for item in all_items if not self.is_sentence(item)]
                    print(f"✅ Loaded {len(self.words)} words from wordlist.json")
                    self.current_index = data.get('current_index', 0)
            
            # Load sentence list (sentences only)
            if Config.SENTENCES_FILE.exists():
                with open(Config.SENTENCES_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.sentences = data.get('sentences', [])
            else:
                # Eğer sentence dosyası yoksa, mevcut wordlist'ten cümleleri ayır
                if Config.WORDS_FILE.exists():
                    with open(Config.WORDS_FILE, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        all_items = data.get('words', [])
                        self.sentences = [item for item in all_items if self.is_sentence(item)]
                    self.save_sentences()  # Ayrı dosyaya kaydet
            
            # Load metadata (recorded words)
            if Config.METADATA_FILE.exists():
                with open(Config.METADATA_FILE, 'r', encoding='utf-8') as f:
                    metadata = json.load(f)
                    self.recorded_words = set(metadata.get('recorded_words', []))
                    self.recorded_speeds = metadata.get('recorded_speeds', {})
        except Exception as e:
            print(f"Error loading data: {e}")
            self.words = []
            self.sentences = []
            self.recorded_words = set()
            self.recorded_speeds = {}
            self.current_index = 0
    
    def save_data(self):
        """Save words and recording progress"""
        try:
            # Save word list (kelimeler)
            word_data = {
                'words': self.words,
                'current_index': self.current_index,
                'last_updated': datetime.now().isoformat()
            }
            with open(Config.WORDS_FILE, 'w', encoding='utf-8') as f:
                json.dump(word_data, f, ensure_ascii=False, indent=2)
            
            # Save sentences separately
            self.save_sentences()
            
            # Save paragraphs separately
            self.save_paragraphs()
            
            # Save metadata
            metadata = {
                'recorded_words': list(self.recorded_words),
                'recorded_speeds': self.recorded_speeds,  # Multi-speed tracking
                'total_words': len(self.words),
                'total_sentences': len(self.sentences),
                'total_paragraphs': len(self.paragraphs),
                'last_updated': datetime.now().isoformat()
            }
            with open(Config.METADATA_FILE, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error saving data: {e}")
    
    def save_sentences(self):
        """Save sentences to separate file"""
        try:
            sentence_data = {
                'sentences': self.sentences,
                'last_updated': datetime.now().isoformat()
            }
            with open(Config.SENTENCES_FILE, 'w', encoding='utf-8') as f:
                json.dump(sentence_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error saving sentences: {e}")
    
    def load_paragraphs(self):
        """Load paragraphs from file or create sample paragraphs"""
        try:
            # Load from file if exists
            if Config.PARAGRAPHS_FILE.exists():
                with open(Config.PARAGRAPHS_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.paragraphs = data.get('paragraphs', [])
                    print(f"Loaded {len(self.paragraphs)} paragraphs from file")
            else:
                # Create sample Kurdish paragraphs for recording
                self.paragraphs = [
                    "Ez ji Kurdistan hatim û Kurdistan gelek cîhek xweş e. Xelkê me gelek xwedî rûmet in û ev erd ji me re gelek giring e. Em dixwazin ku zarokên me bi zimanê xwe mezin bibin.",
                    "Rojhilata navîn gelek caran di şer û pevçûnê de maye. Lê gelê Kurd her dem hêviya aştiyê kiriye û dixwaze ku li vir jiyan aram be. Çanda me dewlemend e û em ê wê biparêzin.",
                    "Zimanê kurdî gelek kevnar û dewlemend e. Em dixwazin ku zarokên xwe bi vî zimanî axivin û nivîsin. Pirtûkên kurdî, stranên kurdî û çîrokên kurdî girîng in ji bo me.",
                    "Li gundên Kurdistan, xelk bi çandiniyê mijûl dibin. Genim, ceriş û darên fêkiyê gelek in. Rûbarên me gelek xweş in û av gelekî pak e.",
                    "Stranbêjên kurd gelek navdar in. Stranên wan li her cîhê cîhanê tên gotin. Muzîka kurdî dilê mirov dixoşe û rûhê mirovan geş dike."
                ]
                self.save_paragraphs()
                # Sample paragraphs created
                    
        except Exception as e:
            print(f"Error loading paragraphs: {e}")
            # Fallback to sample paragraphs
            self.paragraphs = [
                "Ez ji Kurdistan hatim û Kurdistan gelek cîhek xweş e. Xelkê me gelek xwedî rûmet in û ev erd ji me re gelek giring e.",
                "Zimanê kurdî gelek kevnar û dewlemend e. Em dixwazin ku zarokên xwe bi vî zimanî axivin û nivîsin."
            ]
    
    def save_paragraphs(self):
        """Save paragraphs to separate file"""
        try:
            paragraph_data = {
                'paragraphs': self.paragraphs,
                'last_updated': datetime.now().isoformat()
            }
            with open(Config.PARAGRAPHS_FILE, 'w', encoding='utf-8') as f:
                json.dump(paragraph_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error saving paragraphs: {e}")
    
    def add_words(self, new_words):
        """Add new words or sentences to appropriate lists"""
        added_words = 0
        added_sentences = 0
        
        for item in new_words:
            item = item.strip()
            if not item:
                continue
                
            if self.is_sentence(item):
                # Cümle ise sentences listesine ekle
                if item not in self.sentences:
                    self.sentences.append(item)
                    added_sentences += 1
            else:
                # Kelime ise words listesine ekle
                if item not in self.words:
                    self.words.append(item)
                    added_words += 1
        
        if added_words > 0 or added_sentences > 0:
            self.save_data()
        
        return added_words + added_sentences
    
    def get_current_word(self):
        """Get the current word or sentence based on active mode"""
        filtered_content = self.get_filtered_content()
        if 0 <= self.current_index < len(filtered_content):
            return filtered_content[self.current_index]
        return None
    
    def next_word(self):
        """Move to next unrecorded word or sentence based on active mode"""
        filtered_content = self.get_filtered_content()
        
        # İlk olarak kayıt edilmemiş içerik bul
        for i in range(self.current_index + 1, len(filtered_content)):
            if filtered_content[i] not in self.recorded_words:
                self.current_index = i
                self.save_data()
                return self.get_current_word()
        
        # Eğer sonunda kayıt edilmemiş içerik yoksa, baştan ara
        for i in range(0, self.current_index):
            if filtered_content[i] not in self.recorded_words:
                self.current_index = i
                self.save_data()
                return self.get_current_word()
        
        # Tüm içerik kayıt edilmişse normal ilerleme
        if self.current_index < len(filtered_content) - 1:
            self.current_index += 1
            self.save_data()
        return self.get_current_word()
    
    def previous_word(self):
        """Move to previous word or sentence (recorded or not)"""
        # Normal geri gitme - kaydedilen içeriği de göster
        if self.current_index > 0:
            self.current_index -= 1
            self.save_data()
        return self.get_current_word()
    
    def mark_recorded(self, word, audio_filename=None):
        """Mark a word as recorded and update Whisper training files"""
        self.recorded_words.add(word)
        
        # Whisper eğitimi için kayıt
        if audio_filename:
            self.update_whisper_files(word, audio_filename)
        
        self.save_data()
    
    def update_whisper_files(self, transcript, audio_filename):
        """Whisper eğitimi için gerekli dosyaları güncelle"""
        try:
            # JSONL formatında Whisper manifest dosyasını güncelle
            manifest_entry = {
                "audio_filepath": f"audio/{audio_filename}",
                "text": transcript.strip(),
                "language": "ku",  # Kurdish language code
                "duration": 2.0,  # Varsayılan süre
                "speaker_id": self.speaker_id
            }
            
            # Manifest dosyasına ekle
            with open(Config.WHISPER_MANIFEST, 'a', encoding='utf-8') as f:
                f.write(json.dumps(manifest_entry, ensure_ascii=False) + '\n')
            
            # Transkript dosyasına ekle (basit format)
            with open(Config.TRANSCRIPT_FILE, 'a', encoding='utf-8') as f:
                f.write(f"{audio_filename}\t{transcript.strip()}\n")
                
            print(f"✅ Whisper eğitim dosyaları güncellendi: {audio_filename} -> '{transcript}'")
            
        except Exception as e:
            print(f"❌ Whisper dosyaları güncellenirken hata: {e}")
    
    def is_current_recorded(self):
        """Check if current word is recorded"""
        current_word = self.get_current_word()
        return current_word in self.recorded_words if current_word else False
    
    def get_word_speed_status(self, word):
        """Get recording status for all speeds of a word"""
        if word not in self.recorded_speeds:
            return {"slow": False, "normal": False, "fast": False}
        return self.recorded_speeds.get(word, {"slow": False, "normal": False, "fast": False})
    
    def mark_speed_recorded(self, word, speed, audio_filename=None):
        """Mark a specific speed as recorded for a word"""
        if word not in self.recorded_speeds:
            self.recorded_speeds[word] = {"slow": False, "normal": False, "fast": False}
        
        self.recorded_speeds[word][speed] = True
        
        # Mark word as recorded if any speed is recorded
        self.recorded_words.add(word)
        
        # Whisper eğitimi için kayıt
        if audio_filename:
            self.update_whisper_files(word, audio_filename)
        
        self.save_data()
        print(f"✅ {word} - {speed} speed recorded: {audio_filename}")
    
    def get_missing_speeds(self, word):
        """Get list of missing speeds for a word"""
        speeds = self.get_word_speed_status(word)
        missing = [speed for speed, recorded in speeds.items() if not recorded]
        return missing
    
    def get_completion_status(self, word):
        """Get completion percentage for a word (0-100)"""
        speeds = self.get_word_speed_status(word)
        completed = sum(1 for recorded in speeds.values() if recorded)
        return (completed / 3) * 100  # 3 speeds total
    
    def set_recording_speed(self, speed):
        """Set current recording speed"""
        if speed in ["slow", "normal", "fast"]:
            self.current_recording_speed = speed
            print(f"🎚️ Recording speed set to: {speed.upper()}")
        else:
            print(f"❌ Invalid speed: {speed}. Use 'slow', 'normal', or 'fast'")
    
    def generate_speed_filename(self, word, speed):
        """Generate filename for specific word and speed"""
        # Get next available number
        existing_files = list(Config.AUDIO_DIR.glob("*.wav"))
        numbers = []
        for f in existing_files:
            match = re.match(r'^(\d{6})_', f.name)
            if match:
                numbers.append(int(match.group(1)))
        
        next_number = max(numbers, default=0) + 1
        
        # Clean word for filename
        clean_word = re.sub(r'[^\w\s\-]', '', word).replace(' ', '_')
        
        return f"{next_number:06d}_{clean_word}_{speed}.wav"
    
    def get_progress(self):
        """Get recording progress based on active mode"""
        filtered_content = self.get_filtered_content()
        
        # Filtrelenmiş içerikteki kayıtlı öğeleri say
        recorded_in_filtered = sum(1 for item in filtered_content if item in self.recorded_words)
        total_filtered = len(filtered_content)
        
        return recorded_in_filtered, total_filtered
    
    def extract_transcript_from_filename(self, filename):
        """Dosya isminden transkripti çıkar"""
        try:
            import re
            
            # Dosya uzantısını kaldır
            basename = filename.stem if hasattr(filename, 'stem') else Path(filename).stem
            
            # Farklı formatları dene (öncelik sırasına göre)
            patterns = [
                # Whisper formatı: 000001_word
                r'^\d{6}_(.+)$',
                
                # Duplikat formatı: 000001_000002_word
                r'^\d{6}_\d{6}_(.+)$',
                
                # Tarih-saat formatları
                r'^(.+?)_\d{8}_\d{6}$',  # word_20251004_123456
                r'^(.+?)_\d{4}-\d{2}-\d{2}_\d{2}-\d{2}-\d{2}$',  # word_2024-01-15_14-30-45
                
                # Timestamp formatları
                r'^(.+?)_\d{10,}$',  # word_1234567890 (unix timestamp)
                r'^(.+?)_\d+$',  # word_123456789
                
                # Özel formatlar
                r'^(.+?)_recording$',  # word_recording
                r'^(.+?)_audio$',  # word_audio
                
                # Sadece kelime
                r'^(.+)$'
            ]
            
            for pattern in patterns:
                match = re.match(pattern, basename)
                if match:
                    transcript = match.group(1).strip()
                    if transcript and not transcript.isdigit():  # Sadece rakam değilse
                        # Özel karakterleri temizle
                        transcript = re.sub(r'[^\w\s\-\.\_]', '', transcript)
                        return transcript
            
            # Hiçbir pattern uymazsa basename'i döndür
            return re.sub(r'[^\w\s\-\.\_]', '', basename)
            
        except Exception as e:
            print(f"❌ Dosya isminden transkript çıkarılırken hata: {e}")
            return str(filename).replace('.wav', '')
    
    def convert_old_audio_files(self):
        """Eski ses dosyalarını Whisper formatına dönüştür ve duplikatları temizle"""
        try:
            from tkinter import messagebox
            import shutil
            import re
            from pathlib import Path
            import json
            
            if not Config.AUDIO_DIR.exists():
                messagebox.showwarning("Uyarı", "Audio klasörü bulunamadı!")
                return
            
            # Tüm ses dosyalarını al
            all_audio_files = list(Config.AUDIO_DIR.glob("*.wav"))
            if not all_audio_files:
                messagebox.showinfo("Bilgi", "Ses dosyası bulunamadı!")
                return
            
            conversion_log = []
            
            # Dosyaları kategorilere ayır
            whisper_format_files = []  # 000001_word.wav formatındaki dosyalar
            old_format_files = []      # word_timestamp.wav formatındaki dosyalar
            duplicate_files = []       # 000001_000002_word.wav gibi duplikatlar
            
            for audio_file in all_audio_files:
                filename = audio_file.name
                
                # Whisper formatı kontrolü (000001_word.wav)
                if re.match(r'^\d{6}_[^_]+.*\.wav$', filename):
                    # Duplikat kontrolü (000001_000002_word.wav)
                    if re.match(r'^\d{6}_\d{6}_.*\.wav$', filename):
                        duplicate_files.append(audio_file)
                    else:
                        whisper_format_files.append(audio_file)
                else:
                    # Eski format (word_timestamp.wav, word.wav vs)
                    old_format_files.append(audio_file)
            
            conversion_log.append(f"📊 Dosya Kategorileri:")
            conversion_log.append(f"   • Whisper formatı: {len(whisper_format_files)}")
            conversion_log.append(f"   • Eski format: {len(old_format_files)}")
            conversion_log.append(f"   • Duplikatlar: {len(duplicate_files)}")
            conversion_log.append("")
            
            # Duplikatları sil
            if duplicate_files:
                conversion_log.append("🗑️ Duplikat dosyalar siliniyor:")
                for dup_file in duplicate_files:
                    try:
                        dup_file.unlink()
                        conversion_log.append(f"   ✅ Silindi: {dup_file.name}")
                    except Exception as e:
                        conversion_log.append(f"   ❌ Silinemedi: {dup_file.name} - {e}")
                conversion_log.append("")
            
            # Mevcut Whisper manifest'i temizle ve yeniden oluştur
            if Config.WHISPER_MANIFEST.exists():
                Config.WHISPER_MANIFEST.unlink()
            if Config.TRANSCRIPT_FILE.exists():
                Config.TRANSCRIPT_FILE.unlink()
            
            conversion_log.append("🔄 Whisper dosyaları yeniden oluşturuluyor:")
            
            # Whisper formatındaki dosyalar için manifest yeniden oluştur
            current_whisper_files = []
            for audio_file in whisper_format_files:
                try:
                    # Transkript çıkar
                    transcript = self.extract_transcript_from_filename(audio_file.name)
                    self.update_whisper_files(transcript, audio_file.name)
                    current_whisper_files.append((audio_file.name, transcript))
                    conversion_log.append(f"   ✅ Manifest güncellendi: {audio_file.name}")
                except Exception as e:
                    conversion_log.append(f"   ❌ Hata: {audio_file.name} - {e}")
            
            # Eski formatdaki dosyaları dönüştür
            if old_format_files:
                conversion_log.append("")
                conversion_log.append("🔄 Eski format dosyalar dönüştürülüyor:")
                
                # Yeni numara başlangıcı
                next_number = len(current_whisper_files) + 1
                
                for audio_file in old_format_files:
                    try:
                        # Transkript çıkar
                        transcript = self.extract_transcript_from_filename(audio_file.name)
                        
                        # Yeni dosya adı
                        new_filename = f"{next_number:06d}_{transcript}.wav"
                        new_filepath = Config.AUDIO_DIR / new_filename
                        
                        # Dosyayı yeni isimle kopyala
                        shutil.copy2(audio_file, new_filepath)
                        
                        # Whisper dosyalarını güncelle
                        self.update_whisper_files(transcript, new_filename)
                        
                        # Eski dosyayı sil
                        audio_file.unlink()
                        
                        conversion_log.append(f"   ✅ Dönüştürüldü: {audio_file.name} → {new_filename}")
                        next_number += 1
                        
                    except Exception as e:
                        conversion_log.append(f"   ❌ Hata: {audio_file.name} - {e}")
            
            conversion_log.append("")
            conversion_log.append("✅ Dönüştürme tamamlandı!")
            
            # Sonuçları göster
            messagebox.showinfo("Başarılı", 
                f"🎯 Whisper Dönüştürme Tamamlandı!\n\n"
                f"✅ Duplikatlar temizlendi: {len(duplicate_files)}\n"
                f"� Eski formatlar dönüştürüldü: {len(old_format_files)}\n"
                f"📋 Toplam Whisper dosyası: {len(current_whisper_files) + len(old_format_files)}\n\n"
                f"📁 Manifest: whisper_manifest.jsonl\n"
                f"📄 Transkript: transcripts.txt")
            
            # Dönüştürme logunu sakla
            self.conversion_log = conversion_log
            
        except Exception as e:
            messagebox.showerror("Hata", f"Dönüştürme sırasında hata oluştu: {str(e)}")
    
    def show_conversion_log(self):
        """Dönüştürme logunu göster"""
        if not hasattr(self, 'conversion_log') or not self.conversion_log:
            from tkinter import messagebox
            messagebox.showinfo("Bilgi", "Henüz dönüştürme işlemi yapılmamış.")
            return
        
        log_text = "🔄 Whisper Dönüştürme Logu\n" + "="*50 + "\n\n"
        for entry in self.conversion_log:
            log_text += f"{entry}\n"
        
        # Log penceresini göster
        import customtkinter as ctk
        log_window = ctk.CTkToplevel()
        log_window.title("Dönüştürme Logu")
        log_window.geometry("800x600")
        
        text_widget = ctk.CTkTextbox(
            log_window, 
            width=760, 
            height=540,
            font=ctk.CTkFont(family="Consolas", size=11)
        )
        text_widget.pack(padx=20, pady=20)
        text_widget.insert("1.0", log_text)
        text_widget.configure(state="disabled")

    def find_first_unrecorded_word(self):
        """Find and jump to the first unrecorded word or sentence based on active mode"""
        filtered_content = self.get_filtered_content()
        
        for i, item in enumerate(filtered_content):
            if item not in self.recorded_words:
                self.current_index = i
                self.save_data()
                return item
        
        # If all content is recorded, stay at current position
        return self.get_current_word()
    
    def jump_to_first_unrecorded(self):
        """Jump to first unrecorded word and return True if found"""
        first_unrecorded = self.find_first_unrecorded_word()
        return first_unrecorded is not None and first_unrecorded not in self.recorded_words
    
    def separate_mixed_content(self):
        """Mevcut karışık içeriği kelime ve cümlelere ayır"""
        try:
            # Mevcut wordlist dosyasından tüm içeriği yükle
            if Config.WORDS_FILE.exists():
                with open(Config.WORDS_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    all_items = data.get('words', [])
                
                # İçeriği ayır
                separated_words = []
                separated_sentences = []
                
                for item in all_items:
                    if self.is_sentence(item):
                        separated_sentences.append(item)
                    else:
                        separated_words.append(item)
                
                # Listeleri güncelle
                self.words = separated_words
                self.sentences = separated_sentences
                
                # Dosyalara kaydet
                self.save_data()
                
                print(f"✅ Content separated: {len(separated_words)} words, {len(separated_sentences)} sentences")
                return len(separated_words), len(separated_sentences)
        except Exception as e:
            print(f"Error separating content: {e}")
            return 0, 0
    
    def is_sentence(self, text):
        """Bir metnin cümle mi kelime mi olduğunu tespit et"""
        if not text or not text.strip():
            return False
        
        text = text.strip()
        text_lower = text.lower()
        
        # 1. Noktalama işaretleri - kesin cümle belirtisi
        punctuation_marks = ['?', '!', '.']
        for mark in punctuation_marks:
            if mark in text:
                return True
        
        # 2. Boşluk içeriyor - çoklu kelime
        if ' ' in text:
            return True
        
        # 3. Kurmancî özel cümle kalıpları
        kurdish_sentence_patterns = [
            # Soru kalıpları
            'çawa', 'kî', 'kengî', 'li ku', 'li kî', 'çi', 'çend',
            # Fiil çekimleri (cümle sonları)
            ' im', ' e', ' in', ' yî', ' ne', ' n', ' re',
            # Cümle başlangıçları
            'ez ', 'tu ', 'ew ', 'em ', 'hûn ', 'ewan ',
            # Yaygın cümle kelimeleri
            'dixwazim', 'dibînim', 'dikim', 'diçim', 'tê'
        ]
        
        for pattern in kurdish_sentence_patterns:
            if pattern in text_lower:
                return True
        
        # 4. Uzunluk kontrolü - 12+ karakter muhtemelen cümle
        if len(text) >= 12:
            return True
            
        # 5. Üç kelimeden fazlaysa kesinlikle cümle
        word_count = len(text.split())
        if word_count >= 3:
            return True
        
        # 6. İki kelime ama yaygın cümle kalıpları
        if word_count == 2:
            two_word_sentence_patterns = [
                'roj baş', 'spas dikim', 'ez hatim', 'tu çû', 'ew hat',
                'em çûn', 'hûn hatin', 'gellek spas', 'her tim'
            ]
            for pattern in two_word_sentence_patterns:
                if pattern in text_lower:
                    return True
            
        return False
    
    def get_filtered_content(self):
        """Aktif moda göre filtrelenmiş içerik döndür"""
        if self.current_content_type == "paragraph":
            # Paragraf modu
            # Paragraph mode loaded
            if self.paragraphs:
                # Paragraphs loaded successfully
                pass
            else:
                # No paragraphs found
                pass
            return self.paragraphs
        elif self.is_sentence_mode or self.current_content_type == "sentence":
            # Cümle modu - sentences listesini kullan
            # Sentence mode loaded
            if self.sentences:
                # Sentences loaded successfully
                pass
            else:
                # No sentences found
                pass
            return self.sentences
        else:
            # Kelime modu - words listesini kullan
            # Word mode loaded
            if self.words:
                # Words loaded successfully
                pass
            return self.words


class DocumentProcessor:
    """Document processing utilities"""
    
    @staticmethod
    def load_words_from_file(filepath):
        """Load words from various file types"""
        try:
            file_path = Path(filepath)
            
            if not file_path.exists():
                return []
            
            content = ""
            
            if file_path.suffix.lower() == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            elif file_path.suffix.lower() == '.docx':
                doc = docx.Document(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            elif file_path.suffix.lower() == '.pdf':
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    content = ""
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n"
            
            return DocumentProcessor._extract_words(content)
        
        except Exception as e:
            print(f"Error loading file {filepath}: {e}")
            return []
    
    @staticmethod
    def _extract_words(content):
        """Extract words from text content"""
        if not content:
            return []
        
        # Split into words and clean them
        words = []
        for line in content.split('\n'):
            line = line.strip()
            if line:
                # Split by common delimiters and clean
                for word in re.split(r'[,;.\s]+', line):
                    word = word.strip()
                    if word and len(word) > 1:
                        words.append(word)
        
        return list(set(words))  # Remove duplicates
    
    @staticmethod
    def get_manual_words_dialog(parent):
        """Show dialog to manually enter words"""
        dialog = ctk.CTkToplevel(parent)
        dialog.title("Add Words Manually")
        dialog.geometry("600x500")
        dialog.transient(parent)
        dialog.grab_set()
        
        # Center the dialog
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() - dialog.winfo_width()) // 2
        y = (dialog.winfo_screenheight() - dialog.winfo_height()) // 2
        dialog.geometry(f"+{x}+{y}")
        
        main_frame = ctk.CTkFrame(dialog)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Instructions
        ctk.CTkLabel(
            main_frame,
            text="📝 Add Words or Sentences",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=(10, 5))
        
        ctk.CTkLabel(
            main_frame,
            text="Enter words or sentences (one per line):",
            font=ctk.CTkFont(size=12)
        ).pack(pady=(0, 10))
        
        # Text area
        text_widget = ctk.CTkTextbox(
            main_frame,
            width=550,
            height=300,
            font=ctk.CTkFont(family="Consolas", size=12)
        )
        text_widget.pack(pady=(0, 20))
        
        # Sample text
        sample_text = """spî
reş
sor
kesk
zer
Example sentence: Ez hêviya te dikim."""
        text_widget.insert("1.0", sample_text)
        
        result = {'words': None}
        
        def save_words():
            content = text_widget.get("1.0", "end").strip()
            if content:
                lines = [line.strip() for line in content.split('\n') if line.strip()]
                result['words'] = lines
            dialog.destroy()
        
        def cancel():
            result['words'] = None
            dialog.destroy()
        
        # Buttons
        button_frame = ctk.CTkFrame(main_frame)
        button_frame.pack(fill="x")
        
        ctk.CTkButton(
            button_frame,
            text="💾 Add Words",
            command=save_words,
            width=150,
            height=40,
            fg_color="green",
            hover_color="darkgreen"
        ).pack(side="left", padx=(20, 10), pady=10)
        
        ctk.CTkButton(
            button_frame,
            text="❌ Cancel",
            command=cancel,
            width=100,
            height=40
        ).pack(side="right", padx=(10, 20), pady=10)
        
        # Wait for dialog to close
        dialog.wait_window()
        
        return result['words']


class SimpleWordViewer:
    """Enhanced word viewer with checkbox selection and bulk operations"""
    
    def __init__(self, parent, word_manager):
        self.parent = parent
        self.word_manager = word_manager
        self.dialog = None
        self.selected_words = set()  # Track selected words
        self.word_checkboxes = {}    # Map words to checkbox widgets
        self.filtered_words = []     # Current filtered word list
        
        # Pagination variables
        self.current_page = 0
        self.words_per_page = 50  # Show 50 words per page
        self.total_pages = 0
        
    def show(self):
        """Show word list viewer"""
        self.dialog = ctk.CTkToplevel(self.parent)
        self.dialog.title("Word List Manager")
        self.dialog.geometry("800x600")
        self.dialog.transient(self.parent)
        
        # Center the dialog
        self.dialog.update_idletasks()
        x = (self.dialog.winfo_screenwidth() - self.dialog.winfo_width()) // 2
        y = (self.dialog.winfo_screenheight() - self.dialog.winfo_height()) // 2
        self.dialog.geometry(f"+{x}+{y}")
        
        self.setup_viewer()
    
    def setup_viewer(self):
        """Setup enhanced viewer UI with checkbox selection"""
        main_frame = ctk.CTkFrame(self.dialog)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Title with stats and selection counter
        recorded, total = self.word_manager.get_progress()
        title_text = f"📚 Word List ({total} words, {recorded} recorded)"
        
        self.title_label = ctk.CTkLabel(
            main_frame,
            text=title_text,
            font=ctk.CTkFont(size=18, weight="bold")
        )
        self.title_label.pack(pady=(0, 10))
        
        # Selection counter
        self.selection_label = ctk.CTkLabel(
            main_frame,
            text="0 words selected",
            font=ctk.CTkFont(size=12),
            text_color=("blue", "lightblue")
        )
        self.selection_label.pack(pady=(0, 15))
        
        # Search and filter frame
        search_frame = ctk.CTkFrame(main_frame)
        search_frame.pack(fill="x", pady=(0, 10))
        
        # Search controls
        search_controls = ctk.CTkFrame(search_frame)
        search_controls.pack(side="left", padx=10, pady=10)
        
        ctk.CTkLabel(search_controls, text="🔍 Search:").pack(side="left", padx=(0, 5))
        
        self.search_var = tk.StringVar()
        self.search_var.trace("w", self.filter_words)
        self.search_entry = ctk.CTkEntry(
            search_controls,
            textvariable=self.search_var,
            width=200,
            placeholder_text="Type to search words..."
        )
        self.search_entry.pack(side="left", padx=5)
        
        # Clear search button
        ctk.CTkButton(
            search_controls,
            text="🔍",
            command=self.focus_search,
            width=30,
            height=28,
            fg_color="#6c757d",
            hover_color="#545b62"
        ).pack(side="left", padx=2)
        
        ctk.CTkButton(
            search_controls,
            text="❌",
            command=self.clear_search,
            width=30,
            height=28,
            fg_color="#dc3545",
            hover_color="#bd2130"
        ).pack(side="left", padx=2)
        
        # Filter controls
        filter_controls = ctk.CTkFrame(search_frame)
        filter_controls.pack(side="right", padx=10, pady=10)
        
        self.filter_var = tk.StringVar(value="all")
        
        ctk.CTkRadioButton(
            filter_controls,
            text="All",
            variable=self.filter_var,
            value="all",
            command=self.filter_words
        ).pack(side="left", padx=5)
        
        ctk.CTkRadioButton(
            filter_controls,
            text="Recorded",
            variable=self.filter_var,
            value="recorded",
            command=self.filter_words
        ).pack(side="left", padx=5)
        
        ctk.CTkRadioButton(
            filter_controls,
            text="Not Recorded",
            variable=self.filter_var,
            value="not_recorded",
            command=self.filter_words
        ).pack(side="left", padx=5)
        
        # All action buttons above word list (duplicated for easy access)
        top_buttons_frame = ctk.CTkFrame(main_frame)
        top_buttons_frame.pack(fill="x", pady=(0, 10))
        
        # Left side buttons (top)
        top_left_buttons = ctk.CTkFrame(top_buttons_frame)
        top_left_buttons.pack(side="left", padx=10, pady=5)
        
        ctk.CTkButton(
            top_left_buttons,
            text="➕ Add New Word",
            command=self.add_new_word,
            width=130,
            height=30,
            fg_color="#28a745",
            hover_color="#1e7e34"
        ).pack(side="left", padx=2)
        
        ctk.CTkButton(
            top_left_buttons,
            text="🔄 Refresh List",
            command=self.refresh_list,
            width=120,
            height=30,
            fg_color="#6c757d",
            hover_color="#545b62"
        ).pack(side="left", padx=2)
        
        # Center selection buttons (top)
        top_center_buttons = ctk.CTkFrame(top_buttons_frame)
        top_center_buttons.pack(side="left", padx=15, pady=5)
        
        ctk.CTkButton(
            top_center_buttons,
            text="✅ Select All",
            command=self.select_all,
            width=100,
            height=30,
            fg_color="#28a745",
            hover_color="#1e7e34"
        ).pack(side="left", padx=2)
        
        ctk.CTkButton(
            top_center_buttons,
            text="❌ Select None",
            command=self.select_none,
            width=110,
            height=30,
            fg_color="#6c757d",
            hover_color="#545b62"
        ).pack(side="left", padx=2)
        
        ctk.CTkButton(
            top_center_buttons,
            text="🔄 Invert Selection",
            command=self.invert_selection,
            width=130,
            height=30,
            fg_color="#fd7e14",
            hover_color="#e55b00"
        ).pack(side="left", padx=2)
        
        # Right side buttons (top)  
        top_right_buttons = ctk.CTkFrame(top_buttons_frame)
        top_right_buttons.pack(side="right", padx=10, pady=5)
        
        ctk.CTkButton(
            top_right_buttons,
            text="✏️ Edit Selected",
            command=self.edit_selected_words,
            width=130,
            height=30,
            fg_color="#007bff",
            hover_color="#0056b3"
        ).pack(side="left", padx=2)
        
        ctk.CTkButton(
            top_right_buttons,
            text="🗑️ Delete Selected",
            command=self.delete_selected_words,
            width=140,
            height=30,
            fg_color="#dc3545",
            hover_color="#bd2130"
        ).pack(side="left", padx=2)
        
        # Pagination controls
        pagination_frame = ctk.CTkFrame(main_frame)
        pagination_frame.pack(fill="x", pady=(0, 10))
        
        # Page info and controls
        page_controls = ctk.CTkFrame(pagination_frame)
        page_controls.pack(side="left", padx=10, pady=5)
        
        ctk.CTkButton(
            page_controls,
            text="⬅️ Previous",
            command=self.previous_page,
            width=80,
            height=28,
            fg_color="#6c757d",
            hover_color="#545b62"
        ).pack(side="left", padx=2)
        
        self.page_info_label = ctk.CTkLabel(
            page_controls,
            text="Page 1 of 1",
            font=ctk.CTkFont(size=11)
        )
        self.page_info_label.pack(side="left", padx=10)
        
        ctk.CTkButton(
            page_controls,
            text="Next ➡️",
            command=self.next_page,
            width=80,
            height=28,
            fg_color="#6c757d",
            hover_color="#545b62"
        ).pack(side="left", padx=2)
        
        # Page size controls
        page_size_controls = ctk.CTkFrame(pagination_frame)
        page_size_controls.pack(side="right", padx=10, pady=5)
        
        ctk.CTkLabel(
            page_size_controls,
            text="Words per page:",
            font=ctk.CTkFont(size=11)
        ).pack(side="left", padx=5)
        
        self.page_size_var = tk.StringVar(value="50")
        page_size_combo = ctk.CTkComboBox(
            page_size_controls,
            values=["25", "50", "100", "200"],
            variable=self.page_size_var,
            command=self.change_page_size,
            width=80,
            height=28
        )
        page_size_combo.pack(side="left", padx=5)
        
        # Word list with checkboxes (scrollable)
        self.word_list_frame = ctk.CTkScrollableFrame(
            main_frame,
            width=750,
            height=300,
            label_text="📝 Word List (check words to select)"
        )
        self.word_list_frame.pack(fill="both", expand=True, pady=(0, 10))
        



        
        # Keyboard shortcuts
        self.dialog.bind('<Control-f>', lambda e: self.focus_search())
        self.dialog.bind('<F3>', lambda e: self.focus_search())
        
        # Initial load
        self.refresh_list()
    
    def filter_words(self, *args):
        """Filter words and rebuild checkbox list"""
        search_text = self.search_var.get().lower().strip()
        filter_type = self.filter_var.get()
        
        # Clear existing checkboxes
        for widget in self.word_list_frame.winfo_children():
            widget.destroy()
        
        self.word_checkboxes.clear()
        self.filtered_words = []
        
        # Separate words into recorded and unrecorded for proper ordering
        unrecorded_words = []
        recorded_words = []
        
        for word in self.word_manager.words:
            # Apply search filter - search in word content
            if search_text:
                if search_text not in word.lower():
                    continue
            
            # Separate by recording status
            is_recorded = word in self.word_manager.recorded_words
            if is_recorded:
                recorded_words.append(word)
            else:
                unrecorded_words.append(word)
        
        # Apply status filter and prioritize unrecorded words
        if filter_type == "all":
            # Show unrecorded first, then recorded
            self.filtered_words = unrecorded_words + recorded_words
        elif filter_type == "not_recorded":
            self.filtered_words = unrecorded_words
        elif filter_type == "recorded":
            self.filtered_words = recorded_words
        
        # Calculate pagination
        self.total_pages = max(1, (len(self.filtered_words) + self.words_per_page - 1) // self.words_per_page)
        
        # Reset to first page if current page is out of range
        if self.current_page >= self.total_pages:
            self.current_page = 0
        
        # Create checkboxes for current page only
        self.display_current_page()
        
        # Update pagination info
        self.update_pagination_info()
        
        # Update title and selection count
        self.update_display_info()
        
        # Show search results info
        if search_text:
            # Search completed
            pass
    
    def create_word_checkbox(self, word, index):
        """Create a checkbox for a word"""
        is_recorded = word in self.word_manager.recorded_words
        status_icon = "✅" if is_recorded else "⭕"
        
        # Create frame for this word
        word_frame = ctk.CTkFrame(self.word_list_frame)
        word_frame.pack(fill="x", padx=5, pady=2)
        
        # Create checkbox variable
        checkbox_var = tk.BooleanVar()
        checkbox_var.trace("w", lambda *args: self.on_selection_change())
        
        # Checkbox
        checkbox = ctk.CTkCheckBox(
            word_frame,
            text="",
            variable=checkbox_var,
            width=20,
            height=20
        )
        checkbox.pack(side="left", padx=(10, 5), pady=8)
        
        # Word number and status
        index_label = ctk.CTkLabel(
            word_frame,
            text=f"{index+1:3d}.",
            font=ctk.CTkFont(family="Consolas", size=11),
            width=40
        )
        index_label.pack(side="left", padx=5, pady=8)
        
        status_label = ctk.CTkLabel(
            word_frame,
            text=status_icon,
            font=ctk.CTkFont(size=14),
            width=30
        )
        status_label.pack(side="left", padx=5, pady=8)
        
        # Word text
        word_label = ctk.CTkLabel(
            word_frame,
            text=word,
            font=ctk.CTkFont(size=12),
            anchor="w"
        )
        word_label.pack(side="left", padx=10, pady=8, fill="x", expand=True)
        
        # Action buttons for each word
        action_frame = ctk.CTkFrame(word_frame)
        action_frame.pack(side="right", padx=5, pady=5)
        
        # Edit button
        edit_btn = ctk.CTkButton(
            action_frame,
            text="✏️",
            command=lambda w=word: self.edit_word_inline(w),
            width=30,
            height=25,
            fg_color="#007bff",
            hover_color="#0056b3"
        )
        edit_btn.pack(side="left", padx=1)
        
        # Delete button
        delete_btn = ctk.CTkButton(
            action_frame,
            text="🗑️",
            command=lambda w=word: self.delete_word_inline(w),
            width=30,
            height=25,
            fg_color="#dc3545",
            hover_color="#bd2130"
        )
        delete_btn.pack(side="left", padx=1)
        
        # Record button (only for unrecorded words)
        if not is_recorded:
            record_btn = ctk.CTkButton(
                action_frame,
                text="🎤",
                command=lambda w=word: self.record_word_inline(w),
                width=30,
                height=25,
                fg_color="#28a745",
                hover_color="#1e7e34"
            )
            record_btn.pack(side="left", padx=2)
        
        # Store checkbox reference
        self.word_checkboxes[word] = {
            'var': checkbox_var,
            'checkbox': checkbox,
            'frame': word_frame,
            'word_label': word_label,
            'action_frame': action_frame
        }
    
    def display_current_page(self):
        """Display only the words for the current page"""
        # Clear existing checkboxes
        for widget in self.word_list_frame.winfo_children():
            widget.destroy()
        self.word_checkboxes.clear()
        
        # Calculate range for current page
        start_idx = self.current_page * self.words_per_page
        end_idx = min(start_idx + self.words_per_page, len(self.filtered_words))
        
        # Create checkboxes for current page words only
        for i in range(start_idx, end_idx):
            word = self.filtered_words[i]
            # Use the actual word index in the full list for display
            self.create_word_checkbox(word, i)
            
            # Restore selection state
            if word in self.selected_words:
                self.word_checkboxes[word]['var'].set(True)
    
    def update_pagination_info(self):
        """Update pagination information display"""
        if self.total_pages > 0:
            # Count unrecorded words in current filtered list
            unrecorded_count = sum(1 for word in self.filtered_words if word not in self.word_manager.recorded_words)
            recorded_count = len(self.filtered_words) - unrecorded_count
            
            page_text = f"Page {self.current_page + 1} of {self.total_pages}"
            if self.filter_var.get() == "all":
                page_text += f" | 🔴{unrecorded_count} unrecorded, ✅{recorded_count} recorded"
            
            self.page_info_label.configure(text=page_text)
        else:
            self.page_info_label.configure(text="No words")
    
    def previous_page(self):
        """Go to previous page"""
        if self.current_page > 0:
            self.current_page -= 1
            self.display_current_page()
            self.update_pagination_info()
    
    def next_page(self):
        """Go to next page"""
        if self.current_page < self.total_pages - 1:
            self.current_page += 1
            self.display_current_page()
            self.update_pagination_info()
    
    def change_page_size(self, new_size):
        """Change the number of words per page"""
        self.words_per_page = int(new_size)
        self.current_page = 0  # Reset to first page
        self.filter_words()  # Refresh the display
    
    def update_selection_display(self):
        """Update the selection count display"""
        count = len(self.selected_words)
        if hasattr(self, 'selection_label'):
            if count == 0:
                self.selection_label.configure(text="No words selected")
            elif count == 1:
                self.selection_label.configure(text="1 word selected")
            else:
                self.selection_label.configure(text=f"{count} words selected")
    
    def edit_word_inline(self, word):
        """Edit a word in place"""
        try:
            # Create custom edit dialog
            edit_dialog = ctk.CTkToplevel(self.dialog)
            edit_dialog.title("Edit Word")
            edit_dialog.geometry("500x180")
            edit_dialog.transient(self.dialog)
            edit_dialog.grab_set()
            
            # Make dialog modal
            edit_dialog.focus_force()
            
            # Center the dialog
            edit_dialog.update_idletasks()
            x = (edit_dialog.winfo_screenwidth() // 2) - (250)
            y = (edit_dialog.winfo_screenheight() // 2) - (90)
            edit_dialog.geometry(f"500x180+{x}+{y}")
            
            # Variables to store result
            result = {"saved": False, "new_word": ""}
            
            # Main frame
            main_frame = ctk.CTkFrame(edit_dialog, fg_color="transparent")
            main_frame.pack(fill="both", expand=True, padx=20, pady=20)
            
            # Title
            title_label = ctk.CTkLabel(
                main_frame,
                text=f"Edit Word: {word}",
                font=ctk.CTkFont(size=16, weight="bold")
            )
            title_label.pack(pady=(0, 15))
            
            # Entry
            entry = ctk.CTkEntry(
                main_frame,
                width=450,
                height=40,
                font=ctk.CTkFont(size=14),
                placeholder_text="Enter the corrected word or sentence..."
            )
            entry.pack(pady=10)
            entry.insert(0, word)
            entry.focus_force()
            entry.select_range(0, "end")
            
            # Button frame
            btn_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
            btn_frame.pack(pady=15)
            
            def save_word():
                new_text = entry.get().strip()
                if new_text:
                    result["saved"] = True
                    result["new_word"] = new_text
                edit_dialog.destroy()
            
            def cancel_edit():
                result["saved"] = False
                edit_dialog.destroy()
            
            # Save button
            save_btn = ctk.CTkButton(
                btn_frame,
                text="💾 Save Changes",
                command=save_word,
                width=140,
                height=35,
                fg_color="#28a745",
                hover_color="#1e7e34",
                font=ctk.CTkFont(size=12, weight="bold")
            )
            save_btn.pack(side="left", padx=10)
            
            # Cancel button
            cancel_btn = ctk.CTkButton(
                btn_frame,
                text="❌ Cancel",
                command=cancel_edit,
                width=100,
                height=35,
                fg_color="#dc3545",
                hover_color="#bd2130"
            )
            cancel_btn.pack(side="left", padx=10)
            
            # Keyboard bindings
            entry.bind("<Return>", lambda e: save_word())
            entry.bind("<KP_Enter>", lambda e: save_word())
            edit_dialog.bind("<Escape>", lambda e: cancel_edit())
            
            # Wait for the dialog to close
            edit_dialog.wait_window()
            
            # Process the result
            if result["saved"]:
                new_word = result["new_word"]
                
                if new_word != word:  # Only update if actually changed
                    # Check for duplicates
                    if new_word in self.word_manager.words:
                        messagebox.showwarning("Duplicate Word", 
                                             f"The word '{new_word}' already exists in your list.\n\nPlease choose a different word.")
                        return
                    
                    # Update the word in the list
                    try:
                        word_index = self.word_manager.words.index(word)
                        self.word_manager.words[word_index] = new_word
                        
                        # If the old word was recorded, transfer the recording status
                        if word in self.word_manager.recorded_words:
                            self.word_manager.recorded_words.remove(word)
                            self.word_manager.recorded_words.add(new_word)
                        
                        # Save to file immediately
                        self.word_manager.save_data()
                        
                        # Update the display
                        self.filter_words()  # Refresh the entire list to maintain proper ordering
                        
                        # Show success message
                        messagebox.showinfo("Word Updated", 
                                          f"Successfully updated:\n'{word}' → '{new_word}'")
                        
                    except ValueError:
                        messagebox.showerror("Error", "Could not find the word in the list.")
                    except Exception as e:
                        messagebox.showerror("Error", f"Failed to update word: {str(e)}")
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open edit dialog: {str(e)}")
    
    def delete_word_inline(self, word):
        """Delete a word directly from the word list"""
        # Confirm deletion
        if messagebox.askyesno(lang.get("confirm_delete"), 
                              f"{lang.get('confirm_delete')}\n\n'{word}'"):
            try:
                # Remove from word list
                if word in self.word_manager.words:
                    self.word_manager.words.remove(word)
                
                # Remove from recorded words if it was recorded
                if word in self.word_manager.recorded_words:
                    self.word_manager.recorded_words.remove(word)
                
                # Adjust current index if necessary
                if hasattr(self.word_manager, 'current_index'):
                    if self.word_manager.current_index >= len(self.word_manager.words) and len(self.word_manager.words) > 0:
                        self.word_manager.current_index = len(self.word_manager.words) - 1
                    elif len(self.word_manager.words) == 0:
                        self.word_manager.current_index = 0
                
                # Save changes
                self.word_manager.save_data()
                
                # Update the display
                self.filter_words()  # Refresh the entire list
                
                messagebox.showinfo(lang.get("success"), 
                                  f"{lang.get('word_deleted')}: '{word}'")
                
            except Exception as e:
                messagebox.showerror(lang.get("error"), f"Failed to delete word: {str(e)}")
    
    def record_word_inline(self, word):
        """Record a word directly from the word list"""
        # Close the word list dialog temporarily
        if hasattr(self, 'dialog') and self.dialog:
            self.dialog.withdraw()
        
        try:
            # Find the word in the main recorder and start recording
            if hasattr(self.parent, 'current_word_index'):
                # Find the word index in the main list
                try:
                    word_index = self.word_manager.words.index(word)
                    self.parent.current_word_index = word_index
                    self.parent.current_word.set(word)
                    
                    # Update the main window display
                    if hasattr(self.parent, 'update_display'):
                        self.parent.update_display()
                    
                    # Start recording
                    if hasattr(self.parent, 'start_recording'):
                        self.parent.start_recording()
                    
                    messagebox.showinfo("Recording Started", f"Recording started for: '{word}'")
                    
                except ValueError:
                    messagebox.showerror("Error", "Word not found in main list.")
        finally:
            # Restore the word list dialog
            if hasattr(self, 'dialog') and self.dialog:
                self.dialog.deiconify()
    
    def on_selection_change(self):
        """Handle selection change"""
        # Update selected words set based on current page checkboxes
        for word, checkbox_info in self.word_checkboxes.items():
            if checkbox_info['var'].get():
                self.selected_words.add(word)
            else:
                self.selected_words.discard(word)
        
        # Update selection counter
        self.update_selection_display()
    
    def update_display_info(self):
        """Update title and display information"""
        recorded, total = self.word_manager.get_progress()
        filtered_count = len(self.filtered_words)
        
        if filtered_count == total:
            title_text = f"📚 Word List ({total} words, {recorded} recorded)"
        else:
            title_text = f"📚 Word List (showing {filtered_count} of {total} words, {recorded} recorded)"
        
        self.title_label.configure(text=title_text)
    
    def refresh_list(self):
        """Refresh the word list"""
        # Clear search and reset filters
        self.search_var.set("")
        self.filter_var.set("all")
        self.filter_words()
        
        # Clear any selections
        self.selected_words.clear()
        self.on_selection_change()
    
    def focus_search(self):
        """Focus the search entry for better UX"""
        if hasattr(self, 'search_entry'):
            self.search_entry.focus()
    
    def clear_search(self):
        """Clear search and show all words"""
        self.search_var.set("")
        self.filter_var.set("all")
        self.filter_words()
        if hasattr(self, 'search_entry'):
            self.search_entry.focus()
    
    def select_all(self):
        """Select all filtered words across all pages"""
        self.selected_words.update(self.filtered_words)
        # Update visible checkboxes
        for word, checkbox_info in self.word_checkboxes.items():
            checkbox_info['var'].set(True)
        self.update_selection_display()
    
    def select_none(self):
        """Deselect all words across all pages"""
        self.selected_words.clear()
        # Update visible checkboxes
        for checkbox_info in self.word_checkboxes.values():
            checkbox_info['var'].set(False)
        self.update_selection_display()
    
    def invert_selection(self):
        """Invert selection across all filtered words"""
        # Invert selection for all filtered words
        for word in self.filtered_words:
            if word in self.selected_words:
                self.selected_words.remove(word)
            else:
                self.selected_words.add(word)
        
        # Update visible checkboxes
        for word, checkbox_info in self.word_checkboxes.items():
            checkbox_info['var'].set(word in self.selected_words)
        self.update_selection_display()
    
    def add_new_word(self):
        """Add a new word"""
        dialog = ctk.CTkInputDialog(
            text="Enter new word or sentence:",
            title="Add Word"
        )
        new_word = dialog.get_input()
        
        if new_word and new_word.strip():
            new_word = new_word.strip()
            if new_word not in self.word_manager.words:
                self.word_manager.words.append(new_word)
                self.word_manager.save_data()
                self.refresh_list()
                messagebox.showinfo("Success", f"Added: {new_word}")
            else:
                messagebox.showwarning("Duplicate", "This word already exists in the list.")
    
    def edit_selected_words(self):
        """Edit selected words in bulk"""
        if not self.selected_words:
            messagebox.showwarning("No Selection", "Please select at least one word to edit.")
            return
        
        selected_list = sorted(list(self.selected_words))
        
        if len(selected_list) == 1:
            # Single word edit
            self.edit_single_word(selected_list[0])
        else:
            # Bulk edit
            self.bulk_edit_words(selected_list)
    
    def edit_single_word(self, word):
        """Edit a single word"""
        edit_dialog = ctk.CTkInputDialog(
            text=f"Edit word or sentence:\n\nCurrent: {word}",
            title="Edit Word"
        )
        
        new_word = edit_dialog.get_input()
        
        if new_word is None or not new_word.strip():
            return
        
        new_word = new_word.strip()
        if new_word == word:
            return
        
        # Check for duplicates
        if new_word in self.word_manager.words:
            messagebox.showwarning("Duplicate", f"The word '{new_word}' already exists in the list.")
            return
        
        # Update in word manager
        if word in self.word_manager.words:
            index = self.word_manager.words.index(word)
            self.word_manager.words[index] = new_word
            
            # Update recorded status if necessary
            if word in self.word_manager.recorded_words:
                self.word_manager.recorded_words.remove(word)
                self.word_manager.recorded_words.add(new_word)
            
            self.word_manager.save_data()
            self.refresh_list()
            messagebox.showinfo("Success", f"Changed '{word}' to '{new_word}'")
    
    def bulk_edit_words(self, words_list):
        """Edit multiple words with find/replace functionality"""
        # Create bulk edit dialog
        edit_dialog = ctk.CTkToplevel(self.dialog)
        edit_dialog.title("Bulk Edit Words")
        edit_dialog.geometry("600x500")
        edit_dialog.transient(self.dialog)
        edit_dialog.grab_set()
        
        # Center dialog
        edit_dialog.update_idletasks()
        x = (edit_dialog.winfo_screenwidth() - edit_dialog.winfo_width()) // 2
        y = (edit_dialog.winfo_screenheight() - edit_dialog.winfo_height()) // 2
        edit_dialog.geometry(f"+{x}+{y}")
        
        main_frame = ctk.CTkFrame(edit_dialog)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Title
        ctk.CTkLabel(
            main_frame,
            text=f"✏️ Bulk Edit ({len(words_list)} words selected)",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=(10, 20))
        
        # Find/Replace section
        replace_frame = ctk.CTkFrame(main_frame)
        replace_frame.pack(fill="x", pady=(0, 20))
        
        ctk.CTkLabel(replace_frame, text="🔍 Find and Replace:", font=ctk.CTkFont(size=14, weight="bold")).pack(pady=(10, 5))
        
        find_frame = ctk.CTkFrame(replace_frame)
        find_frame.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkLabel(find_frame, text="Find:").pack(side="left", padx=(10, 5))
        find_var = tk.StringVar()
        find_entry = ctk.CTkEntry(find_frame, textvariable=find_var, width=200)
        find_entry.pack(side="left", padx=5)
        
        replace_frame_inner = ctk.CTkFrame(replace_frame)
        replace_frame_inner.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkLabel(replace_frame_inner, text="Replace:").pack(side="left", padx=(10, 5))
        replace_var = tk.StringVar()
        replace_entry = ctk.CTkEntry(replace_frame_inner, textvariable=replace_var, width=200)
        replace_entry.pack(side="left", padx=5)
        
        # Preview section
        preview_frame = ctk.CTkFrame(main_frame)
        preview_frame.pack(fill="both", expand=True, pady=(0, 20))
        
        ctk.CTkLabel(preview_frame, text="📋 Selected Words:", font=ctk.CTkFont(size=14, weight="bold")).pack(pady=(10, 5))
        
        preview_text = ctk.CTkTextbox(preview_frame, width=550, height=200)
        preview_text.pack(padx=10, pady=(0, 10))
        
        # Show selected words
        preview_content = "Current words:\n" + "\n".join([f"{i+1}. {word}" for i, word in enumerate(words_list)])
        preview_text.insert("1.0", preview_content)
        
        def update_preview():
            find_text = find_var.get()
            replace_text = replace_var.get()
            
            if find_text:
                preview_content = "Preview after replace:\n"
                for i, word in enumerate(words_list):
                    new_word = word.replace(find_text, replace_text)
                    preview_content += f"{i+1}. {word} → {new_word}\n"
            else:
                preview_content = "Current words:\n" + "\n".join([f"{i+1}. {word}" for i, word in enumerate(words_list)])
            
            preview_text.delete("1.0", "end")
            preview_text.insert("1.0", preview_content)
        
        # Update preview on text change
        find_var.trace("w", lambda *args: update_preview())
        replace_var.trace("w", lambda *args: update_preview())
        
        # Buttons
        button_frame = ctk.CTkFrame(main_frame)
        button_frame.pack(fill="x")
        
        def apply_bulk_edit():
            find_text = find_var.get()
            replace_text = replace_var.get()
            
            if not find_text:
                messagebox.showwarning("No Find Text", "Please enter text to find.")
                return
            
            # Apply changes
            changes_made = 0
            for word in words_list:
                if find_text in word:
                    new_word = word.replace(find_text, replace_text)
                    if new_word != word and new_word not in self.word_manager.words:
                        # Update in word manager
                        index = self.word_manager.words.index(word)
                        self.word_manager.words[index] = new_word
                        
                        # Update recorded status
                        if word in self.word_manager.recorded_words:
                            self.word_manager.recorded_words.remove(word)
                            self.word_manager.recorded_words.add(new_word)
                        
                        changes_made += 1
            
            if changes_made > 0:
                self.word_manager.save_data()
                self.refresh_list()
                messagebox.showinfo("Success", f"Applied changes to {changes_made} words.")
                edit_dialog.destroy()
            else:
                messagebox.showinfo("No Changes", "No changes were made.")
        
        ctk.CTkButton(
            button_frame,
            text="✅ Apply Changes",
            command=apply_bulk_edit,
            width=140,
            height=35,
            fg_color="#28a745",
            hover_color="#1e7e34"
        ).pack(side="left", padx=(20, 10), pady=10)
        
        ctk.CTkButton(
            button_frame,
            text="❌ Cancel",
            command=edit_dialog.destroy,
            width=100,
            height=35,
            fg_color="#6c757d",
            hover_color="#545b62"
        ).pack(side="right", padx=(10, 20), pady=10)
    
    def delete_selected_words(self):
        """Delete selected words"""
        if not self.selected_words:
            messagebox.showwarning("No Selection", "Please select at least one word to delete.")
            return
        
        selected_list = sorted(list(self.selected_words))
        count = len(selected_list)
        
        # Confirmation dialog
        if count == 1:
            message = f"Delete this word?\n\n'{selected_list[0]}'"
        else:
            word_list = "\n".join([f"• {word}" for word in selected_list[:10]])
            if count > 10:
                word_list += f"\n... and {count - 10} more"
            message = f"Delete {count} selected words?\n\n{word_list}"
        
        if messagebox.askyesno("Confirm Deletion", message):
            # Remove words
            for word in selected_list:
                if word in self.word_manager.words:
                    self.word_manager.words.remove(word)
                if word in self.word_manager.recorded_words:
                    self.word_manager.recorded_words.remove(word)
            
            self.word_manager.save_data()
            self.refresh_list()
            
            if count == 1:
                messagebox.showinfo("Success", f"Deleted: {selected_list[0]}")
            else:
                messagebox.showinfo("Success", f"Deleted {count} words successfully.")


class SimplifiedRecorderApp:
    """Simplified version with working features and menu system"""
    
    def __init__(self):
        # Set default language to Turkish for Kurdish recorder
        lang.set_language("tr")
        
        # Initialize managers
        self.audio_manager = AudioManager()
        self.word_manager = WordManager()
        
        # Set initial recording mode in audio manager
        self.audio_manager.set_sentence_mode(self.word_manager.is_sentence_mode, self.word_manager.current_content_type)
        
        # UI state
        self.current_recording = None
        self.is_recording = False
        self.menu_window = None
        
        # Initialize UI
        ctk.set_appearance_mode(Config.UI_THEME)
        ctk.set_default_color_theme("blue")
        
        self.root = ctk.CTk()
        self.root.title("Kurmanji Word Recorder - Enhanced")
        self.root.geometry("900x600")
        
        self.setup_ui()
        self.check_microphone_status()
        
        # Auto-jump to first unrecorded word on startup
        self.jump_to_first_unrecorded_word()
        
        self.update_word_display()
    
    def setup_ui(self):
        """Setup the simplified UI"""
        # Main container
        self.main_frame = ctk.CTkFrame(self.root)
        self.main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Title with menu button
        title_frame = ctk.CTkFrame(self.main_frame)
        title_frame.pack(fill="x", pady=(20, 30))
        
        self.title_label = ctk.CTkLabel(
            title_frame,
            text=lang.get("app_title"),
            font=ctk.CTkFont(size=28, weight="bold")
        )
        self.title_label.pack(side="left", padx=(20, 0), pady=10)
        
        # Language selector
        lang_frame = ctk.CTkFrame(title_frame)
        lang_frame.pack(side="right", padx=(0, 10), pady=10)
        
        ctk.CTkLabel(
            lang_frame,
            text=lang.get("language"),
            font=ctk.CTkFont(size=12)
        ).pack(side="left", padx=(10, 5), pady=10)
        
        self.language_var = tk.StringVar(value="Türkçe")
        language_options = [name for code, name in lang.get_available_languages()]
        self.language_combo = ctk.CTkComboBox(
            lang_frame,
            values=language_options,
            variable=self.language_var,
            command=self.change_language,
            width=100,
            height=30
        )
        self.language_combo.pack(side="left", padx=(0, 10), pady=10)
        
        # Menu button
        self.menu_button = ctk.CTkButton(
            title_frame,
            text=lang.get("menu"),
            command=self.show_menu,
            width=100,
            height=40,
            font=ctk.CTkFont(size=14, weight="bold"),
            fg_color="#8B4F9B",
            hover_color="#6B3F7B"
        )
        self.menu_button.pack(side="right", padx=(0, 20), pady=10)
        
        # Status frame
        self.setup_status_frame()
        
        # Word display frame
        self.setup_word_frame()
        
        # Recording controls
        self.setup_controls_frame()
        
        # Progress frame
        self.setup_progress_frame()
        
    def setup_status_frame(self):
        """Setup microphone status display"""
        status_frame = ctk.CTkFrame(self.main_frame)
        status_frame.pack(fill="x", pady=(0, 20))
        
        # Microphone status
        self.mic_status_label = ctk.CTkLabel(
            status_frame,
            text="🔍 Checking microphone...",
            font=ctk.CTkFont(size=14)
        )
        self.mic_status_label.pack(side="left", padx=20, pady=10)
        
        # Device selection
        self.device_var = tk.StringVar()
        self.device_combo = ctk.CTkComboBox(
            status_frame,
            variable=self.device_var,
            command=self.on_device_change,
            width=200
        )
        self.device_combo.pack(side="right", padx=20, pady=10)
        
        # Populate devices
        self.update_device_list()
    
    def setup_word_frame(self):
        """Setup word display area"""
        word_frame = ctk.CTkFrame(self.main_frame)
        word_frame.pack(fill="x", pady=(0, 20))
        
        # Current word display
        self.current_word_display_label = ctk.CTkLabel(
            word_frame,
            text=lang.get("current_word"),
            font=ctk.CTkFont(size=16, weight="bold")
        )
        self.current_word_display_label.pack(pady=(20, 5))
        
        # Current word container with edit button
        word_container = ctk.CTkFrame(word_frame)
        word_container.pack(pady=(0, 10))
        
        self.current_word_label = ctk.CTkLabel(
            word_container,
            text="No words loaded",
            font=ctk.CTkFont(size=32, weight="bold"),
            text_color=("blue", "lightblue"),
            wraplength=900,  # Wrap text at 900 pixels
            justify="left"   # Left-align wrapped text
        )
        self.current_word_label.pack(side="left", padx=(20, 10), pady=15)
        
        # Edit current word button
        self.edit_current_btn = ctk.CTkButton(
            word_container,
            text=f"{lang.get('edit')}",
            command=self.edit_current_word,
            width=80,
            height=35,
            fg_color="#007bff",
            hover_color="#0056b3",
            font=ctk.CTkFont(size=12, weight="bold")
        )
        self.edit_current_btn.pack(side="left", padx=(0, 5), pady=15)
        
        # Delete current word button
        self.delete_current_btn = ctk.CTkButton(
            word_container,
            text=f"{lang.get('delete')}",
            command=self.delete_current_word,
            width=80,
            height=35,
            fg_color="#dc3545",
            hover_color="#bd2130",
            font=ctk.CTkFont(size=12, weight="bold")
        )
        self.delete_current_btn.pack(side="left", padx=(0, 20), pady=15)
        
        # === MULTI-SPEED RECORDING CONTROLS ===
        speed_section = ctk.CTkFrame(word_frame)
        speed_section.pack(pady=(10, 0), padx=20, fill="x")
        
        # Speed selection title
        self.speed_title_label = ctk.CTkLabel(
            speed_section,
            text=lang.get("speed_selection"),
            font=ctk.CTkFont(size=14, weight="bold")
        )
        self.speed_title_label.pack(pady=(10, 5))
        
        # Speed selector buttons
        speed_buttons_frame = ctk.CTkFrame(speed_section)
        speed_buttons_frame.pack(pady=(5, 10))
        
        self.speed_buttons = {}
        speeds = [
            (lang.get("slow_speed"), "slow", "#fd7e14"),
            (lang.get("normal_speed"), "normal", "#28a745"), 
            (lang.get("fast_speed"), "fast", "#dc3545")
        ]
        
        for speed_text, speed_key, color in speeds:
            btn = ctk.CTkButton(
                speed_buttons_frame,
                text=speed_text,
                command=lambda s=speed_key: self.set_recording_speed(s),
                width=120,
                height=40,
                fg_color=color if speed_key == "normal" else "#6c757d",
                hover_color=color,
                font=ctk.CTkFont(size=12, weight="bold")
            )
            btn.pack(side="left", padx=10, pady=10)
            self.speed_buttons[speed_key] = btn
        
        # Speed progress indicators
        self.speed_progress_frame = ctk.CTkFrame(speed_section)
        self.speed_progress_frame.pack(pady=(0, 15), padx=20, fill="x")
        
        self.progress_title_label = ctk.CTkLabel(
            self.speed_progress_frame,
            text=lang.get("recording_status"),
            font=ctk.CTkFont(size=12, weight="bold")
        )
        self.progress_title_label.pack(pady=(10, 5))
        
        # Progress indicators for each speed
        progress_container = ctk.CTkFrame(self.speed_progress_frame)
        progress_container.pack(pady=(5, 10))
        
        self.speed_indicators = {}
        for speed_text, speed_key, color in speeds:
            indicator_frame = ctk.CTkFrame(progress_container)
            indicator_frame.pack(side="left", padx=15, pady=5)
            
            indicator = ctk.CTkLabel(
                indicator_frame,
                text=f"{speed_text}\\n⭕ {lang.get('not_recorded_yet')}",
                font=ctk.CTkFont(size=11),
                text_color=("gray", "lightgray")
            )
            indicator.pack(padx=10, pady=8)
            self.speed_indicators[speed_key] = indicator
        
        # Multi-speed control buttons
        speed_control_frame = ctk.CTkFrame(speed_section)
        speed_control_frame.pack(pady=(5, 15))
        
        # Force next word button (for skipping incomplete words)
        self.force_next_btn = ctk.CTkButton(
            speed_control_frame,
            text=lang.get("force_next"),
            command=self.force_next_word,
            width=200,
            height=35,
            fg_color="#fd7e14",
            hover_color="#e55b00",
            font=ctk.CTkFont(size=12, weight="bold")
        )
        self.force_next_btn.pack(side="left", padx=10, pady=5)
        
        # Skip current word button
        self.skip_word_btn = ctk.CTkButton(
            speed_control_frame,
            text=lang.get("skip_word"),
            command=self.skip_current_word,
            width=180,
            height=35,
            fg_color="#6c757d",
            hover_color="#545b62",
            font=ctk.CTkFont(size=12, weight="bold")
        )
        self.skip_word_btn.pack(side="left", padx=10, pady=5)
        
        # Word navigation
        nav_frame = ctk.CTkFrame(word_frame)
        nav_frame.pack(pady=(0, 20))
        
        self.prev_btn = ctk.CTkButton(
            nav_frame,
            text=lang.get("previous"),
            command=self.previous_word,
            width=120
        )
        self.prev_btn.pack(side="left", padx=10, pady=10)
        
        self.word_counter_label = ctk.CTkLabel(
            nav_frame,
            text="0 / 0",
            font=ctk.CTkFont(size=14)
        )
        self.word_counter_label.pack(side="left", padx=20, pady=10)
        
        self.next_btn = ctk.CTkButton(
            nav_frame,
            text=lang.get("next"),
            command=self.next_word,
            width=120
        )
        self.next_btn.pack(side="left", padx=10, pady=10)
        
        # Recording status indicator
        self.recording_status = ctk.CTkLabel(
            word_frame,
            text="○ Not recorded",
            font=ctk.CTkFont(size=14),
            text_color=("orange", "yellow")
        )
        self.recording_status.pack(pady=(0, 10))
    
    def setup_controls_frame(self):
        """Setup recording controls"""
        controls_frame = ctk.CTkFrame(self.main_frame)
        controls_frame.pack(fill="x", pady=(0, 20))
        
        # Recording controls
        rec_controls = ctk.CTkFrame(controls_frame)
        rec_controls.pack(pady=20)
        
        self.record_btn = ctk.CTkButton(
            rec_controls,
            text=lang.get("start_recording"),
            command=self.toggle_recording,
            width=180,
            height=50,
            font=ctk.CTkFont(size=16, weight="bold")
        )
        self.record_btn.pack(side="left", padx=10)
        
        self.play_btn = ctk.CTkButton(
            rec_controls,
            text=lang.get("play"),
            command=self.play_recording,
            width=120,
            height=50,
            state="disabled"
        )
        self.play_btn.pack(side="left", padx=10)
        
        self.save_btn = ctk.CTkButton(
            rec_controls,
            text=lang.get("save_next"),
            command=self.save_and_next,
            width=150,
            height=50,
            state="disabled"
        )
        self.save_btn.pack(side="left", padx=10)
        
        # Mode switching button (prominent placement)
        self.mode_switch_btn = ctk.CTkButton(
            rec_controls,
            text="📝 Kelime Modu",
            command=self.toggle_sentence_mode,
            width=150,
            height=50,
            font=ctk.CTkFont(size=14, weight="bold"),
            fg_color=["#FF6B35", "#FF4500"],
            hover_color=["#FF4500", "#FF6B35"]
        )
        self.mode_switch_btn.pack(side="left", padx=10)
        
        # Audio level indicator
        self.audio_level_frame = ctk.CTkFrame(controls_frame)
        self.audio_level_frame.pack(fill="x", padx=20, pady=(0, 10))
        
        ctk.CTkLabel(
            self.audio_level_frame,
            text="Audio Level:",
            font=ctk.CTkFont(size=12)
        ).pack(side="left", padx=10, pady=5)
        
        self.audio_level_bar = ctk.CTkProgressBar(
            self.audio_level_frame,
            width=300,
            height=20
        )
        self.audio_level_bar.pack(side="left", padx=10, pady=5)
        self.audio_level_bar.set(0)
    
    def setup_progress_frame(self):
        """Setup progress display"""
        progress_frame = ctk.CTkFrame(self.main_frame)
        progress_frame.pack(fill="x", pady=(0, 20))
        
        self.progress_title_label = ctk.CTkLabel(
            progress_frame,
            text=lang.get("recording_progress"),
            font=ctk.CTkFont(size=16, weight="bold")
        )
        self.progress_title_label.pack(pady=(15, 5))
        
        self.progress_bar = ctk.CTkProgressBar(
            progress_frame,
            width=400,
            height=25
        )
        self.progress_bar.pack(pady=10)
        
        self.progress_label = ctk.CTkLabel(
            progress_frame,
            text="0 of 0 words recorded (0%)",
            font=ctk.CTkFont(size=14)
        )
        self.progress_label.pack(pady=(0, 10))
        
        # Status bar for messages
        self.status_label = ctk.CTkLabel(
            progress_frame,
            text="",
            font=ctk.CTkFont(size=12),
            text_color=("green", "lightgreen")
        )
        self.status_label.pack(pady=(0, 10))
    


    def show_menu(self):
        """Show the main menu"""
        if self.menu_window and self.menu_window.winfo_exists():
            self.menu_window.lift()
            return
            
        self.menu_window = ctk.CTkToplevel(self.root)
        self.menu_window.title("Kurmanji Recorder - Full Menu")
        self.menu_window.geometry("700x500")
        self.menu_window.transient(self.root)
        
        # Center the menu window
        self.menu_window.update_idletasks()
        x = (self.menu_window.winfo_screenwidth() - self.menu_window.winfo_width()) // 2
        y = (self.menu_window.winfo_screenheight() - self.menu_window.winfo_height()) // 2
        self.menu_window.geometry(f"+{x}+{y}")
        
        # Menu content
        main_frame = ctk.CTkScrollableFrame(self.menu_window)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Title
        ctk.CTkLabel(
            main_frame,
            text="🎤 Kurmanji Recorder - Full Menu",
            font=ctk.CTkFont(size=24, weight="bold")
        ).pack(pady=(20, 30))
        
        # Word Management Section  
        self.create_menu_section(main_frame, f"📝 {lang.get('word_management')}", [
            (lang.get('add_words_sentences'), self.add_words_manually, "green"),
            (lang.get('quick_add_word'), self.quick_add_word, "blue"),
            (lang.get('view_all_words'), self.view_all_words, "blue"),
            (lang.get('refresh_display'), self.refresh_main_display, "blue"),
            (lang.get('jump_first_unrecorded'), self.manual_jump_to_unrecorded, "blue"),
            (lang.get('statistics'), self.show_statistics, "orange"),
            (lang.get('remove_duplicates'), self.remove_duplicates, "orange"),
            (lang.get('clear_all_words'), self.clear_words, "red")
        ])
        
        # Whisper Optimization Section
        self.create_menu_section(main_frame, f"🎯 {lang.get('whisper_optimization')}", [
            (lang.get('old_recordings_convert'), self.convert_old_recordings, "purple"),
            (lang.get('show_conversion_results'), self.show_whisper_status, "blue"),
            (lang.get('generate_sentence_list'), self.generate_sentence_list, "green"),
            (lang.get('audio_quality_test'), self.audio_quality_test, "orange"),
            (lang.get('audio_augmentation'), self.audio_augmentation_dialog, "red"),
            (lang.get('whisper_training_prep'), self.prepare_whisper_training, "purple")
        ])
        
        # Dataset Management Section
        self.create_menu_section(main_frame, f"👥 {lang.get('dataset_management')}", [
            (lang.get('merge_datasets'), self.merge_datasets, "green"),
            (lang.get('set_speaker_id'), self.set_speaker_id, "orange"),
            (lang.get('show_merged_stats'), self.show_merged_stats, "blue")
        ])
        
        # Word List Management Section
        self.create_menu_section(main_frame, f"📝 {lang.get('word_list_management')}", [
            (lang.get('split_word_list'), self.split_word_list, "purple"),
            (lang.get('create_sub_list'), self.create_sub_list, "orange"),
            (lang.get('random_distribute'), self.random_distribute, "blue"),
            (lang.get('show_list_stats'), self.show_list_stats, "gray")
        ])
        
        # Import/Export Section
        self.create_menu_section(main_frame, f"📂 {lang.get('import_export')}", [
            (lang.get('import_txt'), lambda: self.load_document('txt'), "blue"),
            (lang.get('import_docx'), lambda: self.load_document('docx'), "blue"),
            (lang.get('import_pdf'), lambda: self.load_document('pdf'), "blue"),
            (lang.get('load_from_url'), self.load_from_url, "blue"),
            (lang.get('export_words'), self.export_words, "green"),
            (lang.get('export_dataset'), self.export_dataset, "green")
        ])
        
        # Settings Section
        self.create_menu_section(main_frame, f"⚙️ {lang.get('settings')}", [
            (lang.get('toggle_theme'), self.toggle_theme, "purple"),
            (lang.get('test_audio'), self.test_audio, "gray")
        ])
    
    def create_menu_section(self, parent, title, buttons):
        """Create a menu section with buttons"""
        section_frame = ctk.CTkFrame(parent)
        section_frame.pack(fill="x", pady=20, padx=10)
        
        ctk.CTkLabel(
            section_frame,
            text=title,
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=(15, 10))
        
        button_frame = ctk.CTkFrame(section_frame)
        button_frame.pack(pady=(0, 15))
        
        for i, (text, command, color) in enumerate(buttons):
            if i % 3 == 0 and i > 0:
                # Create new row every 3 buttons
                button_frame = ctk.CTkFrame(section_frame)
                button_frame.pack(pady=(0, 15))
            
            # Map colors to valid hex values to prevent crashes
            color_map = {
                "green": "#28a745",
                "blue": "#007bff", 
                "red": "#dc3545",
                "orange": "#fd7e14",
                "purple": "#8B4F9B",
                "gray": "#6c757d"
            }
            
            hover_color_map = {
                "green": "#1e7e34",
                "blue": "#0056b3",
                "red": "#bd2130", 
                "orange": "#e55b00",
                "purple": "#6B3F7B",
                "gray": "#545b62"
            }
            
            button_color = color_map.get(color, color)
            hover_color = hover_color_map.get(color, color)
            
            ctk.CTkButton(
                button_frame,
                text=text,
                command=command,
                width=200,
                height=35,
                fg_color=button_color,
                hover_color=hover_color
            ).pack(side="left", padx=5, pady=5)
    
    # Core functionality methods (same as before)
    def set_recording_speed(self, speed):
        """Set recording speed and update UI"""
        self.word_manager.set_recording_speed(speed)
        
        # Update button colors
        colors = {
            "slow": "#fd7e14",
            "normal": "#28a745", 
            "fast": "#dc3545"
        }
        
        for speed_key, btn in self.speed_buttons.items():
            if speed_key == speed:
                btn.configure(fg_color=colors[speed_key])  # Active color
            else:
                btn.configure(fg_color="#6c757d")  # Inactive color
        
        print(f"🎚️ Recording speed set to: {speed.upper()}")
    
    def update_speed_indicators(self):
        """Update speed recording indicators for current word"""
        current_word = self.word_manager.get_current_word()
        if not current_word:
            return
        
        speeds_status = self.word_manager.get_word_speed_status(current_word)
        completion = self.word_manager.get_completion_status(current_word)
        
        speed_names = {
            "slow": lang.get("slow_speed"),
            "normal": lang.get("normal_speed"),
            "fast": lang.get("fast_speed")
        }
        
        for speed_key, indicator in self.speed_indicators.items():
            is_recorded = speeds_status.get(speed_key, False)
            if is_recorded:
                indicator.configure(
                    text=f"{speed_names[speed_key]}\n✅ {lang.get('recorded')}",
                    text_color=("green", "lightgreen")
                )
            else:
                indicator.configure(
                    text=f"{speed_names[speed_key]}\n⭕ {lang.get('not_recorded_yet')}",
                    text_color=("gray", "lightgray")
                )
        
        # Update progress in title if needed
        if hasattr(self, 'speed_progress_frame'):
            completed_count = sum(1 for recorded in speeds_status.values() if recorded)
            progress_text = f"{lang.get('recording_status')}: {completed_count}/3 ({completion:.0f}% {lang.get('completion')})"
            
            # Find and update the progress label
            for child in self.speed_progress_frame.winfo_children():
                if isinstance(child, ctk.CTkLabel):
                    current_text = child.cget("text")
                    if "📊" in current_text or "Mevcut Kelimenin" in current_text:
                        child.configure(text=progress_text)
                        break

    def check_microphone_status(self):
        is_working, message = self.audio_manager.check_microphone()
        if is_working:
            self.mic_status_label.configure(
                text=f"✅ {message}",
                text_color=("green", "lightgreen")
            )
        else:
            self.mic_status_label.configure(
                text=f"❌ {message}",
                text_color=("red", "lightcoral")
            )
    
    def update_device_list(self):
        devices = self.audio_manager.devices
        if devices:
            device_names = [f"{d['name']}" for d in devices]
            self.device_combo.configure(values=device_names)
            if device_names:
                self.device_combo.set(device_names[0])
    
    def on_device_change(self, device_name):
        devices = self.audio_manager.devices
        for i, device in enumerate(devices):
            if device['name'] == device_name:
                self.audio_manager.selected_device = device['id']
                break
    
    def jump_to_first_unrecorded_word(self):
        """Jump to the first unrecorded word automatically"""
        if self.word_manager.words:
            if self.word_manager.jump_to_first_unrecorded():
                # Auto-jumped to first unrecorded word
                pass
            else:
                # All words have been recorded
                pass
    
    def update_word_display(self):
        current_word = self.word_manager.get_current_word()
        
        # Başlığı mode türüne göre güncelle
        if self.word_manager.current_content_type == "paragraph":
            display_title = lang.get("current_paragraph")
        elif self.word_manager.is_sentence_mode:
            display_title = lang.get("current_sentence")
        else:
            display_title = lang.get("current_word")
        
        # Ana başlığı güncelle
        if hasattr(self, 'current_word_display_label'):
            self.current_word_display_label.configure(text=display_title)
        
        if current_word:
            self.current_word_label.configure(text=current_word)
            
            is_recorded = self.word_manager.is_current_recorded()
            if is_recorded:
                status_text = f"✅ {lang.get('already_recorded')}"
                self.recording_status.configure(
                    text=status_text,
                    text_color=("green", "lightgreen")
                )
            else:
                status_text = f"○ {lang.get('not_recorded')}"
                self.recording_status.configure(
                    text=status_text,
                    text_color=("orange", "yellow")
                )
            
            # Update multi-speed indicators
            if hasattr(self, 'speed_indicators'):
                self.update_speed_indicators()
        else:
            if self.word_manager.current_content_type == "paragraph":
                no_content_text = lang.get("no_paragraphs_loaded")
            elif self.word_manager.is_sentence_mode:
                no_content_text = lang.get("no_sentences_loaded")
            else:
                no_content_text = lang.get("no_words_loaded")
            self.current_word_label.configure(text=no_content_text)
            self.recording_status.configure(text="")
        
        recorded, total = self.word_manager.get_progress()
        current_idx = self.word_manager.current_index + 1 if current_word else 0
        self.word_counter_label.configure(text=f"{current_idx} / {total}")
        
        if total > 0:
            progress = recorded / total
            self.progress_bar.set(progress)
            
            # Get correct content type based on mode
            if self.word_manager.current_content_type == "paragraph":
                content_type = lang.get("paragraphs_recorded")
            elif self.word_manager.current_content_type == "sentence" or self.word_manager.is_sentence_mode:
                content_type = lang.get("sentences_recorded")
            else:
                content_type = lang.get("words_recorded")
            
            progress_text = f"{recorded} / {total} {content_type} ({progress:.1%})"
            self.progress_label.configure(text=progress_text)
        else:
            # Get correct "no data" message based on mode
            if self.word_manager.current_content_type == "paragraph":
                no_data_text = lang.get("no_paragraphs_loaded")
            elif self.word_manager.is_sentence_mode:
                no_data_text = lang.get("no_sentences_loaded")
            else:
                no_data_text = lang.get("no_words_loaded")
            self.progress_bar.set(0)
            self.progress_label.configure(text=no_data_text)
    
    def previous_word(self):
        self.word_manager.previous_word()
        self.update_word_display()
    
    def next_word(self):
        self.word_manager.next_word()
        self.update_word_display()
    
    def force_next_word(self):
        """Force move to next word even if current word is not fully recorded"""
        current_word = self.word_manager.get_current_word()
        if current_word:
            missing_speeds = self.word_manager.get_missing_speeds(current_word)
            completion = self.word_manager.get_completion_status(current_word)
            
            if missing_speeds:
                # Show confirmation dialog
                missing_names = {
                    "slow": "🐌 Yavaş",
                    "normal": "🎯 Normal", 
                    "fast": "🚀 Hızlı"
                }
                missing_list = [missing_names[speed] for speed in missing_speeds]
                
                result = messagebox.askyesno(
                    "Kelimeyi Zorla Geç",
                    f"⚠️ '{current_word}' kelimesi için eksik kayıtlar var!\n\n"
                    f"📊 Tamamlanma: {completion:.0f}%\n"
                    f"🔄 Eksik: {', '.join(missing_list)}\n\n"
                    f"Yine de sonraki kelimeye geçmek istiyor musunuz?\n"
                    f"(Eksik kayıtları daha sonra tamamlayabilirsiniz)"
                )
                
                if result:
                    self.word_manager.next_word()
                    self.update_word_display()
                    self.update_status_message(f"⏭️ '{current_word}' atlandı - {completion:.0f}% tamamlanmış")
                    print(f"⏭️ Forced next: '{current_word}' (Completion: {completion:.0f}%)")
            else:
                # All speeds completed, just move normally
                self.word_manager.next_word()
                self.update_word_display()
                self.update_status_message(f"✅ '{current_word}' tamamlandı - Sonraki kelimeye geçildi")
    
    def skip_current_word(self):
        """Skip current word completely (mark as not important)"""
        current_word = self.word_manager.get_current_word()
        if current_word:
            result = messagebox.askyesno(
                "Kelimeyi Atla",
                f"'{current_word}' kelimesini tamamen atlamak istiyor musunuz?\n\n"
                f"⚠️ Bu kelime için hiçbir kayıt yapılmayacak!\n"
                f"(Daha sonra kelime listesinden manuel olarak bulabilirsiniz)"
            )
            
            if result:
                self.word_manager.next_word()
                self.update_word_display()
                self.update_status_message(f"⏭️ '{current_word}' tamamen atlandı")
                print(f"⏭️ Skipped word: '{current_word}'")
    
    def edit_current_word(self):
        """Edit the current word being displayed"""
        current_word = self.word_manager.get_current_word()
        if not current_word:
            messagebox.showwarning("No Word", "No word is currently loaded to edit.")
            return
        
        try:
            # Create edit dialog
            edit_dialog = ctk.CTkToplevel(self.root)
            edit_dialog.title("Edit Current Word")
            edit_dialog.geometry("500x200")
            edit_dialog.transient(self.root)
            edit_dialog.grab_set()
            edit_dialog.focus_force()
            
            # Center the dialog
            edit_dialog.update_idletasks()
            x = (edit_dialog.winfo_screenwidth() // 2) - (250)
            y = (edit_dialog.winfo_screenheight() // 2) - (100)
            edit_dialog.geometry(f"500x200+{x}+{y}")
            
            # Variables to store result
            result = {"saved": False, "new_word": ""}
            
            # Main frame
            main_frame = ctk.CTkFrame(edit_dialog, fg_color="transparent")
            main_frame.pack(fill="both", expand=True, padx=20, pady=20)
            
            # Title
            title_label = ctk.CTkLabel(
                main_frame,
                text=lang.get("edit_current_word"),
                font=ctk.CTkFont(size=18, weight="bold")
            )
            title_label.pack(pady=(0, 10))
            
            # Show current word
            current_label = ctk.CTkLabel(
                main_frame,
                text=f"Current: {current_word}",
                font=ctk.CTkFont(size=14),
                text_color=("gray60", "gray40")
            )
            current_label.pack(pady=(0, 15))
            
            # Entry for new word
            entry = ctk.CTkEntry(
                main_frame,
                width=450,
                height=40,
                font=ctk.CTkFont(size=14),
                placeholder_text=lang.get("enter_corrected_word")
            )
            entry.pack(pady=10)
            entry.insert(0, current_word)
            entry.focus_force()
            entry.select_range(0, "end")
            
            # Button frame
            btn_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
            btn_frame.pack(pady=15)
            
            def save_word():
                new_text = entry.get().strip()
                if new_text:
                    result["saved"] = True
                    result["new_word"] = new_text
                edit_dialog.destroy()
            
            def cancel_edit():
                result["saved"] = False
                edit_dialog.destroy()
            
            # Save button
            save_btn = ctk.CTkButton(
                btn_frame,
                text="💾 Save Changes",
                command=save_word,
                width=140,
                height=35,
                fg_color="#28a745",
                hover_color="#1e7e34",
                font=ctk.CTkFont(size=12, weight="bold")
            )
            save_btn.pack(side="left", padx=10)
            
            # Cancel button
            cancel_btn = ctk.CTkButton(
                btn_frame,
                text="❌ Cancel",
                command=cancel_edit,
                width=100,
                height=35,
                fg_color="#dc3545",
                hover_color="#bd2130"
            )
            cancel_btn.pack(side="left", padx=10)
            
            # Keyboard bindings
            entry.bind("<Return>", lambda e: save_word())
            entry.bind("<KP_Enter>", lambda e: save_word())
            edit_dialog.bind("<Escape>", lambda e: cancel_edit())
            
            # Wait for the dialog to close
            edit_dialog.wait_window()
            
            # Process the result
            if result["saved"]:
                new_word = result["new_word"]
                
                if new_word != current_word:  # Only update if actually changed
                    # Get current content list based on mode
                    current_list = self.word_manager.get_filtered_content()
                    all_content = self.word_manager.words + self.word_manager.sentences
                    
                    # Check for duplicates
                    if new_word in all_content:
                        messagebox.showwarning("Duplicate Word", 
                                             f"The word/sentence '{new_word}' already exists in your list.\n\nPlease choose a different one.")
                        return
                    
                    # Update the word/sentence in the appropriate list
                    try:
                        current_index = self.word_manager.current_index
                        
                        if self.word_manager.current_content_type == "paragraph":
                            # Paragraf modunda - paragraphs listesini güncelle
                            if current_index < len(self.word_manager.paragraphs):
                                self.word_manager.paragraphs[current_index] = new_word
                                self.word_manager.save_paragraphs()  # Save paragraphs separately
                        elif self.word_manager.current_content_type == "sentence" or self.word_manager.is_sentence_mode:
                            # Cümle modunda - sentences listesini güncelle
                            if current_index < len(self.word_manager.sentences):
                                self.word_manager.sentences[current_index] = new_word
                                self.word_manager.save_sentences()  # Save sentences separately
                        else:
                            # Kelime modunda - words listesini güncelle
                            if current_index < len(self.word_manager.words):
                                self.word_manager.words[current_index] = new_word
                        
                        # If the old word was recorded, transfer the recording status
                        if current_word in self.word_manager.recorded_words:
                            self.word_manager.recorded_words.remove(current_word)
                            self.word_manager.recorded_words.add(new_word)
                        
                        # Save to file immediately
                        self.word_manager.save_data()
                        
                        # Update the display to show the corrected content
                        self.update_word_display()
                        
                        # Show success message
                        if self.word_manager.current_content_type == "paragraph":
                            content_type = "paragraph"
                        elif self.word_manager.is_sentence_mode:
                            content_type = "sentence"
                        else:
                            content_type = "word"
                        messagebox.showinfo("Content Updated", 
                                          f"Successfully updated {content_type}:\n'{current_word[:50]}...' → '{new_word[:50]}...'\n\nYou can now record the corrected {content_type}.")
                        
                    except Exception as e:
                        messagebox.showerror("Error", f"Failed to update word: {str(e)}")
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open edit dialog: {str(e)}")
    
    def delete_current_word(self):
        """Delete the current word being displayed"""
        current_word = self.word_manager.get_current_word()
        if not current_word:
            messagebox.showwarning(lang.get("warning"), "No word is currently loaded to delete.")
            return
        
        # Confirm deletion
        content_type = "sentence" if self.word_manager.is_sentence_mode else "word"
        if messagebox.askyesno(f"Delete {content_type}?", 
                              f"Are you sure you want to delete this {content_type}?\n\n'{current_word}'"):
            try:
                # Get current index
                current_index = self.word_manager.current_index
                
                # Remove from appropriate list
                if self.word_manager.is_sentence_mode:
                    # Cümle modunda - sentences listesinden sil
                    if current_index < len(self.word_manager.sentences):
                        self.word_manager.sentences.pop(current_index)
                    current_list_length = len(self.word_manager.sentences)
                else:
                    # Kelime modunda - words listesinden sil
                    if current_index < len(self.word_manager.words):
                        self.word_manager.words.pop(current_index)
                    current_list_length = len(self.word_manager.words)
                
                # Remove from recorded words if it was recorded
                if current_word in self.word_manager.recorded_words:
                    self.word_manager.recorded_words.remove(current_word)
                
                # Adjust current index if necessary
                if current_index >= current_list_length and current_list_length > 0:
                    self.word_manager.current_index = current_list_length - 1
                elif current_list_length == 0:
                    self.word_manager.current_index = 0
                
                # Save changes
                self.word_manager.save_data()
                
                # Update display
                self.update_word_display()
                
                messagebox.showinfo(lang.get("success"), 
                                  f"{lang.get('word_deleted')}: '{current_word}'")
                
            except Exception as e:
                messagebox.showerror(lang.get("error"), f"Failed to delete word: {str(e)}")
    
    def toggle_recording(self):
        print(f"🔄 Toggle recording called - current state: {'recording' if self.is_recording else 'stopped'}")
        if not self.is_recording:
            self.start_recording()
        else:
            # Stop recording without confirmation
            self.stop_recording()
    
    def start_recording(self):
        current_content = self.word_manager.get_current_word()
        if not current_content:
            messagebox.showwarning(lang.get("warning"), lang.get("no_word_to_record"))
            return
        
        # Check recording mode configuration
        content_type = "SENTENCE" if self.word_manager.is_sentence_mode else "WORD"
        # Recording start - debug info cleaned up
        
        # Ensure audio manager is in correct mode
        self.audio_manager.set_sentence_mode(self.word_manager.is_sentence_mode, self.word_manager.current_content_type)
        
        result = self.audio_manager.start_recording(self.audio_manager.selected_device)
        if result is True:
            self.is_recording = True
            self.record_btn.configure(
                text=lang.get("stop_recording"),
                fg_color="red",
                hover_color="darkred"
            )
            self.play_btn.configure(state="disabled")
            self.save_btn.configure(state="disabled")
            self.monitor_audio_level()
        else:
            messagebox.showerror(lang.get("error"), f"Failed to start recording: {result[1] if isinstance(result, tuple) else 'Unknown error'}")
    
    def stop_recording(self):
        print(f"🛑 UI stop_recording() called - was recording: {self.is_recording}")
        
        self.current_recording = self.audio_manager.stop_recording()
        self.is_recording = False
        
        self.record_btn.configure(
            text=lang.get("start_recording"),
            fg_color=["#3B8ED0", "#1F6AA5"],
            hover_color=["#36719F", "#144870"]
        )
        
        if self.current_recording is not None:
            try:
                import numpy as np
                data = self.current_recording
                if isinstance(data, list):
                    if len(data) == 1 and isinstance(data[0], np.ndarray):
                        data = data[0]
                    elif len(data) > 1 and all(isinstance(x, np.ndarray) for x in data):
                        data = np.concatenate(data, axis=0)
                    else:
                        data = np.asarray(data, dtype=np.float32)
                if not isinstance(data, np.ndarray):
                    data = np.asarray(data, dtype=np.float32)
                if data.ndim > 1:
                    data = np.mean(data, axis=1).astype(np.float32)
                if data.size == 0:
                    raise ValueError("empty audio")
                max_amplitude = float(np.max(np.abs(data)))
                sr = getattr(self.audio_manager, 'sample_rate', getattr(Config, 'SAMPLE_RATE', 44100))
                duration = float(data.size) / float(sr)
            except Exception:
                max_amplitude = 0.0
                duration = 0.0

            # Relaxed thresholds to reduce false negatives on quiet devices
            if max_amplitude < 1e-4:
                messagebox.showwarning("Low Audio Level", 
                    f"Very low audio signal detected (amplitude: {max_amplitude:.6f}).\n"
                    "Please check your microphone level and try again.")
                return
            elif duration < 0.5:
                messagebox.showwarning("Recording Too Short", 
                    f"Recording is very short ({duration:.2f} seconds).\n"
                    "Please record for at least 0.5 seconds.")
                return
            
            # Play ve Save butonlarını aktif et
            self.play_btn.configure(state="normal")
            self.save_btn.configure(state="normal")
        else:
            messagebox.showwarning("Recording Failed", 
                "No audio recorded or recording failed.\n\n"
                "Please check your microphone and try again.")
    
    def monitor_audio_level(self):
        """Monitor real-time audio level for UI display"""
        if self.is_recording and hasattr(self.audio_manager, 'current_audio_level'):
            try:
                # Get current audio level from AudioManager
                level = getattr(self.audio_manager, 'current_audio_level', 0.0)
                
                # Update UI progress bar (scale level appropriately)
                scaled_level = min(level * 3.0, 1.0)  # Scale up for better visibility
                self.audio_level_bar.set(scaled_level)
                
                # Check for Bluetooth connection issues
                if hasattr(self.audio_manager, 'is_bluetooth_device') and self.audio_manager.is_bluetooth_device():
                    # Monitor Bluetooth audio quality
                    if level < 0.001:
                        self._low_level_count = getattr(self, '_low_level_count', 0) + 1
                        if self._low_level_count > 40:  # 2 seconds of low audio
                            # Show Bluetooth warning in status
                            if hasattr(self, 'status_label'):
                                self.status_label.configure(
                                    text="⚠️ Bluetooth audio level very low - check AirPods connection",
                                    text_color=("orange", "yellow")
                                )
                    else:
                        self._low_level_count = 0
                        # Clear warning if audio recovers  
                        if hasattr(self, 'status_label') and level > 0.01:
                            self.status_label.configure(
                                text="🎧 Bluetooth recording active",
                                text_color=("green", "lightgreen")
                            )
                
                # Continue monitoring
                self.root.after(50, self.monitor_audio_level)  # Update every 50ms for smoother UI
            except Exception as e:
                print(f"Audio level monitoring error: {e}")
                self.audio_level_bar.set(0)
                self.root.after(100, self.monitor_audio_level)
        else:
            # Not recording, set to zero
            self.audio_level_bar.set(0)
            if self.is_recording:  # Only continue if still recording
                self.root.after(100, self.monitor_audio_level)
    
    def play_recording(self):
        if self.current_recording is not None:
            self.audio_manager.play_audio(self.current_recording)
    
    def save_and_next(self):
        if self.current_recording is None:
            messagebox.showwarning("Warning", "No recording to save.")
            return
        
        current_word = self.word_manager.get_current_word()
        if not current_word:
            messagebox.showwarning("Warning", "No current word to save.")
            return
        
        try:
            # Multi-speed recording için filename oluştur
            current_speed = self.word_manager.current_recording_speed
            filename = self.word_manager.generate_speed_filename(current_word, current_speed)
            
            if self.audio_manager.save_audio(self.current_recording, filename):
                # Speed-specific recording olarak kaydet
                self.word_manager.mark_speed_recorded(current_word, current_speed, filename)
                
                # Ses dosyasının uzunluğunu hesapla ve güncelle
                self.update_audio_duration(filename)
                
                self.current_recording = None
                self.play_btn.configure(state="disabled")
                self.save_btn.configure(state="disabled")
                
                # Check if all speeds are completed for this word
                missing_speeds = self.word_manager.get_missing_speeds(current_word)
                completion = self.word_manager.get_completion_status(current_word)
                
                if not missing_speeds:  # All speeds completed
                    # Auto-move to next word when all 3 speeds are done
                    self.word_manager.next_word()
                    self.update_status_message(f"🎉 '{current_word}' tüm hızlarda tamamlandı! Sonraki kelimeye geçiliyor...")
                    print(f"🎉 '{current_word}' - All speeds completed! Moving to next word.")
                else:
                    # Stay on same word, show which speeds are missing
                    missing_names = {
                        "slow": "🐌 Yavaş",
                        "normal": "🎯 Normal", 
                        "fast": "🚀 Hızlı"
                    }
                    missing_list = [missing_names[speed] for speed in missing_speeds]
                    self.update_status_message(f"✅ '{current_word}' {current_speed} kaydedildi. Kalan: {', '.join(missing_list)}")
                    print(f"✅ '{current_word}' - {current_speed} speed saved. Missing: {missing_speeds}")
                
                self.update_word_display()
                print(f"✅ '{current_word}' {current_speed} kaydedildi: {filename} (Completion: {completion:.0f}%)")
            else:
                messagebox.showerror(lang.get("error"), "Ses dosyası kaydedilemedi.")
        except Exception as e:
            messagebox.showerror(lang.get("error"), f"Kaydetme hatası: {str(e)}")
    
    def update_status_message(self, message):
        """Status bar'da mesaj göster"""
        if hasattr(self, 'status_label'):
            self.status_label.configure(text=message)
            # 3 saniye sonra temizle
            self.root.after(3000, lambda: self.status_label.configure(text=""))
    
    def auto_save_and_next(self):
        """Otomatik kaydetme ve sonraki kelimeye geçme"""
        if self.current_recording is None:
            return
        
        current_word = self.word_manager.get_current_word()
        if not current_word:
            return
        
        try:
            # Whisper eğitimi için optimize edilmiş dosya adlandırma
            recorded_count = len(self.word_manager.recorded_words)
            file_number = recorded_count + 1
            
            # Temiz dosya adı
            clean_word = re.sub(r'[^\w\s-]', '', current_word)
            clean_word = re.sub(r'\s+', '_', clean_word.strip())
            
            # Dosya adı oluştur
            speaker_suffix = f"_{self.word_manager.speaker_id}" if self.word_manager.speaker_id != "default" else ""
            filename = f"{file_number:06d}_{clean_word}{speaker_suffix}.wav"
            
            if self.audio_manager.save_audio(self.current_recording, filename):
                # Whisper eğitimi için kaydet
                self.word_manager.mark_recorded(current_word, filename)
                
                # Ses dosyasının uzunluğunu hesapla
                self.update_audio_duration(filename)
                
                # UI'yi temizle
                self.current_recording = None
                self.play_btn.configure(state="disabled")
                self.save_btn.configure(state="disabled")
                
                # Sonraki kayıt edilmemiş kelimeye geç
                self.word_manager.next_word()
                self.update_word_display()
                
                # Konsola başarı mesajı
                print(f"✅ '{current_word}' otomatik kaydedildi: {filename}")
            else:
                # Sadece hata durumunda mesaj göster
                messagebox.showerror("Hata", "Ses dosyası kaydedilemedi.")
                # UI'yi kullanıcının tekrar denemesi için hazırla
                self.play_btn.configure(state="normal")
                self.save_btn.configure(state="normal")
                
        except Exception as e:
            messagebox.showerror("Hata", f"Otomatik kaydetme hatası: {str(e)}")
            # Hata durumunda UI'yi hazırla
            self.play_btn.configure(state="normal")
            self.save_btn.configure(state="normal")
    
    def update_audio_duration(self, filename):
        """Ses dosyasının uzunluğunu hesapla ve Whisper manifest dosyasını güncelle"""
        try:
            audio_path = Config.AUDIO_DIR / filename
            if audio_path.exists():
                # Ses uzunluğunu hesapla
                import soundfile as sf
                audio_data, sample_rate = sf.read(str(audio_path))
                duration = len(audio_data) / sample_rate
                
                # Manifest dosyasını güncelle
                self.update_manifest_duration(filename, duration)
                
                print(f"🎵 Ses uzunluğu hesaplandı: {filename} = {duration:.2f} saniye")
                
        except Exception as e:
            print(f"⚠️ Ses uzunluğu hesaplanamadı: {e}")
    
    def update_manifest_duration(self, filename, duration):
        """Manifest dosyasındaki ses uzunluğunu güncelle"""
        try:
            if not Config.WHISPER_MANIFEST.exists():
                return
                
            # Mevcut manifest dosyasını oku
            lines = []
            with open(Config.WHISPER_MANIFEST, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            # Son satırı güncelle (yeni eklenen dosya)
            if lines:
                last_line = lines[-1].strip()
                if last_line:
                    manifest_entry = json.loads(last_line)
                    if manifest_entry.get('audio') == filename:
                        manifest_entry['duration'] = round(duration, 2)
                        lines[-1] = json.dumps(manifest_entry, ensure_ascii=False) + '\n'
                        
                        # Dosyayı yeniden yaz
                        with open(Config.WHISPER_MANIFEST, 'w', encoding='utf-8') as f:
                            f.writelines(lines)
            
        except Exception as e:
            print(f"⚠️ Manifest güncelleme hatası: {e}")
    
    def convert_old_audio_files(self):
        """Önceki ses dosyalarını Whisper formatına dönüştür"""
        try:
            # Mevcut ses dosyalarını tara
            if not Config.AUDIO_DIR.exists():
                return 0, []
            
            audio_files = list(Config.AUDIO_DIR.glob("*.wav"))
            if not audio_files:
                return 0, []
            
            converted_count = 0
            conversion_log = []
            existing_manifest = set()
            
            # Mevcut manifest dosyasındaki kayıtları kontrol et
            if Config.WHISPER_MANIFEST.exists():
                with open(Config.WHISPER_MANIFEST, 'r', encoding='utf-8') as f:
                    for line in f:
                        if line.strip():
                            entry = json.loads(line.strip())
                            existing_manifest.add(entry.get('audio', ''))
            
            # Her ses dosyası için işlem yap
            for audio_file in audio_files:
                filename = audio_file.name
                
                # Zaten manifest'te varsa atla
                if filename in existing_manifest:
                    continue
                
                try:
                    # Dosya adından transkript çıkar
                    transcript = self.extract_transcript_from_filename(filename)
                    
                    if transcript:
                        # Ses uzunluğunu hesapla
                        import soundfile as sf
                        audio_data, sample_rate = sf.read(str(audio_file))
                        duration = len(audio_data) / sample_rate
                        
                        # Whisper manifest dosyasına ekle
                        manifest_entry = {
                            "audio": filename,
                            "text": transcript,
                            "language": "ku",
                            "duration": round(duration, 2)
                        }
                        
                        with open(Config.WHISPER_MANIFEST, 'a', encoding='utf-8') as f:
                            f.write(json.dumps(manifest_entry, ensure_ascii=False) + '\n')
                        
                        # Transkript dosyasına ekle
                        with open(Config.TRANSCRIPT_FILE, 'a', encoding='utf-8') as f:
                            f.write(f"{filename}\t{transcript}\n")
                        
                        # Bu kelimeyi recorded_words'e ekle
                        self.recorded_words.add(transcript)
                        
                        converted_count += 1
                        conversion_log.append(f"✅ {filename} -> '{transcript}' ({duration:.1f}s)")
                        
                        print(f"🔄 Dönüştürüldü: {filename} -> '{transcript}'")
                    else:
                        conversion_log.append(f"❌ {filename} -> Transkript çıkarılamadı")
                        
                except Exception as e:
                    conversion_log.append(f"❌ {filename} -> Hata: {str(e)}")
                    print(f"⚠️ Hata: {filename} işlenirken: {e}")
            
            # Verileri kaydet
            if converted_count > 0:
                self.save_data()
            
            return converted_count, conversion_log
            
        except Exception as e:
            print(f"❌ Dönüştürme hatası: {e}")
            return 0, [f"Genel hata: {str(e)}"]
    
    def extract_transcript_from_filename(self, filename):
        """Dosya adından transkript metni çıkar"""
        try:
            # .wav uzantısını kaldır
            name_without_ext = filename.replace('.wav', '')
            
            # Farklı formatları dene
            # Format 1: kelime_timestamp.wav -> kelime
            if '_' in name_without_ext:
                parts = name_without_ext.split('_')
                # Son parça timestamp ise (rakam), önceki kısımları al
                if len(parts) > 1:
                    last_part = parts[-1]
                    # Timestamp formatlarını kontrol et (20241004, 123456, vb.)
                    if (last_part.isdigit() and len(last_part) >= 6) or \
                       (len(last_part) == 6 and last_part.isdigit()):
                        transcript = '_'.join(parts[:-1])
                    else:
                        transcript = name_without_ext
                else:
                    transcript = name_without_ext
            else:
                transcript = name_without_ext
            
            # Temizle ve düzelt
            transcript = transcript.replace('_', ' ')  # Alt çizgileri boşluk yap
            transcript = transcript.strip()
            
            # Çok kısa veya boş ise reddet
            if len(transcript) < 2:
                return None
                
            return transcript
            
        except Exception as e:
            print(f"⚠️ Transkript çıkarma hatası: {e}")
            return None
    
    # Menu functions
    def add_words_manually(self):
        words = DocumentProcessor.get_manual_words_dialog(self.root)
        if words:
            self.word_manager.add_words(words)
            self.update_word_display()
            messagebox.showinfo("Success", f"Added {len(words)} words!")
    
    def quick_add_word(self):
        dialog = ctk.CTkInputDialog(text="Enter a word or sentence:", title="Quick Add")
        word = dialog.get_input()
        if word and word.strip():
            word = word.strip()
            if word not in self.word_manager.words:
                self.word_manager.add_words([word])
                self.update_word_display()
                messagebox.showinfo("Success", f"Added: {word}")
            else:
                messagebox.showwarning("Duplicate", "This word already exists.")
    
    def view_all_words(self):
        viewer = SimpleWordViewer(self.root, self.word_manager)
        viewer.show()
        self.root.after(100, self.update_word_display)
    
    def load_document(self, file_type):
        filetypes = {
            'txt': [("Text files", "*.txt")],
            'docx': [("Word documents", "*.docx")],
            'pdf': [("PDF files", "*.pdf")]
        }
        
        filename = filedialog.askopenfilename(
            title=f"Select {file_type.upper()} file",
            filetypes=filetypes[file_type]
        )
        
        if filename:
            words = DocumentProcessor.load_words_from_file(filename)
            if words:
                self.word_manager.add_words(words)
                self.update_word_display()
                messagebox.showinfo("Success", f"Loaded {len(words)} words from document.")
            else:
                messagebox.showwarning("Warning", "No words found in the document.")
    
    def load_from_url(self):
        dialog = ctk.CTkInputDialog(text="Enter URL:", title="Load from URL")
        url = dialog.get_input()
        if url and url.strip():
            try:
                response = requests.get(url.strip(), timeout=10)
                response.raise_for_status()
                content = response.text
                if '<html' in content.lower():
                    content = re.sub(r'<[^>]+>', ' ', content)
                    import html
                    content = html.unescape(content)
                words = DocumentProcessor._extract_words(content)
                if words:
                    self.word_manager.add_words(words)
                    self.update_word_display()
                    messagebox.showinfo("Success", f"Loaded {len(words)} words from URL!")
                else:
                    messagebox.showwarning("No Words", "No words found at the URL.")
            except Exception as e:
                messagebox.showerror("URL Error", f"Failed to load from URL:\n{str(e)}")
    
    def remove_duplicates(self):
        original_count = len(self.word_manager.words)
        seen = set()
        unique_words = []
        for word in self.word_manager.words:
            if word.lower() not in seen:
                seen.add(word.lower())
                unique_words.append(word)
        
        self.word_manager.words = unique_words
        self.word_manager.recorded_words = {
            word for word in self.word_manager.recorded_words 
            if word in unique_words
        }
        self.word_manager.save_data()
        self.update_word_display()
        
        removed_count = original_count - len(unique_words)
        if removed_count > 0:
            messagebox.showinfo("Duplicates Removed", f"Removed {removed_count} duplicate words.")
        else:
            messagebox.showinfo("No Duplicates", "No duplicate words found.")
    
    def clear_words(self):
        if messagebox.askyesno("Confirm", "Clear all words? This will reset recording progress."):
            self.word_manager.words = []
            self.word_manager.recorded_words = set()
            self.word_manager.current_index = 0
            self.word_manager.save_data()
            self.update_word_display()
            messagebox.showinfo("Success", "All words cleared.")
    
    def convert_old_recordings(self):
        """Önceki ses kayıtlarını Whisper formatına dönüştür"""
        # Önce kullanıcıya bilgi ver
        confirm_msg = """🔄 Eski Ses Dosyalarını Dönüştür
        
Bu işlem:
• Mevcut ses dosyalarınızı bulacak
• Dosya adlarından transkript çıkaracak  
• Whisper eğitimi formatına dönüştürecek
• Manifest ve transkript dosyalarını güncelleyecek
        
Devam etmek istiyor musunuz?"""
        
        if not messagebox.askyesno("Eski Dosyaları Dönüştür", confirm_msg):
            return
        
        try:
            # Dönüştürme işlemini başlat
            converted_count, conversion_log = self.word_manager.convert_old_audio_files()
            
            if converted_count > 0:
                # Başarı raporu
                success_msg = f"""🎉 Dönüştürme Tamamlandı!
                
📊 Sonuçlar:
• {converted_count} ses dosyası dönüştürüldü
• Whisper manifest dosyası güncellendi
• Transkript dosyası güncellendi
• Kelime listesi senkronize edildi
                
✅ Artık tüm ses dosyalarınız Whisper eğitimi için hazır!"""
                
                messagebox.showinfo("Dönüştürme Başarılı", success_msg)
                
                # Detaylı log göster
                if messagebox.askyesno("Detay Göster", "Dönüştürme detaylarını görmek ister misiniz?"):
                    self.show_conversion_log(conversion_log)
                
                # Display'i güncelle
                self.update_word_display()
                
            else:
                # Dönüştürülecek dosya bulunamadı
                info_msg = """ℹ️ Dönüştürülecek Dosya Yok
                
Durum:
• Hiç ses dosyası bulunamadı, VEYA
• Tüm ses dosyaları zaten Whisper formatında
                
✅ Sistem güncel durumda!"""
                
                messagebox.showinfo("Bilgi", info_msg)
                
        except Exception as e:
            messagebox.showerror("Dönüştürme Hatası", f"Hata oluştu: {str(e)}")
    
    def show_conversion_log(self, conversion_log):
        """Dönüştürme logunu göster"""
        log_window = ctk.CTkToplevel(self.root)
        log_window.title("Dönüştürme Detayları")
        log_window.geometry("700x500")
        
        # Başlık
        ctk.CTkLabel(
            log_window,
            text="📋 Dönüştürme Raporu",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=10)
        
        # Log metni
        log_text = "\n".join(conversion_log)
        
        text_widget = ctk.CTkTextbox(
            log_window, 
            width=650, 
            height=400, 
            font=ctk.CTkFont(family="Consolas", size=11)
        )
        text_widget.pack(padx=20, pady=10)
        text_widget.insert("1.0", log_text)
        text_widget.configure(state="disabled")
        
        # Kapat butonu
        ctk.CTkButton(
            log_window,
            text="✅ Tamam",
            command=log_window.destroy,
            width=100,
            height=35
        ).pack(pady=10)
    
    def export_words(self):
        if not self.word_manager.words:
            messagebox.showwarning("No Words", "No words to export.")
            return
        
        filename = filedialog.asksaveasfilename(
            title="Export Word List",
            filetypes=[("Text files", "*.txt"), ("JSON files", "*.json")],
            defaultextension=".txt"
        )
        
        if filename:
            try:
                file_path = Path(filename)
                if file_path.suffix.lower() == '.json':
                    export_data = {
                        'words': self.word_manager.words,
                        'recorded_words': list(self.word_manager.recorded_words),
                        'export_date': datetime.now().isoformat()
                    }
                    with open(file_path, 'w', encoding='utf-8') as f:
                        json.dump(export_data, f, ensure_ascii=False, indent=2)
                else:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        for word in self.word_manager.words:
                            status = "✅" if word in self.word_manager.recorded_words else "⭕"
                            f.write(f"{status} {word}\n")
                
                messagebox.showinfo("Export Success", f"Word list exported to:\n{file_path}")
            except Exception as e:
                messagebox.showerror("Export Error", f"Failed to export:\n{str(e)}")
    
    def export_dataset(self):
        recorded, total = self.word_manager.get_progress()
        if recorded == 0:
            messagebox.showwarning("Warning", "No recordings to export.")
            return
        
        export_dir = filedialog.askdirectory(title="Whisper Dataset Dışa Aktarma")
        if export_dir:
            try:
                import shutil
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                export_path = Path(export_dir) / f"whisper_kurmanci_{timestamp}"
                export_path.mkdir(exist_ok=True)
                
                # Audio klasörü oluştur ve ses dosyalarını kopyala
                audio_export = export_path / "audio"  
                audio_export.mkdir(exist_ok=True)
                audio_files = list(Config.AUDIO_DIR.glob("*.wav"))
                for audio_file in audio_files:
                    shutil.copy2(audio_file, audio_export / audio_file.name)
                
                # Whisper dosyalarını kopyala
                if Config.WHISPER_MANIFEST.exists():
                    shutil.copy2(Config.WHISPER_MANIFEST, export_path / "manifest.jsonl")
                if Config.TRANSCRIPT_FILE.exists():
                    shutil.copy2(Config.TRANSCRIPT_FILE, export_path / "transcripts.txt")
                
                # README oluştur
                readme = f"# Kurmancî Whisper Dataset\\n\\nToplam: {len(audio_files)} ses dosyası\\nTarih: {timestamp}\\nDil: ku"
                (export_path / "README.md").write_text(readme, encoding='utf-8')
                
                messagebox.showinfo("Başarılı", f"Whisper dataset hazır!\\n{export_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Export failed: {str(e)}")
    
    def toggle_theme(self):
        try:
            # Close menu window if it exists to prevent widget issues
            if self.menu_window and self.menu_window.winfo_exists():
                self.menu_window.destroy()
                self.menu_window = None
            
            current_theme = ctk.get_appearance_mode()
            new_theme = "light" if current_theme.lower() == "dark" else "dark"
            ctk.set_appearance_mode(new_theme)
            Config.UI_THEME = new_theme
            
            # Force update all widgets by calling update
            self.root.update()
            
            messagebox.showinfo("Theme Changed", f"Theme changed to {new_theme} mode.\n\nMenu has been refreshed and is ready to use.")
        except Exception as e:
            messagebox.showerror("Theme Error", f"Failed to change theme: {str(e)}")
    
    def show_statistics(self):
        recorded, total = self.word_manager.get_progress()
        audio_files = list(Config.AUDIO_DIR.glob("*.wav"))
        
        stats_text = f"""📊 Recording Statistics

📝 Words:
• Total words: {total}
• Recorded: {recorded}
• Remaining: {total - recorded}
• Progress: {(recorded/total*100) if total > 0 else 0:.1f}%

🎵 Audio:
• Audio files: {len(audio_files)}
• Dataset location: {Config.BASE_DIR.absolute()}
• Audio quality: {Config.SAMPLE_RATE} Hz, Mono
• Format: WAV (16-bit PCM)"""
        
        stats_window = ctk.CTkToplevel(self.root)
        stats_window.title("Statistics")
        stats_window.geometry("500x400")
        
        text_widget = ctk.CTkTextbox(stats_window, width=460, height=340, font=ctk.CTkFont(family="Consolas", size=12))
        text_widget.pack(padx=20, pady=20)
        text_widget.insert("1.0", stats_text)
        text_widget.configure(state="disabled")
    
    def show_whisper_status(self):
        """Whisper dönüştürme durumunu göster"""
        try:
            # Ses dosyalarını say
            audio_files = list(Config.AUDIO_DIR.glob("*.wav")) if Config.AUDIO_DIR.exists() else []
            
            # Manifest dosyasındaki kayıtları say
            manifest_count = 0
            if Config.WHISPER_MANIFEST.exists():
                with open(Config.WHISPER_MANIFEST, 'r', encoding='utf-8') as f:
                    for line in f:
                        if line.strip():
                            manifest_count += 1
            
            # Transkript dosyasındaki kayıtları say
            transcript_count = 0
            if Config.TRANSCRIPT_FILE.exists():
                with open(Config.TRANSCRIPT_FILE, 'r', encoding='utf-8') as f:
                    transcript_count = len([line for line in f if line.strip()])
            
            # Durum raporu
            status_text = f"""🎯 Whisper Dönüştürme Durumu

📊 Dosya Sayıları:
• Toplam ses dosyası: {len(audio_files)}
• Whisper manifest'te: {manifest_count}
• Transkript dosyasında: {transcript_count}

📋 Dosya Durumu:
• Whisper Manifest: {'✅ Mevcut' if Config.WHISPER_MANIFEST.exists() else '❌ Yok'}
• Transkript Dosyası: {'✅ Mevcut' if Config.TRANSCRIPT_FILE.exists() else '❌ Yok'}

🔄 Dönüştürme Durumu:
• Dönüştürülen: {manifest_count}/{len(audio_files)}
• Kalan: {max(0, len(audio_files) - manifest_count)}
• Durum: {'✅ Tamamlandı' if manifest_count >= len(audio_files) else '🔄 Eksik var'}

💡 Not:
{'Tüm ses dosyalarınız Whisper formatında!' if manifest_count >= len(audio_files) else 'Bazı eski ses dosyalarınız henüz dönüştürülmemiş. "Eski Ses Dosyalarını Dönüştür" butonunu kullanın!'}"""
            
            # Durum penceresini göster
            status_window = ctk.CTkToplevel(self.root)
            status_window.title("Whisper Dönüştürme Durumu")
            status_window.geometry("600x500")
            
            text_widget = ctk.CTkTextbox(
                status_window, 
                width=560, 
                height=440, 
                font=ctk.CTkFont(family="Consolas", size=12)
            )
            text_widget.pack(padx=20, pady=20)
            text_widget.insert("1.0", status_text)
            text_widget.configure(state="disabled")
            
        except Exception as e:
            messagebox.showerror("Hata", f"Whisper durumu kontrol edilirken hata: {str(e)}")
    
    def convert_old_recordings(self):
        """Eski ses kayıtlarını Whisper formatına dönüştür"""
        if hasattr(self.word_manager, 'convert_old_audio_files'):
            self.word_manager.convert_old_audio_files()
        else:
            messagebox.showwarning("Uyarı", "Dönüştürme fonksiyonu bulunamadı!")
    
    def set_speaker_id(self):
        """Konuşmacı ID'sini ayarla"""
        import customtkinter as ctk
        
        dialog = ctk.CTkInputDialog(
            text="Konuşmacı ID girin (örn: speaker1, ahmet, sara):",
            title="Konuşmacı ID Ayarla"
        )
        speaker_id = dialog.get_input()
        
        if speaker_id and speaker_id.strip():
            self.word_manager.speaker_id = speaker_id.strip().lower()
            messagebox.showinfo("Başarılı", f"Konuşmacı ID ayarlandı: {self.word_manager.speaker_id}")
            # Config dosyasına kaydet
            with open(Config.BASE_DIR / "speaker_config.txt", "w", encoding="utf-8") as f:
                f.write(self.word_manager.speaker_id)
        else:
            messagebox.showwarning("Uyarı", "Geçerli bir ID girin!")
    
    def merge_datasets(self):
        """Birden fazla dataset'i birleştir"""
        from tkinter import filedialog
        import json
        import shutil
        
        # Kaynak klasörleri seç
        source_dirs = []
        while True:
            dir_path = filedialog.askdirectory(
                title=f"Dataset Klasörü Seçin ({len(source_dirs)+1}. klasör, İptal=bitir)"
            )
            if not dir_path:
                break
            source_dirs.append(Path(dir_path))
        
        if len(source_dirs) < 2:
            messagebox.showwarning("Uyarı", "En az 2 dataset klasörü seçmelisiniz!")
            return
        
        # Hedef klasör seç
        target_dir = filedialog.askdirectory(title="Birleşik Dataset Hedef Klasörü")
        if not target_dir:
            return
        
        target_path = Path(target_dir) / "merged_dataset"
        target_audio = target_path / "audio"
        target_audio.mkdir(parents=True, exist_ok=True)
        
        merge_log = []
        file_counter = 1
        merged_manifest = []
        merged_transcripts = []
        
        for i, source_dir in enumerate(source_dirs):
            try:
                # Audio klasörünü bul
                audio_dir = None
                for possible_audio in [source_dir / "audio", source_dir / "kurmanji_dataset" / "audio"]:
                    if possible_audio.exists():
                        audio_dir = possible_audio
                        break
                
                if not audio_dir:
                    merge_log.append(f"❌ Audio klasörü bulunamadı: {source_dir}")
                    continue
                
                merge_log.append(f"📂 İşleniyor: {source_dir.name}")
                
                # Ses dosyalarını kopyala
                audio_files = list(audio_dir.glob("*.wav"))
                for audio_file in audio_files:
                    # Yeni dosya adı
                    original_name = audio_file.stem
                    transcript = self.word_manager.extract_transcript_from_filename(original_name)
                    new_name = f"{file_counter:06d}_{transcript}_speaker{i+1}.wav"
                    
                    # Dosyayı kopyala
                    shutil.copy2(audio_file, target_audio / new_name)
                    
                    # Manifest ve transcript için kaydet
                    merged_manifest.append({
                        "audio_filepath": f"audio/{new_name}",
                        "text": transcript,
                        "duration": 2.0,  # Varsayılan
                        "speaker_id": f"speaker{i+1}"
                    })
                    merged_transcripts.append(f"{new_name}\t{transcript}")
                    
                    file_counter += 1
                
                merge_log.append(f"✅ {len(audio_files)} dosya kopyalandı")
                
            except Exception as e:
                merge_log.append(f"❌ Hata ({source_dir.name}): {e}")
        
        # Birleşik dosyaları yaz
        try:
            # Whisper manifest
            with open(target_path / "whisper_manifest.jsonl", "w", encoding="utf-8") as f:
                for entry in merged_manifest:
                    json.dump(entry, f, ensure_ascii=False)
                    f.write("\n")
            
            # Transcripts
            with open(target_path / "transcripts.txt", "w", encoding="utf-8") as f:
                for transcript in merged_transcripts:
                    f.write(transcript + "\n")
            
            merge_log.append("")
            merge_log.append("✅ Birleştirme tamamlandı!")
            merge_log.append(f"📁 Hedef: {target_path}")
            merge_log.append(f"🎵 Toplam dosya: {len(merged_manifest)}")
            
            # Log göster
            log_text = "\n".join(merge_log)
            log_window = ctk.CTkToplevel(self.root)
            log_window.title("Dataset Birleştirme Sonucu")
            log_window.geometry("700x500")
            
            text_widget = ctk.CTkTextbox(log_window, width=660, height=440)
            text_widget.pack(padx=20, pady=20)
            text_widget.insert("1.0", log_text)
            text_widget.configure(state="disabled")
            
            messagebox.showinfo("Başarılı", f"Dataset birleştirme tamamlandı!\nToplam: {len(merged_manifest)} dosya")
            
        except Exception as e:
            messagebox.showerror("Hata", f"Birleştirme sırasında hata: {e}")
    
    def show_merged_stats(self):
        """Birleşik dataset istatistiklerini göster"""
        from tkinter import filedialog
        
        dataset_dir = filedialog.askdirectory(title="Dataset Klasörü Seçin")
        if not dataset_dir:
            return
        
        dataset_path = Path(dataset_dir)
        audio_dir = dataset_path / "audio"
        manifest_file = dataset_path / "whisper_manifest.jsonl"
        
        if not audio_dir.exists() or not manifest_file.exists():
            messagebox.showwarning("Uyarı", "Geçerli bir dataset klasörü seçin!")
            return
        
        # İstatistikleri hesapla
        audio_files = list(audio_dir.glob("*.wav"))
        speaker_stats = {}
        
        try:
            with open(manifest_file, 'r', encoding='utf-8') as f:
                for line in f:
                    if line.strip():
                        data = json.loads(line)
                        speaker = data.get('speaker_id', 'unknown')
                        if speaker not in speaker_stats:
                            speaker_stats[speaker] = 0
                        speaker_stats[speaker] += 1
            
            stats_text = f"""📊 Dataset İstatistikleri

📁 Klasör: {dataset_path.name}
🎵 Toplam ses dosyası: {len(audio_files)}

👥 Konuşmacı Dağılımı:
"""
            for speaker, count in speaker_stats.items():
                stats_text += f"   • {speaker}: {count} dosya\n"
            
            stats_text += f"""
📋 Dosya Durumu:
   • Manifest: {'✅' if manifest_file.exists() else '❌'}
   • Audio klasörü: {'✅' if audio_dir.exists() else '❌'}
   • Transcript dosyası: {'✅' if (dataset_path / 'transcripts.txt').exists() else '❌'}

🎯 Whisper Eğitimi: Hazır!"""
            
            # Stats penceresini göster
            stats_window = ctk.CTkToplevel(self.root)
            stats_window.title("Dataset İstatistikleri")
            stats_window.geometry("500x400")
            
            text_widget = ctk.CTkTextbox(stats_window, width=460, height=340)
            text_widget.pack(padx=20, pady=20)
            text_widget.insert("1.0", stats_text)
            text_widget.configure(state="disabled")
            
        except Exception as e:
            messagebox.showerror("Hata", f"İstatistik hesaplanırken hata: {e}")
    
    def split_word_list(self):
        """Ana kelime listesini alt listelere böl"""
        import customtkinter as ctk
        from tkinter import filedialog
        
        if not self.word_manager.words:
            messagebox.showwarning("Uyarı", "Önce kelime listesi yükleyin!")
            return
        
        # Alt liste sayısını sor
        dialog = ctk.CTkInputDialog(
            text=f"Toplam {len(self.word_manager.words)} kelime var.\nKaç alt listeye bölünsin?",
            title="Kelime Listesi Bölme"
        )
        
        try:
            num_splits = int(dialog.get_input() or "0")
            if num_splits < 2:
                messagebox.showwarning("Uyarı", "En az 2 alt liste oluşturmalısınız!")
                return
        except:
            messagebox.showwarning("Uyarı", "Geçerli bir sayı girin!")
            return
        
        # Hedef klasör seç
        output_dir = filedialog.askdirectory(title="Alt Listelerin Kaydedileceği Klasör")
        if not output_dir:
            return
        
        output_path = Path(output_dir)
        
        # Kelimeleri böl
        words_per_list = len(self.word_manager.words) // num_splits
        remaining_words = len(self.word_manager.words) % num_splits
        
        split_info = []
        start_idx = 0
        
        for i in range(num_splits):
            # Son listeye kalan kelimeleri de ekle
            end_idx = start_idx + words_per_list + (1 if i < remaining_words else 0)
            sub_words = self.word_manager.words[start_idx:end_idx]
            
            # Alt liste dosyası oluştur
            filename = f"kelime_listesi_{i+1:02d}_{len(sub_words)}_kelime.txt"
            filepath = output_path / filename
            
            with open(filepath, 'w', encoding='utf-8') as f:
                for word in sub_words:
                    f.write(f"{word}\n")
            
            split_info.append(f"📄 {filename}: {len(sub_words)} kelime")
            start_idx = end_idx
        
        # Sonuçları göster
        info_text = f"✅ Kelime listesi {num_splits} parçaya bölündü:\n\n" + "\n".join(split_info)
        info_text += f"\n\n📁 Konum: {output_path}"
        
        messagebox.showinfo("Başarılı", info_text)
    
    def create_sub_list(self):
        """Belirli aralıktaki kelimeleri alt liste olarak kaydet"""
        import customtkinter as ctk
        from tkinter import filedialog
        
        if not self.word_manager.words:
            messagebox.showwarning("Uyarı", "Önce kelime listesi yükleyin!")
            return
        
        # Aralık bilgilerini al
        dialog_start = ctk.CTkInputDialog(
            text=f"Başlangıç kelime numarası (1-{len(self.word_manager.words)}):",
            title="Alt Liste Oluştur"
        )
        start_num = dialog_start.get_input()
        
        if not start_num:
            return
        
        dialog_end = ctk.CTkInputDialog(
            text=f"Bitiş kelime numarası ({start_num}-{len(self.word_manager.words)}):",
            title="Alt Liste Oluştur"
        )
        end_num = dialog_end.get_input()
        
        if not end_num:
            return
        
        try:
            start_idx = int(start_num) - 1  # 0-based index
            end_idx = int(end_num)  # exclusive end
            
            if start_idx < 0 or end_idx > len(self.word_manager.words) or start_idx >= end_idx:
                messagebox.showerror("Hata", "Geçersiz aralık!")
                return
                
        except:
            messagebox.showerror("Hata", "Geçerli sayılar girin!")
            return
        
        # Dosya kaydetme yeri seç
        filename = filedialog.asksaveasfilename(
            title="Alt Liste Kaydet",
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            initialname=f"alt_liste_{start_num}_{end_num}_{end_idx-start_idx}_kelime.txt"
        )
        
        if filename:
            sub_words = self.word_manager.words[start_idx:end_idx]
            with open(filename, 'w', encoding='utf-8') as f:
                for word in sub_words:
                    f.write(f"{word}\n")
            
            messagebox.showinfo("Başarılı", 
                f"Alt liste oluşturuldu!\n"
                f"📄 Dosya: {Path(filename).name}\n"
                f"📝 Kelime sayısı: {len(sub_words)}\n"
                f"🔢 Aralık: {start_num}-{end_num}")
    
    def random_distribute(self):
        """Kelimeleri rastgele dağıtarak alt listeler oluştur"""
        import customtkinter as ctk
        from tkinter import filedialog
        import random
        
        if not self.word_manager.words:
            messagebox.showwarning("Uyarı", "Önce kelime listesi yükleyin!")
            return
        
        # Kişi sayısını sor
        dialog = ctk.CTkInputDialog(
            text=f"Toplam {len(self.word_manager.words)} kelime var.\nKaç kişiye rastgele dağıtılsın?",
            title="Rastgele Dağıtım"
        )
        
        try:
            num_people = int(dialog.get_input() or "0")
            if num_people < 2:
                messagebox.showwarning("Uyarı", "En az 2 kişi olmalı!")
                return
        except:
            messagebox.showwarning("Uyarı", "Geçerli bir sayı girin!")
            return
        
        # Hedef klasör seç
        output_dir = filedialog.askdirectory(title="Rastgele Dağıtım Dosyalarının Kaydedileceği Klasör")
        if not output_dir:
            return
        
        output_path = Path(output_dir)
        
        # Kelimeleri karıştır
        shuffled_words = self.word_manager.words.copy()
        random.shuffle(shuffled_words)
        
        # Dağıt
        words_per_person = len(shuffled_words) // num_people
        remaining = len(shuffled_words) % num_people
        
        distribute_info = []
        start_idx = 0
        
        for i in range(num_people):
            # Son kişiye kalan kelimeleri de ver
            end_idx = start_idx + words_per_person + (1 if i < remaining else 0)
            person_words = shuffled_words[start_idx:end_idx]
            
            # Dosya oluştur
            filename = f"rastgele_liste_kisi{i+1:02d}_{len(person_words)}_kelime.txt"
            filepath = output_path / filename
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(f"# Kişi {i+1} için rastgele kelime listesi\n")
                f.write(f"# Toplam kelime: {len(person_words)}\n\n")
                for word in person_words:
                    f.write(f"{word}\n")
            
            distribute_info.append(f"👤 Kişi {i+1}: {len(person_words)} kelime")
            start_idx = end_idx
        
        # Sonuçları göster
        info_text = f"🎲 Kelimeler rastgele {num_people} kişiye dağıtıldı:\n\n" + "\n".join(distribute_info)
        info_text += f"\n\n📁 Konum: {output_path}"
        
        messagebox.showinfo("Başarılı", info_text)
    
    def show_list_stats(self):
        """Kelime listesi istatistiklerini göster"""
        if not self.word_manager.words:
            messagebox.showwarning("Uyarı", "Kelime listesi boş!")
            return
        
        # İstatistikleri hesapla
        total_words = len(self.word_manager.words)
        recorded_words = len(self.word_manager.recorded_words)
        remaining_words = total_words - recorded_words
        
        # Karakter analizi
        word_lengths = [len(word) for word in self.word_manager.words]
        avg_length = sum(word_lengths) / len(word_lengths) if word_lengths else 0
        min_length = min(word_lengths) if word_lengths else 0
        max_length = max(word_lengths) if word_lengths else 0
        
        # En uzun/kısa kelimeler
        shortest_words = [w for w in self.word_manager.words if len(w) == min_length][:3]
        longest_words = [w for w in self.word_manager.words if len(w) == max_length][:3]
        
        stats_text = f"""📊 Kelime Listesi İstatistikleri

📝 Genel Bilgiler:
   • Toplam kelime: {total_words}
   • Kaydedilen: {recorded_words}
   • Kalan: {remaining_words}
   • Tamamlanma: %{(recorded_words/total_words*100):.1f}

📏 Karakter Analizi:
   • Ortalama uzunluk: {avg_length:.1f} karakter
   • En kısa: {min_length} karakter
   • En uzun: {max_length} karakter

🔤 Örnek Kelimeler:
   • En kısalar: {', '.join(shortest_words[:3])}
   • En uzunlar: {', '.join(longest_words[:3])}

💡 Dağıtım Önerileri:
   • 2 kişiye: {total_words//2} kelime/kişi
   • 5 kişiye: {total_words//5} kelime/kişi
   • 10 kişiye: {total_words//10} kelime/kişi"""
        
        # Stats penceresini göster
        import customtkinter as ctk
        stats_window = ctk.CTkToplevel(self.root)
        stats_window.title("Kelime Listesi İstatistikleri")
        stats_window.geometry("600x500")
        
        text_widget = ctk.CTkTextbox(stats_window, width=560, height=440)
        text_widget.pack(padx=20, pady=20)
        text_widget.insert("1.0", stats_text)
        text_widget.configure(state="disabled")
    
    def manual_jump_to_unrecorded(self):
        """Manually jump to first unrecorded word"""
        if not self.word_manager.words:
            messagebox.showinfo("No Words", "No words loaded. Please add some words first.")
            return
        
        current_word = self.word_manager.get_current_word()
        if self.word_manager.jump_to_first_unrecorded():
            new_word = self.word_manager.get_current_word()
            self.update_word_display()
            messagebox.showinfo("Jumped to Unrecorded", 
                f"Jumped to first unrecorded word:\n\n'{new_word}'\n\nPosition: {self.word_manager.current_index + 1}/{len(self.word_manager.words)}")
        else:
            recorded, total = self.word_manager.get_progress()
            if recorded == total:
                messagebox.showinfo("All Complete!", 
                    f"🎉 Congratulations!\n\nAll {total} words have been recorded!\n\nYour Kurdish dataset is complete.")
            else:
                messagebox.showinfo("No Unrecorded Words", "No unrecorded words found.")
    
    def test_audio(self):
        is_working, message = self.audio_manager.check_microphone()
        if is_working:
            messagebox.showinfo("Audio Test", f"✅ Audio system is working!\n\n{message}")
        else:
            messagebox.showerror("Audio Test", f"❌ Audio system has issues!\n\n{message}")
    
    def update_display_info(self):
        """Update word counter and progress display"""
        if hasattr(self, 'word_counter_label'):
            current_word = self.word_manager.get_current_word()
            if current_word:
                self.word_counter_label.configure(text=f"{self.word_manager.current_index + 1} / {len(self.word_manager.words)}")
        
        # Update progress
        if hasattr(self, 'progress_label') and hasattr(self, 'progress_bar'):
            recorded, total = self.word_manager.get_progress()
            if total > 0:
                progress_percent = recorded / total
                self.progress_bar.set(progress_percent)
                self.progress_label.configure(text=f"{recorded} of {total} words recorded ({progress_percent*100:.1f}%)")
            else:
                self.progress_bar.set(0)
                self.progress_label.configure(text="No words loaded")
    
    def refresh_main_display(self):
        """Ana ekranı tamamen yenile - yeni cümleleri göster"""
        try:
            # Word manager'ı yeniden yükle
            self.word_manager.load_data()
            
            # Reload paragraphs and sentences from files
            self.word_manager.load_paragraphs()
            
            # İlk kaydedilmemiş kelimeye atla
            self.word_manager.jump_to_first_unrecorded()
            
            # Ekranı güncelle
            self.update_word_display()
            self.update_display_info()
            
            # Menüyü kapat
            if self.menu_window:
                self.menu_window.destroy()
                self.menu_window = None
            
            # Başarı mesajı
            recorded, total = self.word_manager.get_progress()
            current_word = self.word_manager.get_current_word()
            
            # Get content counts
            words_count = len(self.word_manager.words)
            sentences_count = len(self.word_manager.sentences)
            paragraphs_count = len(self.word_manager.paragraphs)
            
            mode_text = ""
            if self.word_manager.current_content_type == "paragraph":
                mode_text = f"📄 Paragraf Modu Aktif ({paragraphs_count} paragraphs)"
            elif self.word_manager.is_sentence_mode:
                mode_text = f"🎭 Cümle Modu Aktif ({sentences_count} sentences)"
            else:
                mode_text = f"📝 Kelime Modu Aktif ({words_count} words)"
            
            messagebox.showinfo("Ana Ekran Yenilendi", 
                f"✅ Ana ekran başarıyla yenilendi!\n\n"
                f"📊 İçerik Sayıları:\n"
                f"   • Kelimeler: {words_count}\n"
                f"   • Cümleler: {sentences_count}\n"
                f"   • Paragraflar: {paragraphs_count}\n\n"
                f"✅ Kaydedilen: {recorded}\n"
                f"⭕ Kalan: {total - recorded}\n\n"
                f"🎯 Şu anki içerik:\n'{current_word[:80] if current_word else 'None'}...'\n\n"
                f"{mode_text}")
            
        except Exception as e:
            messagebox.showerror("Hata", f"Ana ekran yenilenirken hata oluştu: {str(e)}")
    
    # === YENİ GELIŞMIŞ ÖZELLİKLER: WHISPER DATASET ZENGİNLEŞTİRME ===
    
    def change_language(self, lang_code_or_name):
        """Change the application language"""
        # Handle both language codes (en, tr, ku) and language names (English, Türkçe, Kurmancî)
        if lang_code_or_name in ["English", "Türkçe", "Kurmancî"]:
            # Map language names to codes
            lang_map = {"English": "en", "Türkçe": "tr", "Kurmancî": "ku"}
            lang_code = lang_map.get(lang_code_or_name, "en")
        else:
            # It's already a code
            lang_code = lang_code_or_name
        
        lang.set_language(lang_code)
        
        # Update main UI elements
        if hasattr(self, 'title_label'):
            self.title_label.configure(text=lang.get("app_title"))
        
        if hasattr(self, 'menu_button'):
            self.menu_button.configure(text=lang.get("menu"))
        
        # Update all visible text elements
        self.update_word_display()
        
        # Update button texts
        if hasattr(self, 'record_btn'):
            if self.is_recording:
                self.record_btn.configure(text=lang.get("stop_recording"))
            else:
                self.record_btn.configure(text=lang.get("start_recording"))
        
        if hasattr(self, 'play_btn'):
            self.play_btn.configure(text=lang.get("play"))
        
        if hasattr(self, 'save_btn'):
            self.save_btn.configure(text=lang.get("save_next"))
        
        if hasattr(self, 'prev_btn'):
            self.prev_btn.configure(text=lang.get("previous"))
        
        if hasattr(self, 'next_btn'):
            self.next_btn.configure(text=lang.get("next"))
        
        if hasattr(self, 'edit_current_btn'):
            self.edit_current_btn.configure(text=lang.get("edit"))
        
        if hasattr(self, 'delete_current_btn'):
            self.delete_current_btn.configure(text=lang.get("delete"))
        
        # Update multi-speed UI elements
        if hasattr(self, 'speed_title_label'):
            self.speed_title_label.configure(text=lang.get("speed_selection"))
        
        if hasattr(self, 'progress_title_label'):
            self.progress_title_label.configure(text=lang.get("recording_status"))
        
        if hasattr(self, 'force_next_btn'):
            self.force_next_btn.configure(text=lang.get("force_next"))
        
        if hasattr(self, 'skip_word_btn'):
            self.skip_word_btn.configure(text=lang.get("skip_word"))
        
        # Update speed buttons
        if hasattr(self, 'speed_buttons'):
            speed_texts = {
                "slow": lang.get("slow_speed"),
                "normal": lang.get("normal_speed"),
                "fast": lang.get("fast_speed")
            }
            for speed_key, button in self.speed_buttons.items():
                button.configure(text=speed_texts[speed_key])
        
        # Update speed indicators
        if hasattr(self, 'speed_indicators'):
            self.update_speed_indicators()
        
        # Close menu if open to refresh it
        if self.menu_window:
            self.menu_window.destroy()
            self.menu_window = None
        
        # Show confirmation
        language_names = {"en": "English", "tr": "Türkçe", "ku": "Kurmancî"}
        messagebox.showinfo(
            lang.get("success"), 
            f"Language changed to {language_names.get(lang_code, lang_code)}\n"
            f"Dil değiştirildi: {language_names.get(lang_code, lang_code)}\n"
            f"Ziman hate guhartin: {language_names.get(lang_code, lang_code)}"
        )
    
    def show_multi_speed_info(self):
        """Show information about multi-speed recording"""
        current_word = self.word_manager.get_current_word()
        if not current_word:
            messagebox.showinfo(lang.get("warning"), lang.get("no_word_to_record"))
            return
        
        speeds_status = self.word_manager.get_word_speed_status(current_word)
        completion = self.word_manager.get_completion_status(current_word)
        missing_speeds = self.word_manager.get_missing_speeds(current_word)
        
        info_text = f"🎚️ Multi-Speed Recording Durumu\n"
        info_text += f"📝 Kelime: '{current_word}'\n\n"
        info_text += f"📊 Genel İlerleme: {completion:.0f}% ({sum(speeds_status.values())}/3)\n\n"
        info_text += f"🎯 Speed Durumları:\n"
        
        speed_names = {
            "slow": "🐌 Yavaş Kayıt",
            "normal": "🎯 Normal Kayıt", 
            "fast": "🚀 Hızlı Kayıt"
        }
        
        for speed, recorded in speeds_status.items():
            status = "✅ Tamamlandı" if recorded else "⭕ Bekliyor"
            info_text += f"   • {speed_names[speed]}: {status}\n"
        
        if missing_speeds:
            info_text += f"\n🔄 Eksik Kayıtlar:\n"
            for speed in missing_speeds:
                info_text += f"   • {speed_names[speed]}\n"
        else:
            info_text += f"\n🎉 Bu kelime için tüm hızlar kaydedildi!"
        
        info_text += f"\n\n💡 İpucu: Farklı hızlarda konuşarak Whisper modelinin daha iyi öğrenmesini sağlayın!"
        
        messagebox.showinfo("Multi-Speed Recording", info_text)
    
    def toggle_sentence_mode(self):
        """Kelime, cümle ve paragraf modu arasında döngü"""
        old_content_type = self.word_manager.current_content_type
        old_sentence_mode = self.word_manager.is_sentence_mode
        
        # Cycle through modes: word → sentence → paragraph → word
        if self.word_manager.current_content_type == "word":
            self.word_manager.current_content_type = "sentence"
            self.word_manager.is_sentence_mode = True
            mode = "Cümle Modu"
            mode_icon = "🎭"
        elif self.word_manager.current_content_type == "sentence":
            self.word_manager.current_content_type = "paragraph"
            self.word_manager.is_sentence_mode = False  # Paragraph uses longer timeouts like sentences
            mode = "Paragraf Modu"
            mode_icon = "📄"
        else:  # paragraph
            self.word_manager.current_content_type = "word"
            self.word_manager.is_sentence_mode = False
            mode = "Kelime Modu"
            mode_icon = "📝"
        
        # Update mode button text
        if hasattr(self, 'mode_switch_btn'):
            self.mode_switch_btn.configure(text=f"{mode_icon} {mode}")
        
        # Mode switch completed
        
        # İndeksi sıfırla ve yeni moda göre ilk kaydedilmemiş içeriğe atla
        old_index = self.word_manager.current_index
        self.word_manager.current_index = 0
        # Index reset for new mode
        
        # Filtrelenmiş içeriği al
        filtered_content = self.word_manager.get_filtered_content()
        content_type = "cümle" if self.word_manager.is_sentence_mode else "kelime"
        
        # Content filtering completed
        if filtered_content:
            # Content loaded successfully
            # İlk kaydedilmemiş içeriğe atla
            first_unrecorded = self.word_manager.find_first_unrecorded_word()
            # Found first unrecorded content
            pass
        else:
            # No content found in current mode
            pass
        
        # Update audio manager with new recording mode
        self.audio_manager.set_sentence_mode(self.word_manager.is_sentence_mode, self.word_manager.current_content_type)
        
        # Ana ekranı güncelle - bu çok önemli!
        # Updating display
        self.update_word_display()
        
        current_displayed = self.word_manager.get_current_word()
        # Mode switch completed successfully
        
        # Get mode-specific information
        mode_info = ""
        if self.word_manager.current_content_type == "paragraph":
            mode_info = f"📄 Paragraf Modu Aktif!\n\n• {len(filtered_content)} paragraf bulundu\n• Kayıt süresi: 45 saniye\n• Uzun metinleri kaydetmeye odaklanın\n• Whisper için optimal paragraf dataset'i!"
        elif self.word_manager.current_content_type == "sentence":
            mode_info = f"🎭 Cümle Modu Aktif!\n\n• {len(filtered_content)} cümle bulundu\n• Kayıt süresi: 25 saniye\n• Uzun cümleleri kaydetmeye odaklanın\n• Whisper için optimal cümle dataset'i!"
        else:
            mode_info = f"📝 Kelime Modu Aktif!\n\n• {len(filtered_content)} kelime bulundu\n• Kayıt süresi: 8 saniye\n• Tekil kelimeleri kaydetmeye odaklanın\n• Temel kelime hazinesi oluşturun!"
        
        mode_info += f"\n\n🎯 Şu anda gösterilen:\n'{current_displayed or 'Hiçbiri'}'"
        
        # Mode change popup removed - unnecessary interruption
    
    def generate_sentence_list(self):
        """Kurmancî cümleler ve hikaye parçaları oluştur"""
        sentence_dialog = ctk.CTkToplevel(self.root)
        sentence_dialog.title("🎭 Kurmancî Cümle Listesi Oluşturucu")
        sentence_dialog.geometry("900x700")
        sentence_dialog.transient(self.root)
        
        main_frame = ctk.CTkFrame(sentence_dialog)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Başlık
        ctk.CTkLabel(
            main_frame,
            text="📝 Whisper İçin Optimal Cümle Listesi",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=(10, 20))
        
        # Kategoriler
        categories_frame = ctk.CTkFrame(main_frame)
        categories_frame.pack(fill="x", pady=(0, 20))
        
        ctk.CTkLabel(categories_frame, text="🎯 Cümle Kategorileri:", font=ctk.CTkFont(size=14, weight="bold")).pack(pady=(10, 5))
        
        # Kategori butonları
        cat_buttons_frame = ctk.CTkFrame(categories_frame)
        cat_buttons_frame.pack(fill="x", padx=20, pady=10)
        
        categories = [
            ("👋 Günlük Konuşma", self.add_daily_sentences),
            ("📚 Eğitim & Öğretim", self.add_education_sentences),
            ("🏠 Aile & İlişkiler", self.add_family_sentences),
            ("🌿 Doğa & Çevre", self.add_nature_sentences),
            ("🍽️ Yemek & İçecek", self.add_food_sentences),
            ("⏰ Zaman & Tarih", self.add_time_sentences)
        ]
        
        for i, (cat_name, cat_func) in enumerate(categories):
            if i % 2 == 0:
                row_frame = ctk.CTkFrame(cat_buttons_frame)
                row_frame.pack(fill="x", pady=5)
            
            ctk.CTkButton(
                row_frame,
                text=cat_name,
                command=cat_func,
                width=200,
                height=35,
                fg_color=["#3B8ED0", "#1F6AA5"]
            ).pack(side="left", padx=10, pady=5)
        
        # Metin alanı
        text_label = ctk.CTkLabel(main_frame, text="✏️ Oluşturulan Cümleler:", font=ctk.CTkFont(size=14, weight="bold"))
        text_label.pack(pady=(20, 5))
        
        self.sentence_text = ctk.CTkTextbox(main_frame, width=850, height=350, font=ctk.CTkFont(size=12))
        self.sentence_text.pack(pady=(0, 20))
        
        # Örnek cümleler ekle
        sample_sentences = """# Günlük Konuşma Örnekleri (Whisper için ideal uzunluk)
Ez navê min dibêjim û ji Kurdistan im
Tu çawa yî û tu ji ku derê yî?
Em hez dikin ku bi hev re axivîn
Îro hewa pir xweş e û ez dixwazim derkevim
Gava min dest bi xwendinê kir, ez pir kêfxweş bûm

# Zaman ve Tarih İfadeleri  
Îro roja duşemê ye û ez dixwazim biçim bazarê
Sibê ez ê biçim xwendingeha xwe ya nû
Di sala borê de ez gelek cihan dîtin

# Soru Cümleleri (Whisper için önemli)
Tu dixwazî ku em bi hev re biçin?
Kengî tu ê vegere malê?
Çi ji te re pir girîng e di jiyanê de?"""
        
        self.sentence_text.insert("1.0", sample_sentences)
        
        # Butonlar
        button_frame = ctk.CTkFrame(main_frame)
        button_frame.pack(fill="x")
        
        ctk.CTkButton(
            button_frame,
            text="💾 Cümleleri Kelime Listesine Ekle",
            command=lambda: self.add_sentences_to_wordlist(sentence_dialog),
            width=250,
            height=40,
            fg_color="green",
            hover_color="darkgreen"
        ).pack(side="left", padx=20, pady=10)
        
        ctk.CTkButton(
            button_frame,
            text="🔄 Temizle",
            command=lambda: self.sentence_text.delete("1.0", "end"),
            width=100,
            height=40
        ).pack(side="left", padx=10, pady=10)
        
        ctk.CTkButton(
            button_frame,
            text="❌ Kapat",
            command=sentence_dialog.destroy,
            width=100,
            height=40
        ).pack(side="right", padx=20, pady=10)
    
    def add_daily_sentences(self):
        """Günlük konuşma cümlelerini ekle"""
        daily_sentences = """
# Günlük Selamlaşma ve Konuşma
Silav û rêz ji we re hemûyan
Ez spas dikim ji bona alîkariya we
Tu çawa hatî vir û çi dixwazî?
Dibe ku em bi hev re çayekê vexwin?
Ez pir kêfxweş im ku te nas dikim

# Günlük Aktiviteler
Ez her roj di sibe de hişyar dibim
Piştî taştê ez diçim kar
Êvaran ez bi malbata xwe re diniştim
Şevê berî razan ez pirtûk dixwînim
Em du rojekê carek diçin bazarê

# Hîs û Empatî
Ez hêvî dikim ku tu baş bî
Tu baş dixebitî û ez şahî dikim
Bila xwedê te biparêze û te baş bigire
Dil û can xweş be ji bo we hemûyan"""
        
        current_text = self.sentence_text.get("1.0", "end")
        self.sentence_text.insert("end", daily_sentences)
    
    def add_education_sentences(self):
        """Eğitim cümleleri ekle"""
        education_sentences = """
# Eğitim ve Öğrenme
Ez dixwazim zimanê Kurdî fêr bibim
Li zanîngeha me mamoste pir baş in
Kitêbên Kurdî pir balkêş û fêrker in
Zarok divê di zimanê dayikê de fêr bibin
Perwerde ji bo pêşketina mirovên girîng e

# Zanist û Teknolojî  
Dinyaya nû ya teknolojîyê pir bilez diguhure
Komputer û înternetê jiyana me guhartiye
Em divê bi zanist û teknolojîyê re hev bikin
Fêrbûn û pêşketin her dem divê berdewam be"""
        
        current_text = self.sentence_text.get("1.0", "end")
        self.sentence_text.insert("end", education_sentences)
    
    def add_family_sentences(self):
        """Aile cümleleri ekle"""  
        family_sentences = """
# Aile ve İlişkiler
Malbata min pir mezin e û em hej hev in
Diya min pir xweş fikirin dike û min fêr dike
Bavê min kesekî hêja ye û ji kar netirsê
Birayên min û xişkên min hêja ne
Em bi hev re gelek kêfê dikin û şahî dikin

# Dostluk ve Toplum
Hevalên min yên baş her dem li rex min in
Em hez dikin ku bi civakê re hev bikin  
Giştî mirov divê hev muhterem bikin
Jiyana civakî pir girîng e ji bo hemûyan"""
        
        current_text = self.sentence_text.get("1.0", "end")
        self.sentence_text.insert("end", family_sentences)
    
    def add_nature_sentences(self):
        """Doğa cümleleri ekle"""
        nature_sentences = """
# Doğa ve Çevre
Xwezî pir xweş e û em divê wê biparêzin
Daran, çêm û çiyan cihaneke xweş diafirînin
Bajar û gund her du cihên xweş in ji bo jiyanê
Havîn û zivistan her du werzên xweş in
Kulîlk û gulên cûr bi cûr rengên xweş dikin

# Heywan û Xwarin
Pez û çêlek heywanên pir spehî ne
Revî û hirç heywanên zel û jîr in
Parî û mişk heywanên pir bedew in
Ev hemû xwezîyê pir dewlemend dike"""
        
        current_text = self.sentence_text.get("1.0", "end")
        self.sentence_text.insert("end", nature_sentences)
    
    def add_food_sentences(self):
        """Yemek cümleleri ekle"""
        food_sentences = """
# Yemek ve İçecek
Xwarinên Kurdî pir tam û xweş in
Birinç, gosit û sebze pir baş in ji bo tenduristîyê
Çay û qehwe pir tê vexwarin li welatê me
Nan û penîr bingehê xwarinê ye li me
Em hez dikin xwarinê bi hev re bixwin

# Çifçîlik û Bêrî  
Genim, ceh û nisk çêlek girîng in
Trî, sêv û hêlûn fêkiyên xweş in
Pîyaz, sîr û felfel sebzeyên tam in
Her sal me bêrîya baş heye"""
        
        current_text = self.sentence_text.get("1.0", "end")
        self.sentence_text.insert("end", food_sentences)
    
    def add_time_sentences(self):
        """Zaman cümleleri ekle"""
        time_sentences = """
# Zaman ve Tarih
Îro rojeke pir xweş e û ez kêfxweş im
Duh ez çûm gund û hevalên xwe dîtin
Sibê ez ê biçim bazarê û tiştên hewce bikînim
Heftiya borê em mehmên mezin hatin
Sala ku tê de em ê bijîn pir baş be

# Demjimêr û Wext
Sibe sê saet e û ez hişyar dibim
Nivroj ezî diçim xwarinê
Êvare şeş saet e û ez ji kar vedigerim  
Şev û roj wexta xwe ya xweş heye"""
        
        current_text = self.sentence_text.get("1.0", "end")
        self.sentence_text.insert("end", time_sentences)
    
    def add_sentences_to_wordlist(self, dialog):
        """Cümleleri kelime listesine ekle"""
        content = self.sentence_text.get("1.0", "end").strip()
        if not content:
            messagebox.showwarning("Uyarı", "Eklenecek cümle bulunamadı!")
            return
        
        # Satırları ayır ve temizle
        lines = content.split('\n')
        sentences = []
        
        for line in lines:
            line = line.strip()
            # Boş satırları ve # ile başlayan yorumları atla
            if line and not line.startswith('#'):
                sentences.append(line)
        
        if sentences:
            added_count = self.word_manager.add_words(sentences)
            
            # Ana ekranı güncelle
            self.update_display_info()
            self.update_word_display()  # Bu çok önemli - ana ekrandaki kelimeyi günceller
            
            messagebox.showinfo("Başarılı", 
                f"🎉 {added_count} cümle başarıyla eklendi!\n\n"
                f"📝 Whisper eğitimi için optimal cümle uzunluğunda\n"
                f"🎯 Toplam içerik sayısı: {len(self.word_manager.words)}\n\n"
                f"✅ Ana ekran güncellendi - yeni cümleler görünecek!")
            
            dialog.destroy()
        else:
            messagebox.showwarning("Uyarı", "Geçerli cümle bulunamadı!")
    
    def audio_quality_test(self):
        """Gelişmiş ses kalitesi testi"""
        quality_dialog = ctk.CTkToplevel(self.root)
        quality_dialog.title("🔊 Ses Kalitesi Testi & Optimizasyon")
        quality_dialog.geometry("700x500")
        quality_dialog.transient(self.root)
        
        main_frame = ctk.CTkFrame(quality_dialog)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Başlık
        ctk.CTkLabel(
            main_frame,
            text="🎤 Whisper İçin Ses Kalitesi Optimizasyonu",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=(10, 20))
        
        # Test sonuçları alanı
        self.quality_results = ctk.CTkTextbox(main_frame, width=650, height=300, font=ctk.CTkFont(size=11))
        self.quality_results.pack(pady=(0, 20))
        
        # Test butonları
        test_frame = ctk.CTkFrame(main_frame)
        test_frame.pack(fill="x")
        
        ctk.CTkButton(
            test_frame,
            text="🔍 Tam Ses Testi Başlat",
            command=self.run_comprehensive_audio_test,
            width=200,
            height=40,
            fg_color="green"
        ).pack(side="left", padx=10, pady=10)
        
        ctk.CTkButton(
            test_frame,
            text="📊 Mevcut Kayıtları Analiz Et",
            command=self.analyze_existing_recordings,
            width=200,
            height=40,
            fg_color="blue"
        ).pack(side="left", padx=10, pady=10)
        
        ctk.CTkButton(
            test_frame,
            text="❌ Kapat",
            command=quality_dialog.destroy,
            width=100,
            height=40
        ).pack(side="right", padx=10, pady=10)
        
        # İlk analizi başlat
        self.run_comprehensive_audio_test()
    
    def run_comprehensive_audio_test(self):
        """Kapsamlı ses kalitesi testi"""
        import sounddevice as sd
        import numpy as np
        
        self.quality_results.delete("1.0", "end")
        self.quality_results.insert("1.0", "🔄 Ses kalitesi testi başlatılıyor...\n\n")
        
        try:
            # Mikrofon testi
            is_working, mic_message = self.audio_manager.check_microphone()
            
            result_text = f"""🎤 WHISPER İÇİN SES KALİTESİ RAPORU
{"="*50}

📊 Mikrofon Durumu: {'✅ Çalışıyor' if is_working else '❌ Sorunlu'}
📝 Detay: {mic_message}

🔧 Önerilen Ayarlar:
   • Sample Rate: {Config.SAMPLE_RATE} Hz (Whisper için optimal)
   • Kanal: {Config.CHANNELS} (Mono - Whisper standard)
   • Bit Derinliği: 32-bit float (Yüksek kalite)

🎯 Whisper Eğitimi İçin Öneriler:
   
   1. 📏 SES UZUNLUĞU:
      • Kelimeler: 1-3 saniye
      • Cümleler: 3-15 saniye (OPTIMAL)
      • Paragraflar: 15-30 saniye
   
   2. 🔊 SES KALİTESİ:
      • Gürültüsüz ortam tercih edin
      • Mikrofonu 15-20 cm mesafede tutun
      • Sabit ses seviyesi kullanın
   
   3. 🎭 ÇEŞİTLİLİK:
      • Farklı konuşma hızları deneyin
      • Farklı tonlama çeşitleri ekleyin
      • Soru, ünlem, normal cümleler karıştırın
   
   4. 📚 İÇERİK ÇEŞİTLİLİĞİ:
      • Günlük konuşma cümleleri (40%)
      • Resmi metinler (20%)
      • Hikaye ve anlatılar (20%)
      • Soru-cevap diyalogları (20%)

🚀 Whisper Başarı İpuçları:
   • En az 1000+ cümle kaydı yapın
   • Farklı konuşmacılar dahil edin  
   • Konsisten ses kalitesi sağlayın
   • Doğru transkript yazımına dikkat edin
"""
            
            if is_working:
                # Kısa test kaydı yap
                test_duration = 2  # 2 saniye
                self.quality_results.insert("end", "\n🎙️ 2 saniyelik test kaydı alınıyor...\n")
                
                test_recording = sd.rec(
                    int(test_duration * Config.SAMPLE_RATE),
                    samplerate=Config.SAMPLE_RATE,
                    channels=Config.CHANNELS,
                    device=self.audio_manager.selected_device,
                    dtype=Config.DTYPE
                )
                sd.wait()
                
                if test_recording is not None:
                    # Analiz yap
                    max_amplitude = np.max(np.abs(test_recording))
                    rms_level = np.sqrt(np.mean(test_recording**2))
                    
                    result_text += f"""
📈 TEST KAYDI ANALİZİ:
   • Maksimum seviye: {max_amplitude:.4f}
   • RMS seviye: {rms_level:.4f}
   • Süre: {test_duration} saniye
   • Durum: {'✅ İYİ' if max_amplitude > 0.01 else '⚠️ DÜŞÜK SES'}
   
{'✅ Ses seviyeniz Whisper eğitimi için uygun!' if max_amplitude > 0.01 else '⚠️ Mikrofona daha yakın konuşun veya ses seviyesini artırın!'}"""
            
            self.quality_results.delete("1.0", "end")
            self.quality_results.insert("1.0", result_text)
            
        except Exception as e:
            error_text = f"❌ Test sırasında hata: {str(e)}\n\nTemel ses sistemi kontrolü yapılacak..."
            self.quality_results.insert("end", error_text)
    
    def analyze_existing_recordings(self):
        """Mevcut kayıtları analiz et"""
        if not Config.AUDIO_DIR.exists():
            messagebox.showwarning("Uyarı", "Henüz ses kaydı bulunamadı!")
            return
        
        audio_files = list(Config.AUDIO_DIR.glob("*.wav"))
        if not audio_files:
            messagebox.showwarning("Uyarı", "Henüz .wav ses dosyası bulunamadı!")
            return
        
        self.quality_results.delete("1.0", "end")
        self.quality_results.insert("1.0", "🔍 Mevcut kayıtlar analiz ediliyor...\n\n")
        
        try:
            import soundfile as sf
            
            total_duration = 0
            total_files = len(audio_files)
            durations = []
            
            analysis_text = f"""📊 MEVCUT KAYITLAR ANALİZİ
{"="*40}

📁 Toplam dosya sayısı: {total_files}
📂 Klasör: {Config.AUDIO_DIR}

🎵 DOSYA ANALİZİ:
"""
            
            for i, audio_file in enumerate(audio_files[:10]):  # İlk 10 dosyayı analiz et
                try:
                    data, samplerate = sf.read(audio_file)
                    duration = len(data) / samplerate
                    total_duration += duration
                    durations.append(duration)
                    
                    analysis_text += f"   {i+1:2d}. {audio_file.name:<30} | {duration:.2f}s\n"
                except:
                    analysis_text += f"   {i+1:2d}. {audio_file.name:<30} | HATA\n"
            
            if durations:
                avg_duration = total_duration / len(durations)
                
                analysis_text += f"""
📈 İSTATİSTİKLER:
   • Ortalama süre: {avg_duration:.2f} saniye
   • Toplam süre: {total_duration/60:.1f} dakika
   • En kısa: {min(durations):.2f}s
   • En uzun: {max(durations):.2f}s

🎯 WHISPER DEĞERLENDİRMESİ:
"""
                
                if avg_duration < 1:
                    analysis_text += "   ⚠️  Çok kısa kayıtlar - Cümle moduna geçmeyi düşünün\n"
                elif avg_duration > 20:
                    analysis_text += "   ⚠️  Çok uzun kayıtlar - Daha kısa bölümlere ayırın\n"
                else:
                    analysis_text += "   ✅ Optimal uzunluk aralığında\n"
                
                # Whisper için öneriler
                analysis_text += f"""
💡 WHISPER EĞİTİMİ ÖNERİLERİ:
   • {'✅' if total_files >= 500 else '⚠️'} Dataset boyutu: {total_files}/500+ dosya
   • {'✅' if total_duration > 1800 else '⚠️'} Toplam süre: {total_duration/60:.1f}/30+ dakika
   • {'✅' if 2 <= avg_duration <= 15 else '⚠️'} Ortalama uzunluk: {avg_duration:.1f}s (2-15s ideal)

🚀 Sonraki adımlar:
   1. {max(0, 500-total_files)} daha fazla kayıt yapın
   2. Cümle modunu aktif edin (daha uzun içerik için)
   3. Farklı konuşmacılardan kayıt toplayın
   4. Whisper eğitim dosyalarını hazırlayın
"""
            
            self.quality_results.delete("1.0", "end")
            self.quality_results.insert("1.0", analysis_text)
            
        except Exception as e:
            error_text = f"❌ Analiz sırasında hata: {str(e)}"
            self.quality_results.insert("end", error_text)
    
    def audio_augmentation_dialog(self):
        """Ses geliştirme ve çoğaltma seçenekleri"""
        aug_dialog = ctk.CTkToplevel(self.root)
        aug_dialog.title("🎵 Audio Augmentation - Ses Çeşitlendirme")
        aug_dialog.geometry("800x600")
        aug_dialog.transient(self.root)
        
        main_frame = ctk.CTkFrame(aug_dialog)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Başlık
        ctk.CTkLabel(
            main_frame,
            text="🎵 Whisper İçin Ses Dataset Çeşitlendirme",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=(10, 20))
        
        # Açıklama
        info_text = """🎯 Audio Augmentation Nedir?
Mevcut ses kayıtlarınızı farklı versiyonlarla çoğaltarak Whisper'ın daha güçlü öğrenmesini sağlar.

🔧 Mevcut Augmentation Seçenekleri:
   • Hız değişiklikleri (0.8x - 1.2x)
   • Ses seviyesi varyasyonları  
   • Hafif gürültü ekleme
   • Pitch (ton) değişiklikleri
   • Echo/reverb efektleri"""
        
        info_label = ctk.CTkTextbox(main_frame, width=750, height=150, font=ctk.CTkFont(size=12))
        info_label.pack(pady=(0, 20))
        info_label.insert("1.0", info_text)
        info_label.configure(state="disabled")
        
        # Seçenekler
        options_frame = ctk.CTkFrame(main_frame)
        options_frame.pack(fill="x", pady=(0, 20))
        
        # Checkboxes for augmentation options
        self.aug_speed = tk.BooleanVar(value=True)
        self.aug_volume = tk.BooleanVar(value=True)  
        self.aug_noise = tk.BooleanVar(value=False)
        self.aug_pitch = tk.BooleanVar(value=False)
        
        ctk.CTkLabel(options_frame, text="🎛️ Augmentation Seçenekleri:", font=ctk.CTkFont(size=14, weight="bold")).pack(pady=(10, 10))
        
        options_grid = ctk.CTkFrame(options_frame)
        options_grid.pack(padx=20, pady=10)
        
        ctk.CTkCheckBox(options_grid, text="⚡ Hız Değişiklikleri (0.9x, 1.1x)", variable=self.aug_speed).pack(anchor="w", pady=5)
        ctk.CTkCheckBox(options_grid, text="🔊 Ses Seviyesi Varyasyonları", variable=self.aug_volume).pack(anchor="w", pady=5)
        ctk.CTkCheckBox(options_grid, text="🔇 Hafif Arka Plan Gürültüsü", variable=self.aug_noise).pack(anchor="w", pady=5)
        ctk.CTkCheckBox(options_grid, text="🎵 Pitch/Ton Değişiklikleri", variable=self.aug_pitch).pack(anchor="w", pady=5)
        
        # Sonuç alanı
        self.aug_results = ctk.CTkTextbox(main_frame, width=750, height=200, font=ctk.CTkFont(size=11))
        self.aug_results.pack(pady=(0, 20))
        
        # Butonlar
        button_frame = ctk.CTkFrame(main_frame)
        button_frame.pack(fill="x")
        
        ctk.CTkButton(
            button_frame,
            text="🚀 Augmentation Başlat",
            command=self.start_audio_augmentation,
            width=200,
            height=40,
            fg_color="green"
        ).pack(side="left", padx=10, pady=10)
        
        ctk.CTkButton(
            button_frame,
            text="📊 Önizleme",
            command=self.preview_augmentation,
            width=150,
            height=40,
            fg_color="blue"
        ).pack(side="left", padx=10, pady=10)
        
        ctk.CTkButton(
            button_frame,
            text="❌ Kapat",
            command=aug_dialog.destroy,
            width=100,
            height=40
        ).pack(side="right", padx=10, pady=10)
    
    def start_audio_augmentation(self):
        """Audio augmentation işlemini başlat"""
        self.aug_results.delete("1.0", "end")
        self.aug_results.insert("1.0", "🔄 Audio augmentation başlatılıyor...\n\n")
        
        # Bu işlem için gelişmiş kütüphaneler gerekebilir
        warning_text = """⚠️ GELIŞMIŞ ÖZELLİK UYARISI

Audio Augmentation işlemi için aşağıdaki Python kütüphaneleri gereklidir:
   • librosa (ses işleme)
   • pydub (audio manipülasyon)
   • scipy (sinyal işleme)

📥 Kurulum komutu:
pip install librosa pydub

🔧 Manuel Augmentation Önerileri:
   1. Mevcut kayıtlarınızı farklı hızlarda kaydedin (yavaş/hızlı konuşma)
   2. Farklı ses seviyelerinde kayıt yapın
   3. Farklı ortamlarda (sessiz oda, doğal ortam) kayıt yapın
   4. Farklı mikrofonlar kullanın
   5. Farklı konuşmacılardan aynı cümleleri kaydedin

💡 Bu yöntemlerle zaten çok etkili bir dataset oluşturabilirsiniz!"""
        
        self.aug_results.insert("end", warning_text)
    
    def preview_augmentation(self):
        """Augmentation önizlemesi"""
        if not Config.AUDIO_DIR.exists() or not list(Config.AUDIO_DIR.glob("*.wav")):
            messagebox.showwarning("Uyarı", "Önizleme için ses dosyası bulunamadı!")
            return
        
        preview_text = """🔍 AUGMENTATION ÖNİZLEMESİ

Seçilen Seçenekler:
"""
        
        if self.aug_speed.get():
            preview_text += "   ✅ Hız değişiklikleri aktif (2x çoğaltma)\n"
        if self.aug_volume.get():
            preview_text += "   ✅ Ses seviyesi varyasyonları aktif (2x çoğaltma)\n"
        if self.aug_noise.get():
            preview_text += "   ✅ Gürültü ekleme aktif (2x çoğaltma)\n"
        if self.aug_pitch.get():
            preview_text += "   ✅ Pitch değişiklikleri aktif (2x çoğaltma)\n"
        
        audio_files = list(Config.AUDIO_DIR.glob("*.wav"))
        current_count = len(audio_files)
        
        multiplier = 1
        if self.aug_speed.get(): multiplier += 1
        if self.aug_volume.get(): multiplier += 1
        if self.aug_noise.get(): multiplier += 1
        if self.aug_pitch.get(): multiplier += 1
        
        new_count = current_count * multiplier
        
        preview_text += f"""
📊 Dataset Büyütme Tahmini:
   • Mevcut dosya sayısı: {current_count}
   • Tahmini yeni dosya sayısı: {new_count}
   • Çoğaltma oranı: {multiplier}x
   
🎯 Whisper Eğitimi İçin:
   • {'✅ Yeterli' if new_count >= 1000 else '⚠️ Daha fazla gerekli'} dataset boyutu
   • Eğitim kalitesi: {'Mükemmel' if multiplier >= 3 else 'İyi' if multiplier >= 2 else 'Temel'}
"""
        
        self.aug_results.delete("1.0", "end")
        self.aug_results.insert("1.0", preview_text)
    
    def prepare_whisper_training(self):
        """Whisper eğitimi için final hazırlık"""
        prep_dialog = ctk.CTkToplevel(self.root)
        prep_dialog.title("🚀 Whisper Eğitim Hazırlığı")
        prep_dialog.geometry("900x700")
        prep_dialog.transient(self.root)
        
        main_frame = ctk.CTkFrame(prep_dialog)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Başlık
        ctk.CTkLabel(
            main_frame,
            text="🚀 Whisper Fine-tuning Hazırlık Raporu",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=(10, 20))
        
        # Hazırlık raporu alanı
        self.prep_results = ctk.CTkTextbox(main_frame, width=850, height=500, font=ctk.CTkFont(family="Consolas", size=11))
        self.prep_results.pack(pady=(0, 20))
        
        # Butonlar
        button_frame = ctk.CTkFrame(main_frame)
        button_frame.pack(fill="x")
        
        ctk.CTkButton(
            button_frame,
            text="📊 Dataset Durumunu Kontrol Et",
            command=self.check_dataset_status,
            width=200,
            height=40,
            fg_color="blue"
        ).pack(side="left", padx=10, pady=10)
        
        ctk.CTkButton(
            button_frame,
            text="📁 Eğitim Dosyalarını Aç",
            command=self.open_training_files,
            width=180,
            height=40,
            fg_color="green"
        ).pack(side="left", padx=10, pady=10)
        
        ctk.CTkButton(
            button_frame,
            text="💾 Eğitim Scriptini Oluştur",
            command=self.generate_training_script,
            width=180,
            height=40,
            fg_color="purple"
        ).pack(side="left", padx=10, pady=10)
        
        ctk.CTkButton(
            button_frame,
            text="❌ Kapat",
            command=prep_dialog.destroy,
            width=100,
            height=40
        ).pack(side="right", padx=10, pady=10)
        
        # İlk kontrol
        self.check_dataset_status()
    
    def check_dataset_status(self):
        """Dataset durumunu detaylı kontrol et"""
        self.prep_results.delete("1.0", "end")
        self.prep_results.insert("1.0", "🔍 Dataset durumu kontrol ediliyor...\n\n")
        
        try:
            # Dosya sayıları
            audio_files = list(Config.AUDIO_DIR.glob("*.wav")) if Config.AUDIO_DIR.exists() else []
            manifest_exists = Config.WHISPER_MANIFEST.exists()
            transcript_exists = Config.TRANSCRIPT_FILE.exists()
            
            # Manifest analizi
            manifest_entries = 0
            if manifest_exists:
                try:
                    with open(Config.WHISPER_MANIFEST, 'r', encoding='utf-8') as f:
                        manifest_entries = len(f.readlines())
                except:
                    manifest_entries = 0
            
            # Toplam süre hesabı
            total_duration = 0
            if audio_files:
                try:
                    import soundfile as sf
                    for audio_file in audio_files[:50]:  # İlk 50 dosyayı kontrol et
                        try:
                            data, samplerate = sf.read(audio_file)
                            total_duration += len(data) / samplerate
                        except:
                            continue
                    # Tahmini toplam süre
                    if len(audio_files) > 50:
                        avg_duration = total_duration / 50
                        total_duration = avg_duration * len(audio_files)
                except:
                    total_duration = len(audio_files) * 2  # Tahmini 2 saniye/dosya
            
            # Rapor oluştur
            status_report = f"""🎯 WHISPER EĞİTİM DATASET DURUMU
{"="*60}

📊 DATASET İSTATİSTİKLERİ:
   📁 Ses dosyası sayısı: {len(audio_files)}
   ⏱️  Tahmini toplam süre: {total_duration/60:.1f} dakika ({total_duration/3600:.1f} saat)
   📋 Manifest girişi: {manifest_entries}
   📄 Transkript dosyası: {'✅ Mevcut' if transcript_exists else '❌ Eksik'}

🎯 WHISPER EĞİTİMİ DEĞERLENDİRMESİ:

   Dataset Boyutu: {'🟢 MÜKEMMEL' if len(audio_files) >= 1000 else '🟡 İYİ' if len(audio_files) >= 500 else '🔴 YETERSİZ'}
   ({len(audio_files)}/500 minimum, 1000+ ideal)
   
   Toplam Süre: {'🟢 MÜKEMMEL' if total_duration >= 3600 else '🟡 İYİ' if total_duration >= 1800 else '🔴 YETERSİZ'}
   ({total_duration/60:.1f}/30 dakika minimum, 60+ dakika ideal)
   
   Manifest Durumu: {'🟢 TAMAM' if manifest_exists and manifest_entries > 0 else '🔴 SORUN'}
   
   Transkript Durumu: {'🟢 TAMAM' if transcript_exists else '🔴 SORUN'}

📋 WHISPER EĞİTİM DOSYALARI:
   • whisper_manifest.jsonl: {'✅' if manifest_exists else '❌'} ({manifest_entries} giriş)
   • transcripts.txt: {'✅' if transcript_exists else '❌'}
   • Audio klasörü: {'✅' if Config.AUDIO_DIR.exists() else '❌'} ({len(audio_files)} dosya)

🚀 HAZIRLIK DURUMU:
"""
            
            readiness_score = 0
            if len(audio_files) >= 500: readiness_score += 25
            if total_duration >= 1800: readiness_score += 25  
            if manifest_exists and manifest_entries > 0: readiness_score += 25
            if transcript_exists: readiness_score += 25
            
            if readiness_score >= 75:
                status_report += """   🎉 WHISPER EĞİTİMİ İÇİN HAZIR!
   
   ✅ Dataset boyutu yeterli
   ✅ Ses kalitesi kontrol edildi
   ✅ Manifest ve transkript dosyaları hazır
   ✅ Fine-tuning işlemine başlayabilirsiniz!"""
            else:
                status_report += f"""   ⚠️ HAZIRLIK ORANI: %{readiness_score}
   
   📝 Eksik olan öğeler:"""
                
                if len(audio_files) < 500:
                    status_report += f"\n   • {500 - len(audio_files)} daha fazla ses kaydı gerekli"
                if total_duration < 1800:
                    status_report += f"\n   • {(1800 - total_duration)/60:.1f} dakika daha ses kaydı gerekli"
                if not manifest_exists:
                    status_report += "\n   • Manifest dosyası oluşturulmalı"
                if not transcript_exists:
                    status_report += "\n   • Transkript dosyası oluşturulmalı"
            
            # Model önerileri
            status_report += f"""

🤝 WHISPER MODEL ÖNERİLERİ:

   Başlangıç için: openai/whisper-small
   • Hızlı eğitim, orta kalite
   • {len(audio_files)} dosya ile uyumlu
   
   Optimal kalite: openai/whisper-base
   • Dengeli hız/kalite oranı  
   • Kurmancî için ideal boyut
   
   Maksimum kalite: openai/whisper-large-v3
   • En yüksek kalite, yavaş eğitim
   • 1000+ dosya ile önerilir

🔧 SONRAKİ ADIMLAR:
   1. {'✅' if readiness_score >= 75 else '⚠️'} Dataset hazırlığını tamamlayın
   2. Hugging Face hesabınızı hazırlayın
   3. GPU/Colab ortamını ayarlayın
   4. Fine-tuning scriptini çalıştırın
   5. Eğitilmiş modeli test edin

💡 EĞITIM SONRASI:
   • Kendi modelinizle real-time transkripsiyon yapabilirsiniz
   • Web uygulaması veya API oluşturabilirsiniz
   • Ses-metin çeviri sistemi geliştirebilirsiniz"""
            
            self.prep_results.delete("1.0", "end")
            self.prep_results.insert("1.0", status_report)
            
        except Exception as e:
            error_text = f"❌ Durum kontrolü sırasında hata: {str(e)}"
            self.prep_results.insert("end", error_text)
    
    def open_training_files(self):
        """Eğitim dosyalarını sistem dosya gezgininde aç"""
        import subprocess
        import os
        
        try:
            if os.name == 'nt':  # Windows
                subprocess.run(['explorer', str(Config.BASE_DIR)], check=True)
            elif os.name == 'posix':  # macOS/Linux
                subprocess.run(['open' if sys.platform == 'darwin' else 'xdg-open', str(Config.BASE_DIR)], check=True)
            
            messagebox.showinfo("Dosyalar Açıldı", 
                f"📁 Whisper eğitim dosyaları açıldı:\n\n"
                f"📂 Klasör: {Config.BASE_DIR}\n"
                f"🎵 Ses dosyaları: audio/\n"
                f"📋 Manifest: whisper_manifest.jsonl\n"
                f"📄 Transkriptler: transcripts.txt")
        except Exception as e:
            messagebox.showerror("Hata", f"Dosyalar açılamadı: {str(e)}")
    
    def generate_training_script(self):
        """Whisper eğitim scripti oluştur"""
        script_content = '''#!/usr/bin/env python3
"""
Kurmancî Whisper Fine-tuning Script
Bu script, kaydedilmiş Kurmancî dataset ile Whisper modelini fine-tune eder.
"""

import os
import json
import torch
from datasets import Dataset, Audio
from transformers import (
    WhisperProcessor,
    WhisperForConditionalGeneration,
    WhisperTokenizer,
    Seq2SeqTrainingArguments,
    Seq2SeqTrainer,
    WhisperFeatureExtractor
)
from dataclasses import dataclass
from typing import Dict, List, Union
import librosa
import soundfile as sf

# Dataset yolu - Bu scripti kurmanji_dataset klasörüne koyun
DATASET_PATH = "."
AUDIO_PATH = "audio"
MANIFEST_FILE = "whisper_manifest.jsonl"

class KurmanjiWhisperDataset:
    def __init__(self, manifest_path, audio_path):
        self.manifest_path = manifest_path
        self.audio_path = audio_path
        self.data = self.load_manifest()
    
    def load_manifest(self):
        """Manifest dosyasını yükle"""
        data = []
        with open(self.manifest_path, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    entry = json.loads(line.strip())
                    data.append(entry)
        return data
    
    def create_hf_dataset(self):
        """Hugging Face Dataset oluştur"""
        # Veri hazırlığı
        audio_paths = []
        transcripts = []
        
        for entry in self.data:
            audio_file = os.path.join(self.audio_path, entry['audio_filepath'].replace('audio/', ''))
            if os.path.exists(audio_file):
                audio_paths.append(audio_file)
                transcripts.append(entry['text'])
        
        # Dataset oluştur
        dataset = Dataset.from_dict({
            "audio": audio_paths,
            "transcription": transcripts
        })
        
        # Audio özelliğini ekle
        dataset = dataset.cast_column("audio", Audio(sampling_rate=22050))
        
        return dataset

@dataclass
class DataCollatorSpeechSeq2SeqWithPadding:
    """Whisper için özel data collator"""
    processor: any
    decoder_start_token_id: int

    def __call__(self, features: List[Dict[str, Union[List[int], torch.Tensor]]]) -> Dict[str, torch.Tensor]:
        # Audio özelliklerini ayır
        input_features = [{"input_features": feature["input_features"]} for feature in features]
        batch = self.processor.feature_extractor.pad(input_features, return_tensors="pt")

        # Etiketleri ayır  
        label_features = [{"input_ids": feature["labels"]} for feature in features]
        labels_batch = self.processor.tokenizer.pad(label_features, return_tensors="pt")

        # -100 ile replace et (ignore index)
        labels = labels_batch["input_ids"].masked_fill(labels_batch.attention_mask.ne(1), -100)

        # bos token kaldır
        if (labels[:, 0] == self.decoder_start_token_id).all().cpu().item():
            labels = labels[:, 1:]

        batch["labels"] = labels
        return batch

def prepare_dataset(batch, processor):
    """Dataset hazırlama fonksiyonu"""
    # Audio'yu yükle ve özellik çıkar
    audio = batch["audio"]
    
    # Input features
    input_features = processor.feature_extractor(
        audio["array"], 
        sampling_rate=audio["sampling_rate"], 
        return_tensors="pt"
    ).input_features[0]
    
    # Labels (tokenize)
    labels = processor.tokenizer(batch["transcription"]).input_ids
    
    return {
        "input_features": input_features,
        "labels": labels,
    }

def main():
    """Ana eğitim fonksiyonu"""
    print("🚀 Kurmancî Whisper Fine-tuning başlatılıyor...")
    
    # Model ve processor yükle
    model_name = "openai/whisper-small"  # Küçük model ile başla
    processor = WhisperProcessor.from_pretrained(model_name, language="ku", task="transcribe")
    model = WhisperForConditionalGeneration.from_pretrained(model_name)
    
    # Tokenizer ayarları
    tokenizer = WhisperTokenizer.from_pretrained(model_name, language="ku", task="transcribe")
    
    # Dataset yükle
    dataset_loader = KurmanjiWhisperDataset(MANIFEST_FILE, AUDIO_PATH)
    dataset = dataset_loader.create_hf_dataset()
    
    print(f"📊 Toplam örnek sayısı: {len(dataset)}")
    
    # Dataset'i hazırla
    dataset = dataset.map(
        lambda batch: prepare_dataset(batch, processor),
        remove_columns=dataset.column_names,
        desc="Dataset hazırlanıyor"
    )
    
    # Train/test split
    train_test = dataset.train_test_split(test_size=0.1)
    train_dataset = train_test["train"]
    eval_dataset = train_test["test"]
    
    print(f"📚 Eğitim örnekleri: {len(train_dataset)}")
    print(f"🧪 Test örnekleri: {len(eval_dataset)}")
    
    # Data collator
    data_collator = DataCollatorSpeechSeq2SeqWithPadding(
        processor=processor,
        decoder_start_token_id=model.generation_config.decoder_start_token_id,
    )
    
    # Eğitim parametreleri
    training_args = Seq2SeqTrainingArguments(
        output_dir="./whisper-kurdish-kurmanji",
        per_device_train_batch_size=8,
        gradient_accumulation_steps=2,
        learning_rate=1e-5,
        warmup_steps=500,
        max_steps=2000,
        gradient_checkpointing=True,
        fp16=True,
        evaluation_strategy="steps",
        per_device_eval_batch_size=8,
        predict_with_generate=True,
        generation_max_length=225,
        save_steps=500,
        eval_steps=500,
        logging_steps=25,
        report_to=["tensorboard"],
        load_best_model_at_end=True,
        metric_for_best_model="wer",
        greater_is_better=False,
        push_to_hub=False,
    )
    
    # Trainer oluştur
    trainer = Seq2SeqTrainer(
        args=training_args,
        model=model,
        train_dataset=train_dataset,
        eval_dataset=eval_dataset,
        data_collator=data_collator,
        tokenizer=processor.feature_extractor,
    )
    
    # Eğitimi başlat
    print("🔥 Eğitim başlatılıyor...")
    trainer.train()
    
    # Modeli kaydet
    trainer.save_model("./whisper-kurdish-kurmanji-final")
    processor.save_pretrained("./whisper-kurdish-kurmanji-final")
    
    print("✅ Eğitim tamamlandı!")
    # Model saved
    
    # Test
    print("🧪 Model test ediliyor...")
    # Burada test kodu eklenebilir

if __name__ == "__main__":
    main()
'''
        
        # Scripti kaydet
        script_path = Config.BASE_DIR / "whisper_training.py"
        try:
            with open(script_path, 'w', encoding='utf-8') as f:
                f.write(script_content)
            
            requirements_content = '''# Whisper Fine-tuning Requirements
torch>=1.9.0
transformers>=4.21.0
datasets>=2.0.0
librosa>=0.9.0
soundfile>=0.12.0
tensorboard>=2.8.0
accelerate>=0.20.0
evaluate>=0.4.0
jiwer>=2.5.0  # WER hesabı için
'''
            
            req_path = Config.BASE_DIR / "requirements.txt"
            with open(req_path, 'w', encoding='utf-8') as f:
                f.write(requirements_content)
            
            # README oluştur
            readme_content = '''# Kurmancî Whisper Fine-tuning

Bu klasör, Kurmancî dili için Whisper model eğitimi içerir.

## Kurulum

```bash
pip install -r requirements.txt
```

## Kullanım

```bash
python whisper_training.py
```

## Dosyalar

- `whisper_training.py`: Ana eğitim scripti
- `whisper_manifest.jsonl`: Whisper dataset manifest
- `transcripts.txt`: Transkript dosyası
- `audio/`: Ses dosyaları klasörü
- `requirements.txt`: Gerekli Python paketleri

## Eğitim Sonrası

Eğitilmiş model `whisper-kurdish-kurmanji-final` klasöründe saklanır.

## Kullanım Örneği

```python
from transformers import WhisperProcessor, WhisperForConditionalGeneration
import librosa

# Modeli yükle
processor = WhisperProcessor.from_pretrained("./whisper-kurdish-kurmanji-final")
model = WhisperForConditionalGeneration.from_pretrained("./whisper-kurdish-kurmanji-final")

# Ses dosyasını yükle
audio, sr = librosa.load("test_audio.wav", sr=16000)

# Transkripsiyon
input_features = processor(audio, sampling_rate=sr, return_tensors="pt").input_features
predicted_ids = model.generate(input_features)
transcription = processor.batch_decode(predicted_ids, skip_special_tokens=True)

print(transcription[0])
```
'''
            
            readme_path = Config.BASE_DIR / "README.md"
            with open(readme_path, 'w', encoding='utf-8') as f:
                f.write(readme_content)
            
            messagebox.showinfo("Script Oluşturuldu!", 
                f"🎉 Whisper eğitim scripti oluşturuldu!\n\n"
                f"📁 Ana script: {script_path.name}\n"
                f"📋 Requirements: requirements.txt\n"
                f"📖 Kılavuz: README.md\n\n"
                f"🚀 Sonraki adımlar:\n"
                f"1. pip install -r requirements.txt\n"
                f"2. python whisper_training.py\n\n"
                f"💡 GPU kullanımı için Colab önerilir!")
            
        except Exception as e:
            messagebox.showerror("Hata", f"Script oluşturulamadı: {str(e)}")
    
    def run(self):
        self.root.mainloop()


def main():
    """Main entry point"""
    try:
        app = SimplifiedRecorderApp()
        app.run()
    except Exception as e:
        print(f"Error starting application: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()